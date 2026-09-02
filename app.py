# -*- coding: utf-8 -*-
"""
ORG_FULL_NAME · 局域网协同录入系统  主服务
================================================
架构：Flask + SQLite（单文件数据库，免安装服务）
部署：财务部电脑运行本程序即作为本地服务器，局域网内其他人员通过浏览器访问。

模块：
 1. 案件管理（自动案号、开发人/负责人/小组成员/客户/代理费/代垫诉讼费/开票金额）
 2. 薪酬计算（自定义公式 + 固定金额，按人员独立计算）
 3. 批量结算单生成（python-docx，支持自定义 Word 模板占位符）
 4. 文件导入导出（Excel 导入清洗 + 导出）
 5. 权限管理（多角色登录 + 基于角色/字段的访问控制）
 6. AI 接口集成（电子发票识别 + 费用科目自动归集，可配置任意兼容 OpenAI 格式的多模态 API）
协同：乐观锁版本号 + 字段级合并（冲突字段最后写入优先），全部操作写入审计日志。
"""

import ast
import csv
import io
import json
import os
import re
import secrets
import shutil
import sqlite3
import sys
import threading
import time
import webbrowser
import zipfile
from datetime import datetime, date, timedelta
from functools import wraps

from flask import (Flask, g, jsonify, request, send_file, session,
                   render_template)

# 多云 AI 发票识别适配器（密钥仅来自环境变量 / 本地配置文件，绝不硬编码、不入库明文）
import ai_providers

# ---------------------------------------------------------------- 基础配置
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# 品牌 / 机构配置（coop-base 通用分身：已去品牌化，请按需修改）
# 可通过环境变量覆盖：COOP_ORG_NAME / COOP_ORG_FULL / COOP_APP_DATA
# ============================================================
ORG_NAME = os.environ.get("COOP_ORG_NAME", "ORG")                  # 机构简称（默认销售方）
ORG_FULL_NAME = os.environ.get("COOP_ORG_FULL", "协同录入系统")       # 机构完整名称（如：XX 律师事务所），可用环境变量覆盖
APP_DATA_NAME = os.environ.get("COOP_APP_DATA", "协同录入系统")      # 应用 / 数据目录名（位于 %LOCALAPPDATA%）


def resource_path(rel):
    """资源定位：PyInstaller 打包后位于 sys._MEIPASS，开发态位于项目目录。"""
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, rel)
    return os.path.join(BASE_DIR, rel)


if getattr(sys, "frozen", False):
    # 打包发布态：数据写入用户可写目录，避免装在 Program Files 时无写权限
    _data = os.environ.get("COOP_DATA")
    if not _data:
        _data = os.path.join(os.path.expandvars("%LOCALAPPDATA%"), APP_DATA_NAME)
    BASE_DIR = _data
    DATA_DIR = _data
else:
    DATA_DIR = os.path.join(BASE_DIR, "data")
DB_PATH = os.path.join(DATA_DIR, "coop.db")
EXPORT_DIR = os.path.join(DATA_DIR, "exports")
SETTLE_DIR = os.path.join(DATA_DIR, "settlements")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
TEMPLATE_DIR = os.path.join(DATA_DIR, "templates")
for d in (DATA_DIR, EXPORT_DIR, SETTLE_DIR, UPLOAD_DIR, TEMPLATE_DIR):
    os.makedirs(d, exist_ok=True)

SERVER_PORT = int(os.environ.get("COOP_PORT", "8100"))  # 副本默认 8100，避免与原项目 8000 冲突

app = Flask(__name__, template_folder=resource_path("templates"))
app.secret_key = os.environ.get("COOP_SECRET", secrets.token_hex(32))
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB 上传上限
# 开发态模板即时重载：避免修改 index.html 后必须重启服务
app.config["TEMPLATES_AUTO_RELOAD"] = True

# ---------------------------------------------------------------- 角色与权限
# 角色定义：admin=财务经理（最高权限） finance=财务专员 case=案件专员 viewer=查看
#            institution_head=机构负责人（机构全局管理 + 审批权限）
#            lawyer=律师（案件基本信息录入与维护，可查看参与的案件）
ROLES = {
    "admin":   {"name": "财务经理", "desc": "最高权限：全部数据、用户、系统配置"},
    "finance": {"name": "财务专员", "desc": "案件财务字段、薪酬、结算单、发票识别"},
    "case":    {"name": "案件专员", "desc": "案件基本信息录入与维护"},
    "viewer":  {"name": "查看人员", "desc": "只读查看案件台账"},
    "institution_head": {"name": "机构负责人", "desc": "机构全局管理：查看全部数据，审批结案，管理本机构人员与角色"},
    "lawyer":  {"name": "律师", "desc": "案件基本信息录入与维护，可查看/编辑参与的案件"},
}

# 模块级权限矩阵
PERMS = {
    "case.view":      {"admin", "finance", "case", "viewer", "institution_head", "lawyer"},
    "case.edit":      {"admin", "finance", "case", "institution_head", "lawyer"},  # 基本字段
    "case.finance":   {"admin", "finance", "institution_head"},                    # 财务字段（代理费/代垫/开票）
    "case.stop":      {"admin", "finance", "institution_head"},                    # 停止案件（停用/暂停）
    "case.enable":    {"admin", "finance", "institution_head"},                    # 启用案件（恢复）
    "case.delete":    {"admin", "institution_head"},
    "salary.view":    {"admin", "finance", "institution_head"},
    "salary.edit":    {"admin", "finance", "institution_head"},
    "settle.view":    {"admin", "finance", "institution_head"},
    "settle.edit":    {"admin", "finance", "institution_head"},
    "file.import":    {"admin", "finance", "institution_head"},
    "file.export":    {"admin", "finance", "case", "institution_head", "lawyer"},
    "invoice.edit":   {"admin", "finance", "institution_head"},
    "user.manage":    {"admin", "institution_head"},
    "log.view":       {"admin", "institution_head"},
    "setting.manage": {"admin"},
}

# 权限中文标签（用于角色配置管理 UI）
PERM_LABELS = {
    "case.view":      "查看案件",
    "case.edit":      "编辑案件基本信息",
    "case.finance":   "编辑案件财务字段",
    "case.stop":      "停止案件",
    "case.enable":    "启用案件",
    "case.delete":    "删除案件",
    "salary.view":    "查看薪酬",
    "salary.edit":    "编辑薪酬",
    "settle.view":    "查看结算单",
    "settle.edit":    "生成/编辑结算单",
    "file.import":    "导入文件",
    "file.export":    "导出文件",
    "invoice.edit":   "发票识别与维护",
    "user.manage":    "用户与角色管理",
    "log.view":       "查看操作日志",
    "setting.manage": "系统设置",
}

# 角色配置持久化文件（启动时加载，运行时修改后保存）
ROLE_CONFIG_PATH = os.path.join(DATA_DIR, "role_config.json")


def _perms_to_dict():
    """将 PERMS 集合转成 JSON 可序列化的 dict。"""
    return {k: sorted(v) for k, v in PERMS.items()}


def load_role_config():
    """启动时从 data/role_config.json 加载角色-权限矩阵，覆盖内存中的 PERMS。
    配置文件不存在或格式错误时，保持代码内默认 PERMS 不变。"""
    if not os.path.isfile(ROLE_CONFIG_PATH):
        return
    try:
        with open(ROLE_CONFIG_PATH, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        perms_cfg = cfg.get("perms") if isinstance(cfg, dict) else None
        if not isinstance(perms_cfg, dict):
            return
        # 合并：仅覆盖已声明的权限项；角色名仅接受 ROLES 中已定义的有效值
        for k, roles in perms_cfg.items():
            if k in PERMS and isinstance(roles, list):
                PERMS[k] = {r for r in roles if r in ROLES}
    except Exception as e:
        print("[角色配置] 加载失败，保持默认:", e)


def save_role_config():
    """将当前 PERMS 持久化到 data/role_config.json。"""
    try:
        cfg = {
            "perms": _perms_to_dict(),
            "roles": {k: v["name"] for k, v in ROLES.items()},
            "updated_at": now(),
        }
        with open(ROLE_CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print("[角色配置] 保存失败:", e)
        return False

# 案件表字段分组（字段级权限）
CASE_BASE_FIELDS = ["developer", "owner", "members", "client_name",
                    "case_type", "status", "remark"]
CASE_FIN_FIELDS = ["agent_fee", "court_fee_advance", "invoice_amount"]
CASE_EXT_BASE_FIELDS = ["plaintiff", "defendant", "section", "case_cause",
                        "project", "fee_standard", "maintain_user",
                        "archive_status", "archive_borrow", "borrow_user",
                        "borrow_time", "borrow_return_time",
                        "recover_month", "settle_month"]
CASE_EXT_FIN_FIELDS = ["subject_amount", "expected_profit", "confirmed_amount"]
# 财务文本字段（需 case.finance 权限，但不做数字强转）：预计收益文本说明 / 财务备注
CASE_EXT_FIN_TEXT_FIELDS = ["expected_profit_text", "finance_remark"]

# 实体字段元数据：驱动「导入模板字段映射」「批量导入校验」「前端表单」三处。
# type: str=文本 num=数字 date=日期
ENTITY_DEFS = {
    "case": {
        "name": "案件台账",
        "fields": [
            {"key": "case_no", "label": "案号", "required": False, "type": "str"},
            {"key": "client_name", "label": "客户名称", "required": True, "type": "str"},
            {"key": "developer", "label": "案件开发人", "required": False, "type": "str"},
            {"key": "owner", "label": "负责人", "required": False, "type": "str"},
            {"key": "members", "label": "小组成员", "required": False, "type": "str"},
            {"key": "case_type", "label": "案件类型", "required": False, "type": "str"},
            {"key": "status", "label": "状态", "required": False, "type": "str"},
            {"key": "agent_fee", "label": "代理费金额", "required": False, "type": "num"},
            {"key": "court_fee_advance", "label": "代垫诉讼费", "required": False, "type": "num"},
            {"key": "invoice_amount", "label": "开票金额", "required": False, "type": "num"},
            {"key": "remark", "label": "备注", "required": False, "type": "str"},
        ],
    },
    "staff": {
        "name": "人员花名册",
        "fields": [
            {"key": "staff_no", "label": "工号", "required": False, "type": "str"},
            {"key": "name", "label": "姓名", "required": True, "type": "str"},
            {"key": "department", "label": "部门", "required": False, "type": "str"},
            {"key": "position", "label": "岗位", "required": False, "type": "str"},
            {"key": "hire_date", "label": "入职日期", "required": False, "type": "date"},
            {"key": "phone", "label": "联系电话", "required": False, "type": "str"},
            {"key": "id_card", "label": "身份证号", "required": False, "type": "str"},
            {"key": "status", "label": "在职状态", "required": False, "type": "str"},
            {"key": "commission_rate", "label": "案件提成率", "required": False, "type": "num"},
            {"key": "remark", "label": "备注", "required": False, "type": "str"},
        ],
    },
    "customer": {
        "name": "客户信息",
        "fields": [
            {"key": "name", "label": "客户名称", "required": True, "type": "str"},
            {"key": "credit_code", "label": "统一社会信用代码", "required": False, "type": "str"},
            {"key": "contact", "label": "联系人", "required": False, "type": "str"},
            {"key": "phone", "label": "联系电话", "required": False, "type": "str"},
            {"key": "cust_type", "label": "客户类型", "required": False, "type": "str"},
            {"key": "industry", "label": "所属行业", "required": False, "type": "str"},
            {"key": "address", "label": "地址", "required": False, "type": "str"},
            {"key": "remark", "label": "备注", "required": False, "type": "str"},
        ],
    },
}

# ---------------------------------------------------------------- 案件状态与执行状态枚举
# 2026-08-31 新增：支持合伙人运营看板按状态分列、流转校验、操作日志。
# 编码采用英文 code，前端/后端统一；中文 label 仅用于展示；
# 历史数据中的中文状态通过 LEGACY_STATUS_MAP 无损映射。

CASE_STATUS_META = [
    # 正常流转（蓝色系）
    {"code": "pending_contact",   "label": "待接洽",       "color": "#3B82F6", "category": "flow",      "next": ["consulting"]},
    {"code": "consulting",        "label": "咨询评估中",   "color": "#3B82F6", "category": "flow",      "next": ["consulted", "rejected_conflict"]},
    {"code": "consulted",         "label": "已咨询待签约", "color": "#3B82F6", "category": "flow",      "next": ["contracting", "rejected_fee", "rejected_conflict"]},
    {"code": "contracting",       "label": "已签约待收案", "color": "#3B82F6", "category": "flow",      "next": ["filing_prepare", "mediation_closed", "withdrawn"]},
    {"code": "filing_prepare",    "label": "立案准备中",   "color": "#3B82F6", "category": "flow",      "next": ["first_instance", "mediation_closed", "withdrawn"]},
    {"code": "first_instance",    "label": "一审",         "color": "#3B82F6", "category": "flow",      "next": ["second_instance", "execution", "closed", "mediation_closed", "withdrawn", "dismissed"]},
    {"code": "second_instance",   "label": "二审",         "color": "#3B82F6", "category": "flow",      "next": ["execution", "closed", "mediation_closed", "withdrawn", "dismissed"]},
    {"code": "execution",         "label": "执行",         "color": "#F59E0B", "category": "flow",      "next": ["closed", "execution_ended"]},
    # 结果类（绿色系）
    {"code": "closed",            "label": "结案",         "color": "#10B981", "category": "result",    "next": ["archived"]},
    {"code": "archived",          "label": "归档",         "color": "#10B981", "category": "result",    "next": []},
    {"code": "mediation_closed",  "label": "调解结案",     "color": "#10B981", "category": "result",    "next": ["archived"]},
    # 终止类（灰色系）
    {"code": "withdrawn",         "label": "撤诉",         "color": "#9CA3AF", "category": "terminated", "next": ["archived"]},
    {"code": "dismissed",         "label": "驳回",         "color": "#9CA3AF", "category": "terminated", "next": ["archived"]},
    {"code": "terminated",        "label": "终止",         "color": "#9CA3AF", "category": "terminated", "next": ["archived"]},
    # 异常类（红色系）
    {"code": "rejected_conflict", "label": "拒接（利益冲突）", "color": "#EF4444", "category": "abnormal",  "next": []},
    {"code": "rejected_fee",      "label": "拒接（费用未达成）", "color": "#EF4444", "category": "abnormal",  "next": []},
]

EXECUTION_STATUS_META = [
    {"code": "not_applied",     "label": "未申请执行", "color": "#3B82F6", "next": ["filing"]},
    {"code": "filing",          "label": "立案中",     "color": "#3B82F6", "next": ["executing"]},
    {"code": "executing",       "label": "执行中",     "color": "#F59E0B", "next": ["recovered", "execution_ended"]},
    {"code": "recovered",       "label": "已执行到位", "color": "#10B981", "next": ["closed"]},
    {"code": "execution_ended", "label": "终本",       "color": "#9CA3AF", "next": ["resumed"]},
    {"code": "resumed",         "label": "恢复执行",   "color": "#3B82F6", "next": ["executing"]},
]

# 历史中文状态 → 新 code 的映射（无损兼容）
LEGACY_STATUS_MAP = {
    "进行中": "first_instance",
    "已结案": "closed",
    "暂停":   "terminated",
    "":       "pending_contact",
}

CASE_STATUS_CODES = {s["code"] for s in CASE_STATUS_META}
EXECUTION_STATUS_CODES = {s["code"] for s in EXECUTION_STATUS_META}
CASE_STATUS_BY_CODE = {s["code"]: s for s in CASE_STATUS_META}
EXECUTION_STATUS_BY_CODE = {s["code"]: s for s in EXECUTION_STATUS_META}
CASE_STATUS_LABEL_TO_CODE = {s["label"]: s["code"] for s in CASE_STATUS_META}


def norm_case_status(raw):
    """把数据库里可能的中文/英文/空状态统一归一化为新 code。
    支持：旧中文状态、新英文 code、新中文 label。"""
    raw = (raw or "").strip()
    if raw in LEGACY_STATUS_MAP:
        return LEGACY_STATUS_MAP[raw]
    if raw in CASE_STATUS_CODES:
        return raw
    if raw in CASE_STATUS_LABEL_TO_CODE:
        return CASE_STATUS_LABEL_TO_CODE[raw]
    return "pending_contact"


def status_label(code):
    """code → 中文展示 label（兼容未知 code 原样返回）。"""
    return CASE_STATUS_BY_CODE.get(code, {}).get("label", code)


def is_flow_status(code):
    """是否仍在正常推进中（未结项、未终止、未异常）。"""
    meta = CASE_STATUS_BY_CODE.get(code)
    return meta is not None and meta["category"] == "flow"


# 为导入/旧表单提供允许的字符串集合（新旧均可）
VALID_STATUS_STRINGS = set(CASE_STATUS_CODES) | set(LEGACY_STATUS_MAP.keys())


def entity_fields(entity):
    """返回实体字段定义列表，case 兼容旧 IMPORT_MAP"""
    return ENTITY_DEFS.get(entity, {}).get("fields", [])


def role():
    return session.get("role")


def perms_of_role(role_key):
    """返回某角色默认拥有的全部权限键列表。"""
    return [p for p in PERMS if role_key in PERMS.get(p, set())]


def validate_perms(keys):
    """校验并归一化前端传入的权限键列表。

    返回 None 表示非法输入（非列表 / 含未知权限键）；返回列表（可能为空）表示合法。
    空列表表示「清除用户级覆盖、继承角色默认权限」。
    """
    if keys is None:
        return None
    if not isinstance(keys, list):
        return None
    out = []
    for k in keys:
        if k not in PERMS:
            return None  # 含未知权限键，整体拒绝，由调用方返回 400
        if out.count(k) == 0:
            out.append(k)
    return out


def resolve_perms(u):
    """解析用户最终权限集合：优先用户级覆盖，否则继承角色默认权限。

    u 可为 dict 或 sqlite3.Row（均支持下标访问；注意 sqlite3.Row 不支持 `in` 判断列名）。
    """
    try:
        raw = u["perms"]
    except (KeyError, IndexError):
        raw = None
    if raw:
        try:
            s = set(json.loads(raw)) & set(PERMS.keys())
            if s:
                return list(s)
        except Exception:
            pass
    return perms_of_role(u["role"])


def has_perm(p):
    perms = session.get("perms")
    if perms is not None:
        return p in perms
    return role() in PERMS.get(p, set())


def login_required(f):
    @wraps(f)
    def wrapper(*a, **kw):
        if not session.get("uid"):
            return jsonify({"ok": False, "msg": "未登录或会话已过期"}), 401
        return f(*a, **kw)
    return wrapper


def perm_required(p):
    def deco(f):
        @wraps(f)
        def wrapper(*a, **kw):
            if not session.get("uid"):
                return jsonify({"ok": False, "msg": "未登录"}), 401
            if not has_perm(p):
                return jsonify({"ok": False, "msg": "无此操作权限"}), 403
            return f(*a, **kw)
        return wrapper
    return deco


# ---------------------------------------------------------------- 数据库
def get_db():
    if "db" not in g:
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA journal_mode=WAL")       # 支持并发读写
        g.db.execute("PRAGMA busy_timeout=5000")      # 写冲突时等待而非立即失败
        g.db.execute("PRAGMA foreign_keys=ON")
    return g.db


@app.teardown_appcontext
def close_db(_=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


SCHEMA = """
CREATE TABLE IF NOT EXISTS users (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  username TEXT UNIQUE NOT NULL,
  password_hash TEXT NOT NULL,
  display_name TEXT NOT NULL,
  role TEXT NOT NULL DEFAULT 'viewer',
  perms TEXT DEFAULT NULL,              -- 用户级权限覆盖（JSON 数组）；为空则继承所选角色默认权限
  active INTEGER NOT NULL DEFAULT 1,
  created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS cases (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_no TEXT UNIQUE NOT NULL,          -- 系统自动生成案号（如 #260901）
  developer TEXT DEFAULT '',             -- 案件开发人
  owner TEXT DEFAULT '',                 -- 负责人
  members TEXT DEFAULT '',               -- 小组成员（逗号分隔）
  client_name TEXT NOT NULL,             -- 客户名称
  case_type TEXT DEFAULT '',             -- 案件类型
  status TEXT DEFAULT '进行中',          -- 案件主状态（兼容旧中文 + 新英文 code）
  execution_status TEXT DEFAULT '',      -- 执行状态 code（仅 case_status=execution 时有效）
  agent_fee REAL DEFAULT 0,              -- 代理费金额
  court_fee_advance REAL DEFAULT 0,      -- 代垫诉讼费
  invoice_amount REAL DEFAULT 0,         -- 财务开票金额
  plaintiff TEXT DEFAULT '',             -- 原告
  defendant TEXT DEFAULT '',           -- 被告
  section TEXT DEFAULT '',               -- 部门
  case_cause TEXT DEFAULT '',            -- 案由
  subject_amount REAL DEFAULT 0,         -- 标的
  project TEXT DEFAULT '',               -- 项目
  fee_standard TEXT DEFAULT '',          -- 收费标准
  expected_profit REAL DEFAULT 0,        -- 预期收益
  maintain_user TEXT DEFAULT '',         -- 维护人
  archive_status TEXT DEFAULT '未归档',  -- 档案是否已归档：已归档/未归档
  archive_borrow TEXT DEFAULT '',        -- 原件调档借阅情况
  borrow_user TEXT DEFAULT '',           -- 借阅人
  borrow_time TEXT DEFAULT '',           -- 借阅时间
  borrow_return_time TEXT DEFAULT '',    -- 预计归还时间
  recover_month TEXT DEFAULT '',         -- 财务回款月（YYYY-MM）
  settle_month TEXT DEFAULT '',          -- 提成结算月（YYYY-MM）
  confirmed_amount REAL DEFAULT 0,       -- 财务确认回款金额（到账金额）
  settle_ready INTEGER DEFAULT 0,        -- 是否满足提成结算条件（0/1）
  remark TEXT DEFAULT '',
  version INTEGER NOT NULL DEFAULT 1,    -- 乐观锁版本号
  created_by TEXT, created_at TEXT,
  updated_by TEXT, updated_at TEXT,
  deleted INTEGER NOT NULL DEFAULT 0
);
CREATE TABLE IF NOT EXISTS fee_recoveries (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_id INTEGER NOT NULL,
  recovery_date TEXT DEFAULT '',         -- 回款/费用日期 YYYY-MM-DD
  item TEXT DEFAULT '',                  -- 款项说明
  amount REAL DEFAULT 0,                 -- 金额（正=收入，负=支出/垫付）
  confirm INTEGER DEFAULT 0,
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS case_commissions (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_id INTEGER NOT NULL,
  person TEXT NOT NULL,                  -- 提成人员姓名
  role TEXT DEFAULT '',                  -- 角色：承办人/维护人/其他
  rate REAL DEFAULT 0,                   -- 提成比例（如 0.05）
  base_amount REAL DEFAULT 0,            -- 提成基数
  amount REAL DEFAULT 0,                 -- 提成金额（四舍五入取整后）
  raw_amount REAL DEFAULT 0,             -- 取整前精确金额（用于复核）
  remark TEXT DEFAULT '',                -- 备注，如「（取整）」
  calc_process TEXT DEFAULT '',          -- 完整计算过程：基数 × 比例 = 精确值 → 取整 金额
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS case_members (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_id INTEGER NOT NULL,              -- 案件 ID
  person TEXT NOT NULL,                  -- 成员姓名
  role TEXT DEFAULT '',                  -- 角色：维护人/主承办人/协办人
  rate REAL DEFAULT 0,                   -- 预计提成比例（默认从花名册引入，可手动修改）
  guaranteed_amount REAL DEFAULT 600,    -- 保底金额（维护人默认 0，其他默认 600，可修改）
  fixed_amount REAL DEFAULT 0,           -- 固定金额
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS salary_adjustments (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_id INTEGER NOT NULL,
  person TEXT NOT NULL,
  month TEXT DEFAULT '',                 -- 结算月 YYYY-MM
  amount REAL DEFAULT 0,                 -- 手动调整金额（覆盖自动计算）
  created_by TEXT, created_at TEXT,
  UNIQUE(case_id, person, month)
);
CREATE TABLE IF NOT EXISTS salary_rules (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  user_name TEXT NOT NULL,               -- 适用人员（与案件 开发人/负责人/成员 对应）
  mode TEXT NOT NULL DEFAULT 'formula',  -- formula=自定义公式 / fixed=固定金额
  formula TEXT DEFAULT '',               -- 如 agent_fee*0.1 + invoice_amount*0.02
  fixed_amount REAL DEFAULT 0,
  effective_from TEXT DEFAULT '',        -- 生效日期 YYYY-MM-DD，空=立即
  note TEXT DEFAULT '',
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS settlements (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  batch_no TEXT NOT NULL,                -- 批次号
  case_id INTEGER NOT NULL,
  case_no TEXT, user_name TEXT,
  amount REAL DEFAULT 0,
  doc_path TEXT DEFAULT '',
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS invoices (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  case_id INTEGER,
  case_no TEXT DEFAULT '',               -- 案件号（多对一发票索引）
  invoice_no TEXT DEFAULT '',
  client_name TEXT DEFAULT '',           -- 客户名称（由发票「购买方」校正而来）
  seller TEXT DEFAULT '',                -- 销售方（销售方台账；默认取「ORG_NAME」）
  seller_raw TEXT DEFAULT '',            -- 销售方原始识别值（仅追溯，不被覆盖）
  buyer_raw TEXT DEFAULT '',             -- 购买方原始识别值（仅追溯，不被覆盖）
  review_flag INTEGER DEFAULT 0,         -- 1=待人工核对（销售方与默认值不一致等识别异常）
  review_note TEXT DEFAULT '',           -- 待核对原因说明
  amount REAL DEFAULT 0,
  tax_amount REAL DEFAULT 0,
  category TEXT DEFAULT '',              -- 费用科目（优先取发票「项目名称」去编码前缀，无则关键词归集）
  invoice_date TEXT DEFAULT '',
  invoice_type TEXT DEFAULT '',          -- 发票类型（专票/普票，由票面识别）
  item_name TEXT DEFAULT '',             -- 发票项目名称（原始，含税收分类编码前缀）
  file_path TEXT DEFAULT '',
  ai_raw TEXT DEFAULT '',                -- AI 原始返回
  confirm INTEGER DEFAULT 0,             -- 人工确认
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS audit_logs (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  user_id INTEGER, username TEXT,
  action TEXT NOT NULL,                  -- create/update/delete/query/login/export/import...
  object_type TEXT NOT NULL,             -- case/user/salary_rule/settlement/invoice/system
  object_id TEXT DEFAULT '',
  detail TEXT DEFAULT '',
  ip TEXT DEFAULT '',
  created_at TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS settings (
  key TEXT PRIMARY KEY,
  value TEXT DEFAULT ''
);
CREATE TABLE IF NOT EXISTS case_no_seq (
  ym TEXT PRIMARY KEY,                -- 案号序号计数键：年两位+月两位，如 2409
  seq INTEGER NOT NULL DEFAULT 0      -- 当前年月已使用的最大顺序号（按年月重置）
);
CREATE TABLE IF NOT EXISTS settle_templates (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  name TEXT NOT NULL,               -- 模板名称
  is_default INTEGER NOT NULL DEFAULT 0,  -- 是否默认结算模板
  config TEXT NOT NULL DEFAULT '{}',-- JSON：title/fields/show_* /footer
  created_by TEXT, created_at TEXT,
  updated_by TEXT, updated_at TEXT
);
CREATE TABLE IF NOT EXISTS import_templates (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  name TEXT NOT NULL,               -- 模板名称
  entity TEXT NOT NULL DEFAULT 'case',  -- 目标实体：case/staff/customer
  mapping TEXT NOT NULL DEFAULT '{}',   -- JSON：{源列名: 目标字段}
  created_by TEXT, created_at TEXT
);
CREATE TABLE IF NOT EXISTS staff (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  staff_no TEXT DEFAULT '',         -- 工号
  name TEXT NOT NULL,               -- 姓名
  department TEXT DEFAULT '',       -- 部门
  position TEXT DEFAULT '',         -- 岗位
  hire_date TEXT DEFAULT '',        -- 入职日期 YYYY-MM-DD
  phone TEXT DEFAULT '',            -- 联系电话
  id_card TEXT DEFAULT '',          -- 身份证号
  status TEXT DEFAULT '在职',
  commission_rate REAL DEFAULT 0,    -- 案件提成率（如 0.08，薪酬计算自动引用）
  remark TEXT DEFAULT '',
  created_by TEXT, created_at TEXT,
  updated_by TEXT, updated_at TEXT,
  deleted INTEGER NOT NULL DEFAULT 0
);
CREATE TABLE IF NOT EXISTS customers (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  name TEXT NOT NULL,               -- 客户名称
  credit_code TEXT DEFAULT '',      -- 统一社会信用代码
  contact TEXT DEFAULT '',          -- 联系人
  phone TEXT DEFAULT '',            -- 联系电话
  cust_type TEXT DEFAULT '',        -- 客户类型
  industry TEXT DEFAULT '',         -- 所属行业
  address TEXT DEFAULT '',          -- 地址
  remark TEXT DEFAULT '',
  created_by TEXT, created_at TEXT,
  updated_by TEXT, updated_at TEXT,
  deleted INTEGER NOT NULL DEFAULT 0
);
CREATE INDEX IF NOT EXISTS idx_cases_client ON cases(client_name);
CREATE INDEX IF NOT EXISTS idx_audit_time ON audit_logs(created_at);
CREATE INDEX IF NOT EXISTS idx_staff_name ON staff(name);
CREATE INDEX IF NOT EXISTS idx_customers_name ON customers(name);
"""


def init_db():
    db = sqlite3.connect(DB_PATH)
    db.executescript(SCHEMA)
    # 迁移：为已存在的库补充 users.perms 列（幂等，列已存在则忽略）
    try:
        db.execute("ALTER TABLE users ADD COLUMN perms TEXT DEFAULT NULL")
    except sqlite3.OperationalError:
        pass
    db.commit()
    # 默认管理员：admin / 123456（首次登录后请修改）
    if not db.execute("SELECT 1 FROM users WHERE username='admin'").fetchone():
        db.execute(
            "INSERT INTO users(username,password_hash,display_name,role,created_at)"
            " VALUES(?,?,?,?,?)",
            ("admin", hash_pw("123456"), "财务经理", "admin", now()))
    defaults = {
        "pair_code": "%06d" % secrets.randbelow(1000000),  # 局域网配网数字
        "firm_name": ORG_FULL_NAME,
        # 激活的 AI 识别服务商（火山引擎/阿里百炼/腾讯云/兼容OpenAI），密钥在本地配置文件或环境变量中
        "ai_provider": "bailian",
        # 以下为兼容旧版单端点配置的兜底值（非密钥项）；新版按服务商在本地配置文件中分别设置
        "ai_base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        # 多模态视觉模型（兼容 OpenAI 模式兜底），阿里百炼原生接口默认用 qwen-vl-max-latest
        "ai_model": "qwen-vl-max-latest",
        "fee_categories": json.dumps(
            ["诉讼费", "律师代理费", "差旅费", "公证费", "鉴定费", "办公费", "其他"],
            ensure_ascii=False),
        "settle_template": "",  # 自定义结算单模板文件名（置于 data/templates/）
        "case_no_seq_start": "1",  # 案号起始序号（当前年月首次生成案号时生效，支持调整）
        # 民事案件类型下拉选项（可在系统设置中扩展，JSON 数组）
        "case_types": json.dumps([
            "婚姻家庭纠纷", "继承纠纷", "合同纠纷", "侵权责任纠纷", "劳动争议",
            "人事争议", "物权纠纷", "担保物权纠纷", "知识产权纠纷", "不正当竞争纠纷",
            "公司有关纠纷", "合伙企业纠纷", "破产纠纷", "建设工程合同纠纷",
            "房屋买卖合同纠纷", "房屋租赁合同纠纷", "物业服务合同纠纷", "民间借贷纠纷",
            "保证合同纠纷", "其他民事纠纷"], ensure_ascii=False),
        # 案件类型「板块」二级选项（JSON 对象：类型→板块列表）。
        # 建设工程合同纠纷 / 买卖合同纠纷 预置固定 20 项；其余类型为空（支持用户自定义录入）。
        "case_type_sections": json.dumps({
            "建设工程合同纠纷": ["木方", "钢材", "钢模", "工伤人损", "混凝土",
                              "混凝土（普混）", "混凝土（透水）", "建设工程", "脚手架",
                              "劳动争议", "劳务", "铝模", "民间借贷", "票据纠纷", "其他",
                              "其他建材", "砌体", "砂浆", "塔吊", "外加剂"],
            "买卖合同纠纷": ["木方", "钢材", "钢模", "工伤人损", "混凝土",
                          "混凝土（普混）", "混凝土（透水）", "建设工程", "脚手架",
                          "劳动争议", "劳务", "铝模", "民间借贷", "票据纠纷", "其他",
                          "其他建材", "砌体", "砂浆", "塔吊", "外加剂"],
        }, ensure_ascii=False),
    }
    for k, v in defaults.items():
        db.execute("INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)", (k, v))
    db.commit()
    db.close()
    migrate_db()   # 为旧库补齐新增字段


def migrate_db():
    """轻量迁移：为已存在的旧库补齐新增字段（不会删除任何已有数据）"""
    db = sqlite3.connect(DB_PATH)
    try:
        # cases 表：结案审批表所需扩展字段
        cols = {r[1] for r in db.execute("PRAGMA table_info(cases)")}
        new_cols = [
            ("plaintiff", "TEXT DEFAULT ''"),
            ("defendant", "TEXT DEFAULT ''"),
            ("section", "TEXT DEFAULT ''"),
            ("case_cause", "TEXT DEFAULT ''"),
            ("subject_amount", "REAL DEFAULT 0"),
            ("project", "TEXT DEFAULT ''"),
            ("fee_standard", "TEXT DEFAULT ''"),
            ("expected_profit", "REAL DEFAULT 0"),
            ("maintain_user", "TEXT DEFAULT ''"),
            ("archive_status", "TEXT DEFAULT '未归档'"),
            ("archive_borrow", "TEXT DEFAULT ''"),
            ("borrow_user", "TEXT DEFAULT ''"),
            ("borrow_time", "TEXT DEFAULT ''"),
            ("borrow_return_time", "TEXT DEFAULT ''"),
            ("recover_month", "TEXT DEFAULT ''"),
            ("settle_month", "TEXT DEFAULT ''"),
            ("confirmed_amount", "REAL DEFAULT 0"),
            ("settle_ready", "INTEGER DEFAULT 0"),
            ("expected_profit_text", "TEXT DEFAULT ''"),  # 财务信息：预计收益金额（文本说明）
            ("finance_remark", "TEXT DEFAULT ''"),         # 财务信息：备注内容
            ("execution_status", "TEXT DEFAULT ''"),   # 2026-08-31 执行状态 code
            ("active", "INTEGER NOT NULL DEFAULT 1"),   # 案件启用/停止（1=启用，0=停止）
        ]
        for name, ddl in new_cols:
            if name not in cols:
                db.execute("ALTER TABLE cases ADD COLUMN %s %s" % (name, ddl))
        # staff 表：案件提成率
        scols = {r[1] for r in db.execute("PRAGMA table_info(staff)")}
        if "commission_rate" not in scols:
            db.execute("ALTER TABLE staff ADD COLUMN commission_rate REAL DEFAULT 0")
        # case_commissions 表：取整相关字段
        ccols = {r[1] for r in db.execute("PRAGMA table_info(case_commissions)")}
        new_ccols = [
            ("raw_amount", "REAL DEFAULT 0"),
            ("remark", "TEXT DEFAULT ''"),
        ]
        for name, ddl in new_ccols:
            if name not in ccols:
                db.execute("ALTER TABLE case_commissions ADD COLUMN %s %s" % (name, ddl))
        # case_members 表：保底金额 / 固定金额
        mcols = {r[1] for r in db.execute("PRAGMA table_info(case_members)")}
        if "guaranteed_amount" not in mcols:
            db.execute("ALTER TABLE case_members ADD COLUMN guaranteed_amount REAL DEFAULT 600")
        if "fixed_amount" not in mcols:
            db.execute("ALTER TABLE case_members ADD COLUMN fixed_amount REAL DEFAULT 0")
            # 旧数据按角色回填默认保底：维护人 0，其他 600
        try:
            db.execute("UPDATE case_members SET role='maintainer' WHERE role='developer'")
            db.execute("UPDATE case_members SET guaranteed_amount=0 WHERE LOWER(role) LIKE 'maintainer'")
            db.execute("UPDATE case_members SET guaranteed_amount=600 WHERE LOWER(role) NOT LIKE 'maintainer' AND (guaranteed_amount IS NULL OR guaranteed_amount=0 AND role NOT LIKE 'maintainer')")
        except Exception:
            pass
        # invoices 表：案件号 + 客户名称/销售方台账/追溯字段 + 发票类型/项目名称
        icols = {r[1] for r in db.execute("PRAGMA table_info(invoices)")}
        if "case_no" not in icols:
            db.execute("ALTER TABLE invoices ADD COLUMN case_no TEXT DEFAULT ''")
        for name, ddl in [
            ("client_name", "TEXT DEFAULT ''"),   # 客户名称（由发票「购买方」校正而来）
            ("seller_raw", "TEXT DEFAULT ''"),    # 销售方原始识别值（仅追溯）
            ("buyer_raw", "TEXT DEFAULT ''"),     # 购买方原始识别值（仅追溯）
            ("review_flag", "INTEGER DEFAULT 0"), # 1=待人工核对
            ("review_note", "TEXT DEFAULT ''"),   # 待核对原因
            ("invoice_type", "TEXT DEFAULT ''"),  # 发票类型（专票/普票，由票面识别）
            ("item_name", "TEXT DEFAULT ''"),     # 发票项目名称（原始，含税收分类编码前缀）
            ("total_amount", "REAL DEFAULT 0"),   # 价税合计（票面小写金额）
            ("total_amount_cn", "TEXT DEFAULT ''"),  # 价税合计大写
            ("tax_rate", "TEXT DEFAULT ''"),      # 税率/征收率（如 6%、13%、1%、免税）
            ("buyer_tax_id", "TEXT DEFAULT ''"),  # 购买方统一社会信用代码/纳税人识别号
            ("seller_tax_id", "TEXT DEFAULT ''"), # 销售方统一社会信用代码/纳税人识别号
            ("drawer", "TEXT DEFAULT ''"),        # 开票人
            ("remark", "TEXT DEFAULT ''"),        # 发票备注栏
        ]:
            if name not in icols:
                db.execute("ALTER TABLE invoices ADD COLUMN %s %s" % (name, ddl))
        # 旧数据校正：历史 seller 字段实为「购买方」展示值 → 校正为客户名称；
        # 销售方统一记为默认「ORG_NAME」，原始值保留到 seller_raw 以便追溯（不静默覆盖）
        try:
            db.execute(
                "UPDATE invoices SET client_name=seller, seller_raw=seller, "
                "buyer_raw=seller, seller=ORG_NAME, review_flag=0, review_note='' "
                "WHERE (client_name IS NULL OR client_name='') AND seller<>''")
        except Exception:
            pass
        # 案号序号计数器表（按年月重置，并发安全）
        db.execute("""CREATE TABLE IF NOT EXISTS case_no_seq (
            ym TEXT PRIMARY KEY,
            seq INTEGER NOT NULL DEFAULT 0
        )""")
        # 发票金额口径迁移（规范对齐）：历史 amount 实为「价税合计」→ 迁入 total_amount；
        # 新数据由 AI 识别按规范入库（amount=不含税、total_amount=价税合计），不受影响
        try:
            db.execute(
                "UPDATE invoices SET total_amount=amount "
                "WHERE (total_amount IS NULL OR total_amount=0) AND amount>0")
        except Exception:
            pass
        # 案件类型「板块」二级选项：为已存在的库补齐预置（新库由 init_db 写入）
        _sect_preset = json.dumps({
            "建设工程合同纠纷": ["木方", "钢材", "钢模", "工伤人损", "混凝土",
                              "混凝土（普混）", "混凝土（透水）", "建设工程", "脚手架",
                              "劳动争议", "劳务", "铝模", "民间借贷", "票据纠纷", "其他",
                              "其他建材", "砌体", "砂浆", "塔吊", "外加剂"],
            "买卖合同纠纷": ["木方", "钢材", "钢模", "工伤人损", "混凝土",
                          "混凝土（普混）", "混凝土（透水）", "建设工程", "脚手架",
                          "劳动争议", "劳务", "铝模", "民间借贷", "票据纠纷", "其他",
                          "其他建材", "砌体", "砂浆", "塔吊", "外加剂"],
        }, ensure_ascii=False)
        try:
            db.execute("INSERT OR IGNORE INTO settings(key,value) VALUES('case_type_sections', ?)",
                       (_sect_preset,))
        except Exception:
            pass
        # 结算模板表：Word 模板四部分结构所需字段（幂等补齐）
        tcols = {r[1] for r in db.execute("PRAGMA table_info(settle_templates)")}
        for name, ddl in [
            ("type", "TEXT DEFAULT 'salary_settlement'"),   # 模板类型：salary_settlement / case_approval
            ("is_system", "INTEGER DEFAULT 0"),             # 是否系统内置（内置禁止删除）
            ("title", "TEXT DEFAULT ''"),                   # 文档标题
            ("fields", "TEXT DEFAULT '[]'"),                # 第1部分案件基本信息字段配置 JSON
            ("commission_config", "TEXT DEFAULT '{}'"),     # 第2部分提成比例配置 JSON（handler_rate/maintainer_rate）
        ]:
            if name not in tcols:
                db.execute("ALTER TABLE settle_templates ADD COLUMN %s %s" % (name, ddl))
        # 案件费用明细表（第2部分自动生成的数据源：正=收入，负=垫付）
        db.execute("""CREATE TABLE IF NOT EXISTS case_fee_detail (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            case_id INTEGER NOT NULL,        -- 关联案件 ID
            fee_date TEXT DEFAULT '',        -- 回款/垫付日期 YYYY-MM-DD
            description TEXT DEFAULT '',     -- 款项说明
            amount REAL DEFAULT 0,           -- 金额（正数为收入，负数为垫付）
            created_at TEXT, updated_at TEXT
        )""")
        # 清理重复字段：维护人与开发人统一，cases.maintain_user 由成员表同步，不再手工维护
        db.commit()
        # 结算模板管理注册（ORG_NAME）案件结案审批表自定义模板（幂等）
        ensure_close_approval_template(db)
        db.commit()
    except Exception as e:
        print("[迁移] 字段补齐跳过:", e)
    finally:
        db.close()


def now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def round_half_up(v):
    """金额四舍五入取整（标准 ROUND_HALF_UP）。
    Python 内置 round() 为「银行家舍入」（.5 取偶），不符合财务习惯，故显式实现。"""
    from decimal import Decimal, ROUND_HALF_UP
    return float(Decimal(str(v)).quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def hash_pw(pw):
    import hashlib
    salt = "coop-salt-2026"
    return hashlib.sha256((salt + pw).encode()).hexdigest()


def get_setting(key, default=""):
    row = get_db().execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    return row["value"] if row else default


# ---------------------------------------------------------------- 审计日志
def audit(action, obj_type, obj_id="", detail=""):
    try:
        get_db().execute(
            "INSERT INTO audit_logs(user_id,username,action,object_type,object_id,detail,ip,created_at)"
            " VALUES(?,?,?,?,?,?,?,?)",
            (session.get("uid"), session.get("username", "-"), action,
             obj_type, str(obj_id), detail[:2000], request.remote_addr or "", now()))
        get_db().commit()
    except Exception:
        pass


# ---------------------------------------------------------------- 案号生成
def _fmt_case_no(ym, seq):
    """将年月(4位: 年两位+月两位)与顺序号格式化为案号。
    顺序号 <100 固定占 2 位（不足补 0）；>=100 按实际位数输出。
    例：#240506（seq=6）、#2405101（seq=101）。"""
    seq_str = "%02d" % seq if seq < 100 else str(seq)
    return "#%s%s" % (ym, seq_str)


def next_case_no(db):
    """生成案号：# + 年两位 + 月两位 + 顺序号。
    规则：顺序号 <100 固定 2 位（补 0，如 06）；>=100 按实际位数（如 101）。
    编号规则与唯一性保证：
      1) 按年月重置：计数键 ym = 年两位+月两位（如 2409），新月份从起始序号重新计；
      2) 起始序号可调：settings.case_no_seq_start（默认 1），仅在该年月首次生成时生效；
      3) 并发去重：BEGIN IMMEDIATE 加锁完成「读-改-写」原子自增，配合 cases.case_no
         的 UNIQUE 约束与调用方重试，杜绝重复案号。"""
    ym = date.today().strftime("%y%m")            # 2 位年 + 2 位月
    start = 1
    try:
        start = int(get_setting("case_no_seq_start", "1") or "1")
    except (TypeError, ValueError):
        start = 1
    if getattr(db, "in_transaction", False):
        db.commit()                               # 提交导入循环等已开启的未决事务，避免「事务中再开事务」报错
    db.execute("BEGIN IMMEDIATE")                 # 立即加写锁，避免并发读-改-写竞态
    try:
        row = db.execute("SELECT seq FROM case_no_seq WHERE ym=?", (ym,)).fetchone()
        if row is None:
            # 首次进入该年月：扫描已有同前缀案号取最大序号作为起点（兼容手动录入）
            mx = start - 1
            for r in db.execute(
                    "SELECT case_no FROM cases WHERE case_no LIKE ? AND deleted=0",
                    ("#%s%%" % ym,)).fetchall():
                m = re.match(r"^#\d{4}(\d+)$", r["case_no"] or "")
                if m:
                    mx = max(mx, int(m.group(1)))
            seq = mx + 1
            db.execute("INSERT INTO case_no_seq(ym,seq) VALUES(?,?)", (ym, seq))
        else:
            seq = row[0] + 1
            db.execute("UPDATE case_no_seq SET seq=? WHERE ym=?", (seq, ym))
        db.commit()
    except Exception:
        try:
            db.rollback()
        except Exception:
            pass
        raise
    return _fmt_case_no(ym, seq)


# ---------------------------------------------------------------- 安全公式求值
_ALLOWED_BIN = {ast.Add: lambda a, b: a + b, ast.Sub: lambda a, b: a - b,
                ast.Mult: lambda a, b: a * b, ast.Div: lambda a, b: a / b if b else 0,
                ast.Mod: lambda a, b: a % b if b else 0}
_ALLOWED_FN = {"min": min, "max": max, "round": round, "abs": abs}


def safe_eval(expr, names):
    """仅允许 数字/变量/+ - * / % 括号/min max round abs 的公式求值"""
    tree = ast.parse(expr, mode="eval")

    def ev(node):
        if isinstance(node, ast.Expression):
            return ev(node.body)
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
            return node.value
        if isinstance(node, ast.Name):
            if node.id in names:
                return float(names[node.id] or 0)
            raise ValueError("未知变量: %s" % node.id)
        if isinstance(node, ast.BinOp) and type(node.op) in _ALLOWED_BIN:
            return _ALLOWED_BIN[type(node.op)](ev(node.left), ev(node.right))
        if isinstance(node, ast.UnaryOp) and isinstance(node.op, (ast.USub, ast.UAdd)):
            v = ev(node.operand)
            return -v if isinstance(node.op, ast.USub) else v
        if isinstance(node, ast.Call) and isinstance(node.func, ast.Name) \
                and node.func.id in _ALLOWED_FN:
            return _ALLOWED_FN[node.func.id](*[ev(a) for a in node.args])
        raise ValueError("公式含有不允许的内容")

    return ev(tree)


# ================================================================ 页面
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/pair/info")
def pair_info():
    """登录页获取配网信息（用于展示，不泄露敏感配置）"""
    return jsonify({"ok": True, "firm": get_setting("firm_name", ORG_FULL_NAME)})


# ---------------------------------------------------------------- 全局错误处理
# 保证任何异常都返回 JSON（而非 HTML 错误页），前端才能解析出明确提示
@app.errorhandler(400)
def err_400(e):
    return jsonify({"ok": False, "msg": "请求格式错误，请刷新页面后重试"}), 400


@app.errorhandler(404)
def err_404(e):
    return jsonify({"ok": False, "msg": "接口不存在（404）"}), 404


@app.errorhandler(405)
def err_405(e):
    return jsonify({"ok": False, "msg": "请求方式不允许（405）"}), 405


@app.errorhandler(500)
def err_500(e):
    return jsonify({"ok": False, "msg": "服务器内部错误，请联系管理员查看日志"}), 500


@app.errorhandler(Exception)
def err_any(e):
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return jsonify({"ok": False, "msg": "请求失败（%d）" % e.code}), e.code
    import traceback
    traceback.print_exc()
    return jsonify({"ok": False, "msg": "服务器内部错误，请联系管理员"}), 500


# ================================================================ 认证
@app.route("/api/login", methods=["POST"])
def login():
    # 容错解析：即使请求体不是合法 JSON 也不抛 400 HTML 页
    d = request.get_json(force=True, silent=True) or {}
    username = str(d.get("username") or "").strip()
    password = str(d.get("password") or "")
    if not username or not password:
        return jsonify({"ok": False, "msg": "请输入用户名和密码"}), 400

    u = get_db().execute("SELECT * FROM users WHERE username=?",
                         (username,)).fetchone()
    if not u:
        audit("login", "system", username, "登录失败：账户不存在")
        return jsonify({"ok": False, "msg": "账户不存在，请核对用户名或联系财务经理开通"}), 401
    if not u["active"]:
        audit("login", "system", username, "登录失败：账户已停用")
        return jsonify({"ok": False, "msg": "账户已被停用，请联系财务经理"}), 403
    if u["password_hash"] != hash_pw(password):
        audit("login", "system", username, "登录失败：密码错误")
        return jsonify({"ok": False, "msg": "用户名或密码错误"}), 401

    session.update(uid=u["id"], username=u["username"],
                   name=u["display_name"], role=u["role"],
                   perms=resolve_perms(u))
    audit("login", "system", u["username"], "登录成功")
    return jsonify({"ok": True, "user": {"name": u["display_name"],
                                         "role": u["role"],
                                         "role_name": ROLES[u["role"]]["name"]}})


@app.route("/api/logout", methods=["POST"])
def logout():
    audit("logout", "system", session.get("username", "-"), "退出登录")
    session.clear()
    return jsonify({"ok": True})


@app.route("/api/me")
@login_required
def me():
    perms = [p for p in PERMS if has_perm(p)]
    return jsonify({"ok": True,
                    "user": {"name": session["name"], "username": session["username"],
                             "role": role(), "role_name": ROLES[role()]["name"]},
                    "perms": perms,
                    "firm": get_setting("firm_name")})


@app.route("/api/password", methods=["POST"])
@login_required
def change_pw():
    d = request.get_json(force=True)
    u = get_db().execute("SELECT * FROM users WHERE id=?", (session["uid"],)).fetchone()
    if u["password_hash"] != hash_pw(d.get("old", "")):
        return jsonify({"ok": False, "msg": "原密码错误"}), 400
    if len(d.get("new", "")) < 6:
        return jsonify({"ok": False, "msg": "新密码至少 6 位"}), 400
    get_db().execute("UPDATE users SET password_hash=? WHERE id=?",
                     (hash_pw(d["new"]), session["uid"]))
    get_db().commit()
    audit("update", "user", session["username"], "修改本人密码")
    return jsonify({"ok": True})


# ================================================================ 用户管理（admin）
@app.route("/api/users", methods=["GET", "POST"])
@perm_required("user.manage")
def users():
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT id,username,display_name,role,perms,active,created_at"
                          " FROM users ORDER BY id").fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows],
                        "roles": {k: v["name"] for k, v in ROLES.items()}})
    d = request.get_json(force=True)
    if d.get("role") not in ROLES:
        return jsonify({"ok": False, "msg": "角色无效"}), 400
    perms = validate_perms(d.get("perms"))
    if perms is None and d.get("perms") is not None:
        return jsonify({"ok": False, "msg": "权限格式错误"}), 400
    try:
        cur = db.execute(
            "INSERT INTO users(username,password_hash,display_name,role,perms,created_at)"
            " VALUES(?,?,?,?,?,?)",
            (d["username"].strip(), hash_pw(d.get("password") or "123456"),
             d["display_name"].strip(), d["role"],
             json.dumps(perms, ensure_ascii=False) if perms else None,
             now()))
        db.commit()
        audit("create", "user", d["username"],
              "新建用户 %s 角色=%s%s" % (
                  d["display_name"], d["role"],
                  " 权限=%s" % ",".join(perms) if perms else " 权限=继承角色默认"))
        return jsonify({"ok": True, "id": cur.lastrowid})
    except sqlite3.IntegrityError:
        return jsonify({"ok": False, "msg": "用户名已存在"}), 400


@app.route("/api/users/<int:uid>", methods=["PUT", "DELETE"])
@perm_required("user.manage")
def user_op(uid):
    db = get_db()
    u = db.execute("SELECT * FROM users WHERE id=?", (uid,)).fetchone()
    if not u:
        return jsonify({"ok": False, "msg": "用户不存在"}), 404
    if request.method == "DELETE":
        if u["username"] == "admin":
            return jsonify({"ok": False, "msg": "内置管理员不可删除"}), 400
        # 数据清理：移除该用户的个人薪酬规则与其账号生命周期审计记录；
        # 案件 / 结算单 / 操作日志等业务历史数据予以保留，确保可追溯。
        db.execute("DELETE FROM salary_rules WHERE user_name=? OR user_name=?",
                   (u["username"], u["display_name"]))
        db.execute("DELETE FROM audit_logs WHERE object_type='user' AND object_id=?",
                   (u["username"],))
        db.execute("DELETE FROM users WHERE id=?", (uid,))
        db.commit()
        audit("delete", "user", u["username"],
              "删除用户 %s（已清理其个人薪酬规则与账号记录）" % u["display_name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    fields, vals = [], []
    for k in ("display_name", "role"):
        if k in d:
            fields.append(k + "=?"); vals.append(d[k])
    if d.get("password"):
        fields.append("password_hash=?"); vals.append(hash_pw(d["password"]))
    if "active" in d:
        fields.append("active=?"); vals.append(1 if d["active"] else 0)
    if "perms" in d:
        if d["perms"] is None:
            fields.append("perms=?"); vals.append(None)
        else:
            pv = validate_perms(d.get("perms"))
            if pv is None:
                return jsonify({"ok": False, "msg": "权限格式错误"}), 400
            fields.append("perms=?"); vals.append(json.dumps(pv, ensure_ascii=False) if pv else None)
    if fields:
        vals.append(uid)
        db.execute("UPDATE users SET %s WHERE id=?" % ",".join(fields), vals)
        db.commit()
        audit("update", "user", u["username"], "修改: %s" % ",".join(f.split("=")[0] for f in fields))
    return jsonify({"ok": True})


# ================================================================ 角色权限配置
# 「机构负责人」与「律师」等新增角色默认在代码内 PERMS 中预置；财务经理可在系统设置
# 页面调整每个角色拥有的权限（仅修改「谁能做什么」，不修改角色名/角色描述）。
@app.route("/api/role-config", methods=["GET"])
@perm_required("user.manage")
def get_role_config():
    return jsonify({
        "ok": True,
        "roles": {k: {"name": v["name"], "desc": v["desc"]} for k, v in ROLES.items()},
        "perms": _perms_to_dict(),
        "perm_labels": PERM_LABELS,
        "config_path": ROLE_CONFIG_PATH,
    })


@app.route("/api/role-config", methods=["PUT"])
@perm_required("user.manage")
def update_role_config():
    """更新角色-权限矩阵。body: { perms: { perm_key: [role_key, ...] } }。
    仅接受已在 ROLES 中声明的角色名；其他角色会被忽略。仅更新已声明的权限项。"""
    d = request.get_json(force=True) or {}
    perms_in = d.get("perms")
    if not isinstance(perms_in, dict):
        return jsonify({"ok": False, "msg": "参数格式错误：需要 perms 字典"}), 400

    old = {k: set(v) for k, v in PERMS.items()}
    for k, roles in perms_in.items():
        if k not in PERMS or not isinstance(roles, list):
            continue
        PERMS[k] = {r for r in roles if r in ROLES}

    # 「setting.manage」必须始终包含 admin，防止误操作导致无人能进设置页
    if "admin" not in PERMS["setting.manage"]:
        PERMS["setting.manage"] |= {"admin"}

    ok = save_role_config()
    if not ok:
        # 保存失败：回滚到旧值
        for k, v in old.items():
            PERMS[k] = v
        return jsonify({"ok": False, "msg": "保存失败：无法写入角色配置文件"}), 500
    audit("update", "role_config", "perms", "更新角色权限矩阵")
    return jsonify({"ok": True, "perms": _perms_to_dict()})


# ================================================================ 案件管理
def case_to_dict(r):
    return {k: r[k] for k in r.keys()}


@app.route("/api/cases")
@perm_required("case.view")
def case_list():
    audit("query", "case", "", "查询案件列表")
    db = get_db()
    kw = request.args.get("kw", "").strip()
    sql = "SELECT * FROM cases WHERE deleted=0"
    args = []
    if kw:
        sql += (" AND (client_name LIKE ? OR case_no LIKE ? OR owner LIKE ?"
                " OR developer LIKE ? OR case_cause LIKE ? OR plaintiff LIKE ? OR defendant LIKE ?)")
        args += ["%" + kw + "%"] * 7
    sql += " ORDER BY id DESC"
    rows = db.execute(sql, args).fetchall()
    return jsonify({"ok": True, "data": [case_to_dict(r) for r in rows]})


@app.route("/api/cases", methods=["POST"])
@perm_required("case.edit")
def case_create():
    d = request.get_json(force=True)
    if not d.get("client_name"):
        return jsonify({"ok": False, "msg": "客户名称必填"}), 400
    if any(k in d for k in CASE_FIN_FIELDS + CASE_EXT_FIN_FIELDS + CASE_EXT_FIN_TEXT_FIELDS) and not has_perm("case.finance"):
        return jsonify({"ok": False, "msg": "无财务字段填写权限"}), 403
    db = get_db()
    # 案号：支持手动录入（case_no 非空则用），否则按自定义规则自动生成
    case_no = (d.get("case_no") or "").strip()
    if not case_no:
        case_no = next_case_no(db)
    elif db.execute("SELECT 1 FROM cases WHERE case_no=?", (case_no,)).fetchone():
        return jsonify({"ok": False, "msg": "案号已存在，请更换"}), 400
    cur = db.execute(
        """INSERT INTO cases(case_no,developer,owner,members,client_name,case_type,status,
           agent_fee,court_fee_advance,invoice_amount,
           plaintiff,defendant,section,case_cause,subject_amount,project,fee_standard,expected_profit,maintain_user,
           archive_status,archive_borrow,borrow_user,borrow_time,borrow_return_time,
           recover_month,settle_month,confirmed_amount,settle_ready,
           expected_profit_text,finance_remark,
           remark,created_by,created_at,updated_by,updated_at)
           VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (case_no, d.get("developer", ""), d.get("owner", ""), d.get("members", ""),
         d["client_name"], d.get("case_type", ""), d.get("status", "进行中"),
         float(d.get("agent_fee") or 0), float(d.get("court_fee_advance") or 0),
         float(d.get("invoice_amount") or 0),
         d.get("plaintiff", ""), d.get("defendant", ""), d.get("section", ""),
         d.get("case_cause", ""), float(d.get("subject_amount") or 0), d.get("project", ""),
         d.get("fee_standard", ""), float(d.get("expected_profit") or 0),
         d.get("maintain_user", ""),
         d.get("archive_status") or "未归档", d.get("archive_borrow", ""),
         d.get("borrow_user", ""), d.get("borrow_time", ""), d.get("borrow_return_time", ""),
         d.get("recover_month", ""), d.get("settle_month", ""),
         float(d.get("confirmed_amount") or 0), 0,
         d.get("expected_profit_text", ""), d.get("finance_remark", ""),
         d.get("remark", ""),
         session["username"], now(), session["username"], now()))
    db.commit()
    audit("create", "case", case_no, "新建案件 客户=%s" % d["client_name"])
    return jsonify({"ok": True, "case_no": case_no, "id": cur.lastrowid})


@app.route("/api/cases/<int:cid>", methods=["GET", "PUT", "DELETE"])
def case_op(cid):
    """GET：单案详情（第1部分自动填充用）；PUT：乐观锁字段级合并；DELETE：软删除。"""
    if not session.get("uid"):
        return jsonify({"ok": False, "msg": "未登录"}), 401
    db = get_db()
    row = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not row:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404

    if request.method == "GET":
        if not has_perm("case.view"):
            return jsonify({"ok": False, "msg": "无权限"}), 403
        return jsonify({"ok": True, "data": dict(row)})

    if request.method == "DELETE":
        if not has_perm("case.delete"):
            return jsonify({"ok": False, "msg": "无删除案件权限（需「删除案件」权限）"}), 403
        db.execute("UPDATE cases SET deleted=1, updated_by=?, updated_at=? WHERE id=?",
                   (session["username"], now(), cid))
        db.commit()
        audit("delete", "case", row["case_no"], "删除案件 客户=%s" % row["client_name"])
        return jsonify({"ok": True})

    d = request.get_json(force=True)
    base_version = int(d.get("base_version") or 0)
    changes = d.get("fields") or {}
    conflict = base_version < row["version"]   # 他人已先保存

    # 字段级权限过滤（含扩展字段）
    allowed = {}
    for k, v in changes.items():
        if k in (CASE_BASE_FIELDS + CASE_EXT_BASE_FIELDS) and has_perm("case.edit"):
            allowed[k] = v
        elif k in (CASE_FIN_FIELDS + CASE_EXT_FIN_FIELDS) and has_perm("case.finance"):
            try:
                allowed[k] = float(v or 0)
            except (TypeError, ValueError):
                return jsonify({"ok": False, "msg": "金额字段必须为数字"}), 400
        elif k in CASE_EXT_FIN_TEXT_FIELDS and has_perm("case.finance"):
            allowed[k] = v
    if not allowed:
        return jsonify({"ok": False, "msg": "没有可更新的字段或权限不足"}), 400

    sets, vals, overwritten = [], [], []
    for k, v in allowed.items():
        if conflict and str(row[k]) != str(v):
            overwritten.append("%s: %s -> %s" % (k, row[k], v))  # 最后写入优先，记录旧值
        sets.append(k + "=?"); vals.append(v)
    sets += ["version=version+1", "updated_by=?", "updated_at=?"]
    vals += [session["username"], now(), cid]
    db.execute("UPDATE cases SET %s WHERE id=?" % ",".join(sets), vals)
    db.commit()

    detail = "更新字段: %s" % ",".join(allowed.keys())
    if conflict:
        detail += " | 并发冲突合并(最后写入优先)，覆盖: " + "; ".join(overwritten)
    audit("update", "case", row["case_no"], detail)
    return jsonify({"ok": True, "conflict_merged": conflict,
                    "overwritten": overwritten})


@app.route("/api/cases/<int:cid>/status", methods=["PATCH"])
@perm_required("case.edit")
def case_status_update(cid):
    """案件状态变更（含流转校验、执行状态、原因、审计日志）。
    新工作流使用英文 code 存储，旧中文状态通过 norm_case_status 无损兼容。"""
    d = request.get_json(force=True) or {}
    target = (d.get("status") or "").strip()
    reason = (d.get("reason") or "").strip()
    exec_status = (d.get("execution_status") or "").strip() or None

    if not target:
        return jsonify({"ok": False, "msg": "目标状态不能为空"}), 400
    if target not in CASE_STATUS_CODES:
        return jsonify({"ok": False, "msg": "无效的案件状态: " + target}), 400
    if exec_status and exec_status not in EXECUTION_STATUS_CODES:
        return jsonify({"ok": False, "msg": "无效的执行状态: " + exec_status}), 400

    db = get_db()
    row = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not row:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404

    current = norm_case_status(row["status"])
    current_meta = CASE_STATUS_BY_CODE.get(current)
    target_meta = CASE_STATUS_BY_CODE[target]

    # 流转校验：允许保持当前状态；若目标在 next 列表外则拒绝
    allowed = set(current_meta.get("next", [])) if current_meta else set()
    if target != current and target not in allowed:
        return jsonify({"ok": False,
                        "msg": "状态流转非法：%s（%s）不能到 %s（%s）" % (
                            status_label(current), current,
                            status_label(target), target)}), 400

    # 结果类/终止类/异常类必须填写原因
    if target_meta["category"] in ("result", "terminated", "abnormal") and not reason:
        return jsonify({"ok": False, "msg": "进入结果/终止/异常状态需填写原因"}), 400

    # 当目标不是 "execution" 时，清空执行状态；否则沿用或更新
    final_exec = ""
    if target == "execution":
        final_exec = exec_status or row.get("execution_status") or "not_applied"

    new_label = target_meta["label"]
    db.execute(
        """UPDATE cases
           SET status=?, execution_status=?, version=version+1,
               updated_by=?, updated_at=?
           WHERE id=?""",
        (new_label, final_exec, session["username"], now(), cid))
    db.commit()

    detail = "状态变更 %s(%s) -> %s(%s)" % (status_label(current), current, new_label, target)
    if reason:
        detail += " 原因：" + reason
    if final_exec:
        detail += " 执行状态：" + EXECUTION_STATUS_BY_CODE.get(final_exec, {}).get("label", final_exec)
    audit("update", "case", row["case_no"], detail)
    return jsonify({"ok": True, "status": target, "execution_status": final_exec,
                    "label": new_label})


@app.route("/api/cases/<int:cid>/active", methods=["PUT"])
def case_active_toggle(cid):
    """案件启用/停止切换：target=1 启用（需 case.enable），target=0 停止（需 case.stop）。
    权限细分，确保「停止」「启用」各自受对应权限控制。"""
    if not session.get("uid"):
        return jsonify({"ok": False, "msg": "未登录"}), 401
    d = request.get_json(force=True) or {}
    target = 1 if d.get("active") else 0
    needed = "case.enable" if target == 1 else "case.stop"
    if not has_perm(needed):
        return jsonify({"ok": False,
                        "msg": ("无启用案件权限" if target == 1 else "无停止案件权限")}), 403
    db = get_db()
    row = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not row:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    if row["active"] == target:
        return jsonify({"ok": True, "active": target, "msg": "状态未变化"})
    db.execute(
        "UPDATE cases SET active=?, version=version+1, updated_by=?, updated_at=? WHERE id=?",
        (target, session["username"], now(), cid))
    db.commit()
    audit("update", "case", row["case_no"], ("启用案件" if target == 1 else "停止案件"))
    return jsonify({"ok": True, "active": target})



# 角色统一命名：维护人（原开发人/维护人合并）、主承办人（原负责人）、协办人（原助理）
MEMBER_ROLES = {
    "maintainer": "维护人",
    "owner": "主承办人",
    "assistant": "协办人",
}
# 兼容旧数据里的角色键
ROLE_ALIAS = {"developer": "maintainer", "maintain": "maintainer", "负责人": "owner"}


def norm_role(role):
    """归一化角色键：开发人→维护人，兼容历史数据"""
    role = (role or "").strip()
    return ROLE_ALIAS.get(role, role)


def default_guaranteed(role):
    """默认保底金额：维护人 0 元，其他成员 600 元"""
    return 0.0 if norm_role(role) == "maintainer" else 600.0


def sync_case_members(db, cid):
    """根据 case_members 同步 cases 冗余字段（developer/owner/members/maintain_user），
    用于列表展示与旧逻辑兼容。人员信息唯一来源是 case_members，不再手工编辑这些字段。"""
    rows = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()

    def names(role):
        return [r["person"] for r in rows if norm_role(r["role"]) == role]

    devs = names("maintainer")
    owns = names("owner")
    asst = names("assistant")
    db.execute("UPDATE cases SET developer=?, owner=?, members=?, maintain_user=?, updated_at=? WHERE id=?",
               (",".join(devs), ",".join(owns), ",".join(asst),
                (devs[0] if devs else ""), now(), cid))


@app.route("/api/cases/<int:cid>/members", methods=["GET", "POST"])
@perm_required("case.view")
def case_members(cid):
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM case_members WHERE case_id=? ORDER BY id",
                          (cid,)).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows],
                        "roles": MEMBER_ROLES})
    if not has_perm("case.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    person = (d.get("person") or "").strip()
    role = norm_role(d.get("role") or "assistant")
    if not person:
        return jsonify({"ok": False, "msg": "请填写成员姓名"}), 400
    if role not in MEMBER_ROLES:
        return jsonify({"ok": False, "msg": "角色无效"}), 400
    cnt = db.execute("SELECT COUNT(*) n FROM case_members WHERE case_id=?", (cid,)).fetchone()["n"]
    if cnt >= 10:
        return jsonify({"ok": False, "msg": "小组成员最多 10 人"}), 400
    # 提成率：未显式提供时默认从花名册自动引入
    rate = d.get("rate")
    if rate is None or rate == "":
        s = db.execute("SELECT commission_rate FROM staff WHERE name=? AND deleted=0 LIMIT 1",
                       (person,)).fetchone()
        rate = float(s["commission_rate"] or 0) if s else 0.0
    else:
        rate = float(rate or 0)
    gua = d.get("guaranteed_amount")
    gua = default_guaranteed(role) if (gua is None or gua == "") else float(gua or 0)
    fixed = float(d.get("fixed_amount") or 0)
    cur = db.execute(
        "INSERT INTO case_members(case_id,person,role,rate,guaranteed_amount,fixed_amount,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?)",
        (cid, person, role, rate, gua, fixed, session["username"], now()))
    db.commit()
    sync_case_members(db, cid)
    db.commit()
    audit("create", "case_member", cur.lastrowid,
          "案件 %d 新增成员 %s(%s) 提成率 %.2f%%" % (cid, person, role, rate * 100))
    return jsonify({"ok": True, "id": cur.lastrowid, "rate": rate,
                    "guaranteed_amount": gua, "fixed_amount": fixed})


@app.route("/api/cases/<int:cid>/members/<int:mid>", methods=["PUT", "DELETE"])
@perm_required("case.edit")
def case_member_op(cid, mid):
    db = get_db()
    r = db.execute("SELECT * FROM case_members WHERE id=? AND case_id=?", (mid, cid)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "成员不存在"}), 404
    if request.method == "DELETE":
        db.execute("DELETE FROM case_members WHERE id=?", (mid,))
        db.commit()
        sync_case_members(db, cid)
        db.commit()
        audit("delete", "case_member", mid, "案件 %d 移除成员 %s" % (cid, r["person"]))
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    sets, vals = [], []
    for k in ("person", "role"):
        if k in d:
            sets.append(k + "=?"); vals.append(d[k])
    if "rate" in d:
        sets.append("rate=?"); vals.append(float(d.get("rate") or 0))
    if "guaranteed_amount" in d:
        sets.append("guaranteed_amount=?"); vals.append(float(d.get("guaranteed_amount") or 0))
    if "fixed_amount" in d:
        sets.append("fixed_amount=?"); vals.append(float(d.get("fixed_amount") or 0))
    if sets:
        vals.append(mid)
        db.execute("UPDATE case_members SET %s WHERE id=?" % ",".join(sets), vals)
        db.commit()
        sync_case_members(db, cid)
        db.commit()
        audit("update", "case_member", mid, "案件 %d 修改成员 %s" % (cid, r["person"]))
    return jsonify({"ok": True})


@app.route("/api/cases/<int:cid>/members/sync", methods=["POST"])
@perm_required("case.edit")
def sync_members(cid):
    """全量重建案件成员（保存案件时调用）。成员最多 10 人；
    提成率默认从花名册引入；保底金额默认维护人 0 元、其他成员 600 元，均可手动修改。"""
    d = request.get_json(force=True)
    members = d.get("members") or []
    if len(members) > 10:
        return jsonify({"ok": False, "msg": "小组成员最多 10 人"}), 400
    db = get_db()
    db.execute("DELETE FROM case_members WHERE case_id=?", (cid,))
    rows = []
    for m in members:
        person = (m.get("person") or "").strip()
        role = norm_role(m.get("role") or "assistant")
        if role not in MEMBER_ROLES:
            role = "assistant"
        if not person:
            continue
        # 提成比例：未填则默认从花名册引入
        rate = m.get("rate")
        if rate is None or rate == "":
            s = db.execute("SELECT commission_rate FROM staff WHERE name=? AND deleted=0 LIMIT 1",
                           (person,)).fetchone()
            rate = float(s["commission_rate"] or 0) if s else 0.0
        else:
            rate = float(rate or 0)
        # 保底金额：未填则按角色给默认（维护人 0，其他 600）
        gua = m.get("guaranteed_amount")
        gua = default_guaranteed(role) if (gua is None or gua == "") else float(gua or 0)
        fixed = float(m.get("fixed_amount") or 0)
        db.execute(
            "INSERT INTO case_members(case_id,person,role,rate,guaranteed_amount,fixed_amount,created_by,created_at)"
            " VALUES(?,?,?,?,?,?,?,?)",
            (cid, person, role, rate, gua, fixed, session["username"], now()))
        rows.append({"person": person, "role": role, "rate": rate,
                     "guaranteed_amount": gua, "fixed_amount": fixed})
    db.commit()
    sync_case_members(db, cid)
    db.commit()
    audit("update", "case_member", cid, "案件 %d 同步成员 %d 人" % (cid, len(rows)))
    return jsonify({"ok": True, "count": len(rows), "members": rows})


# ---------------------------------------------------------------- 案件费用收回明细
@app.route("/api/cases/<int:cid>/recoveries", methods=["GET", "POST"])
@perm_required("case.view")
def case_recoveries(cid):
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM fee_recoveries WHERE case_id=? ORDER BY id",
                          (cid,)).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows]})
    if not has_perm("case.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    cur = db.execute(
        "INSERT INTO fee_recoveries(case_id,recovery_date,item,amount,confirm,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?)",
        (cid, d.get("recovery_date", ""), d.get("item", ""),
         float(d.get("amount") or 0), 1 if d.get("confirm") else 0,
         session["username"], now()))
    db.commit()
    audit("create", "fee_recovery", cur.lastrowid,
          "案件 %d 新增费用收回 %s %.2f" % (cid, d.get("item", ""), float(d.get("amount") or 0)))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/cases/<int:cid>/recoveries/<int:rid>", methods=["DELETE"])
@perm_required("case.edit")
def del_recovery(cid, rid):
    db = get_db()
    db.execute("DELETE FROM fee_recoveries WHERE id=? AND case_id=?", (rid, cid))
    db.commit()
    audit("delete", "fee_recovery", rid, "案件 %d 删除费用收回" % cid)
    return jsonify({"ok": True})


# ---------------------------------------------------------------- 案件提成明细
@app.route("/api/cases/<int:cid>/commissions", methods=["GET", "POST"])
@perm_required("case.view")
def case_commissions(cid):
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM case_commissions WHERE case_id=? ORDER BY id",
                          (cid,)).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows]})
    if not has_perm("salary.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    base_amount = float(d.get("base_amount") or 0)
    rate = float(d.get("rate") or 0)
    raw = round(base_amount * rate, 2)              # 取整前精确金额（留痕可复核）
    amount = round_half_up(raw)                     # 四舍五入取整后的提成金额
    remark = "（取整）" if abs(raw - amount) >= 0.005 else ""
    calc = "%.2f × %.0f%% = %.2f" % (base_amount, rate * 100, raw)
    calc += (" → 取整 %.2f 元" % amount) if remark else " 元"
    cur = db.execute(
        "INSERT INTO case_commissions(case_id,person,role,rate,base_amount,amount,raw_amount,remark,calc_process,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?)",
        (cid, d.get("person", ""), d.get("role", ""), rate, base_amount, amount, raw,
         remark, calc, session["username"], now()))
    db.commit()
    audit("create", "commission", cur.lastrowid,
          "案件 %d 提成 %s %.2f (%s)" % (cid, d.get("person", ""), amount, calc))
    return jsonify({"ok": True, "id": cur.lastrowid, "amount": amount,
                    "raw_amount": raw, "remark": remark, "calc_process": calc})


@app.route("/api/cases/<int:cid>/commissions/calc", methods=["POST"])
@perm_required("salary.edit")
def calc_case_commissions(cid):
    """根据费用收回合计（实际收益）+ 薪酬规则自动计算提成明细，并展示完整计算过程"""
    db = get_db()
    case = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not case:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    recs = db.execute("SELECT amount FROM fee_recoveries WHERE case_id=?", (cid,)).fetchall()
    base_amount = sum(float(r["amount"] or 0) for r in recs)
    # 清除旧提成，按薪酬规则重新生成
    db.execute("DELETE FROM case_commissions WHERE case_id=?", (cid,))
    commissions = []
    # 案件相关人员
    persons = set([case["owner"] or "", case["maintain_user"] or "", case["developer"] or ""] +
                  [m.strip() for m in (case["members"] or "").split(",") if m.strip()])
    persons.discard("")
    for person in persons:
        rules = db.execute("SELECT * FROM salary_rules WHERE user_name=? ORDER BY id",
                           (person,)).fetchall()
        for r in rules:
            if r["effective_from"] and r["effective_from"] > date.today().isoformat():
                continue
            if r["mode"] == "fixed":
                raw = round(float(r["fixed_amount"] or 0), 2)
                rate = 0.0
                calc = "固定金额 = %.2f" % raw
            else:
                try:
                    raw = round(safe_eval(r["formula"],
                        {"agent_fee": float(case["agent_fee"] or 0),
                         "court_fee_advance": float(case["court_fee_advance"] or 0),
                         "invoice_amount": float(case["invoice_amount"] or 0)}), 2)
                except Exception:
                    continue
                rate = 0.0
                calc = "按公式：%s = %.2f" % (r["formula"], raw)
            amount = round_half_up(raw)              # 四舍五入取整
            remark = "（取整）" if abs(raw - amount) >= 0.005 else ""
            calc += (" → 取整 %.2f 元" % amount) if remark else " 元"
            role = "承办人" if person == case["owner"] else ("维护人" if person == case["maintain_user"] else "参与人")
            cur = db.execute(
                "INSERT INTO case_commissions(case_id,person,role,rate,base_amount,amount,raw_amount,remark,calc_process,created_by,created_at)"
                " VALUES(?,?,?,?,?,?,?,?,?,?,?)",
                (cid, person, role, rate, base_amount, amount, raw, remark, calc,
                 session["username"], now()))
            commissions.append({"id": cur.lastrowid, "person": person, "role": role,
                                "amount": amount, "raw_amount": raw, "remark": remark,
                                "calc_process": calc})
    db.commit()
    audit("create", "commission", cid, "案件 %d 自动计算提成，基数 %.2f，共 %d 条" % (cid, base_amount, len(commissions)))
    return jsonify({"ok": True, "base_amount": base_amount, "commissions": commissions})


@app.route("/api/cases/<int:cid>/commissions/<int:cmid>", methods=["DELETE"])
@perm_required("salary.edit")
def del_commission(cid, cmid):
    db = get_db()
    db.execute("DELETE FROM case_commissions WHERE id=? AND case_id=?", (cmid, cid))
    db.commit()
    audit("delete", "commission", cmid, "案件 %d 删除提成" % cid)
    return jsonify({"ok": True})


# ================================================================ 薪酬计算
@app.route("/api/salary/rules", methods=["GET", "POST"])
@perm_required("salary.view")
def salary_rules():
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM salary_rules ORDER BY id DESC").fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows]})
    if not has_perm("salary.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    if d.get("mode") == "formula":
        try:  # 保存前先校验公式合法性
            safe_eval(d.get("formula", ""), {"agent_fee": 1, "invoice_amount": 1,
                                             "court_fee_advance": 1})
        except Exception as e:
            return jsonify({"ok": False, "msg": "公式无效: %s" % e}), 400
    cur = db.execute(
        "INSERT INTO salary_rules(user_name,mode,formula,fixed_amount,effective_from,note,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?)",
        (d["user_name"].strip(), d.get("mode", "formula"), d.get("formula", ""),
         float(d.get("fixed_amount") or 0), d.get("effective_from", ""),
         d.get("note", ""), session["username"], now()))
    db.commit()
    audit("create", "salary_rule", cur.lastrowid, "人员=%s 模式=%s" % (d["user_name"], d.get("mode")))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/salary/rules/<int:rid>", methods=["DELETE"])
@perm_required("salary.edit")
def salary_rule_del(rid):
    db = get_db()
    r = db.execute("SELECT * FROM salary_rules WHERE id=?", (rid,)).fetchone()
    if r:
        db.execute("DELETE FROM salary_rules WHERE id=?", (rid,))
        db.commit()
        audit("delete", "salary_rule", rid, "删除 %s 的薪酬规则" % r["user_name"])
    return jsonify({"ok": True})


def person_in_case(case_row, person):
    """判断人员是否参与案件：开发人/负责人/小组成员（均为逗号分隔）"""
    def in_field(field):
        return person in [x.strip() for x in (case_row[field] or "").split(",") if x.strip()]
    return in_field("developer") or in_field("owner") or in_field("members")


@app.route("/api/salary/calc")
@perm_required("salary.view")
def salary_calc():
    """按案件 + 成员提成率计算薪酬，以案号为索引。
    公式：max(到账金额 × 提成率, 600) + 固定金额，四舍五入取整。
    到账金额 = 财务确认回款金额（confirmed_amount，缺省回退到已确认费用收回合计）；
    提成率 = 案件成员预计提成率（默认从花名册引入，可手动修改）。"""
    db = get_db()
    # 按自定义年月归集（month=YYYY-MM，按提成结算月筛选）
    month = request.args.get("month", "").strip()
    sql = "SELECT * FROM cases WHERE deleted=0"
    args = []
    if month:
        sql += " AND settle_month=?"
        args.append(month)
    sql += " ORDER BY id DESC"
    cases = db.execute(sql, args).fetchall()
    audit("query", "salary_rule", "", "执行薪酬试算")
    # 固定金额仅在案件成员(case_members)中维护，不再使用 salary_rules 跨案件规则
    # 手动调整金额（按案件+人员+结算月覆盖自动计算结果）
    adj_map = {}
    for a in db.execute("SELECT * FROM salary_adjustments").fetchall():
        adj_map[(a["case_id"], a["person"], a["month"])] = float(a["amount"] or 0)
    result = []
    for c in cases:
        members = db.execute("SELECT * FROM case_members WHERE case_id=?", (c["id"],)).fetchall()
        recovered = float(c["confirmed_amount"] or 0)
        if recovered <= 0:
            rec = db.execute(
                "SELECT COALESCE(SUM(amount),0) s FROM fee_recoveries WHERE case_id=? AND confirm=1",
                (c["id"],)).fetchone()["s"]
            recovered = float(rec or 0)
        items = []
        for m in members:
            rate = float(m["rate"] or 0)
            # 保底金额：成员自定义优先，否则按角色默认（维护人 0，其他 600）
            gua = m["guaranteed_amount"]
            gua = default_guaranteed(m["role"]) if gua is None else float(gua or 0)
            # 固定金额：仅取自案件成员配置
            fixed = float(m["fixed_amount"] or 0)
            # 统一薪酬公式：max(到账金额 × 提成率, 保底金额) + 固定金额
            base = max(recovered * rate, gua)
            auto_amount = round_half_up(base + fixed)
            # 手动调整覆盖
            adj_key = (c["id"], m["person"], c["settle_month"])
            adjusted = adj_key in adj_map
            amount = adj_map[adj_key] if adjusted else auto_amount
            items.append({
                "member_id": m["id"], "person": m["person"], "role": m["role"],
                "role_name": MEMBER_ROLES.get(norm_role(m["role"]), m["role"]),
                "rate": rate, "guaranteed_amount": gua, "fixed": fixed,
                "base": round(base, 2), "amount": amount, "adjusted": adjusted,
                "calc": "max(%.2f × %.2f%%, 保底%.2f) + 固定%.2f = %.2f → 取整 %d"
                        % (recovered, rate * 100, gua, fixed, base + fixed, auto_amount)})
        result.append({
            "case_id": c["id"], "case_no": c["case_no"], "client_name": c["client_name"],
            "recovered": round(recovered, 2), "recover_month": c["recover_month"],
            "settle_month": c["settle_month"], "settle_ready": bool(c["settle_ready"]),
            "member_count": len(members), "items": items,
            "total": sum(i["amount"] for i in items)})
    return jsonify({"ok": True, "data": result})


@app.route("/api/salary/adjust", methods=["POST"])
@perm_required("salary.edit")
def salary_adjust():
    """手动调整某案件某成员在某结算月的薪酬金额"""
    d = request.get_json(force=True)
    case_id = int(d.get("case_id") or 0)
    person = (d.get("person") or "").strip()
    month = (d.get("month") or "").strip()
    if not case_id or not person or not month:
        return jsonify({"ok": False, "msg": "参数缺失"}), 400
    db = get_db()
    amount = d.get("amount")
    if amount is None or amount == "":
        # 删除调整，恢复自动计算
        db.execute("DELETE FROM salary_adjustments WHERE case_id=? AND person=? AND month=?",
                   (case_id, person, month))
        db.commit()
        return jsonify({"ok": True, "cleared": True})
    amount = float(amount)
    db.execute(
        "INSERT INTO salary_adjustments(case_id,person,month,amount,created_by,created_at) VALUES(?,?,?,?,?,?)"
        " ON CONFLICT(case_id,person,month) DO UPDATE SET amount=excluded.amount",
        (case_id, person, month, amount, session["username"], now()))
    db.commit()
    audit("update", "salary_adjustment", case_id,
          "手动调整 %s %s 薪酬=%.2f" % (person, month, amount))
    return jsonify({"ok": True, "amount": amount})


# ---------------------------------------------------------------- 提成结算判断
@app.route("/api/cases/<int:cid>/check-settle", methods=["POST"])
@perm_required("case.edit")
def check_settle(cid):
    """判断案件提成是否可结算：财务已确认回款金额(confirmed_amount>0) 且有成员"""
    db = get_db()
    c = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not c:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    members = db.execute("SELECT COUNT(*) n FROM case_members WHERE case_id=?", (cid,)).fetchone()["n"]
    ready = 1 if (float(c["confirmed_amount"] or 0) > 0 and members > 0) else 0
    db.execute("UPDATE cases SET settle_ready=? WHERE id=?", (ready, cid))
    db.commit()
    audit("update", "case", c["case_no"], "判断提成可结算=%d" % ready)
    return jsonify({"ok": True, "settle_ready": bool(ready),
                    "msg": "满足结算条件" if ready else "未满足：需财务确认回款金额且案件有成员"})


@app.route("/api/cases/check-settle-all", methods=["POST"])
@perm_required("case.edit")
def check_settle_all():
    """批量判断全部案件是否可结算，返回达标数量"""
    db = get_db()
    cases = db.execute("SELECT * FROM cases WHERE deleted=0").fetchall()
    ready_cnt = 0
    for c in cases:
        members = db.execute("SELECT COUNT(*) n FROM case_members WHERE case_id=?", (c["id"],)).fetchone()["n"]
        ready = 1 if (float(c["confirmed_amount"] or 0) > 0 and members > 0) else 0
        if ready != c["settle_ready"]:
            db.execute("UPDATE cases SET settle_ready=? WHERE id=?", (ready, c["id"]))
        if ready:
            ready_cnt += 1
    db.commit()
    audit("update", "case", "", "批量判断提成可结算，达标 %d 件" % ready_cnt)
    return jsonify({"ok": True, "ready_count": ready_cnt})


# ================================================================ 批量结算单（Word）
SETTLE_FIELDS = ["case_no", "client_name", "developer", "owner", "members",
                 "agent_fee", "court_fee_advance", "invoice_amount",
                 "user_name", "amount", "batch_no", "date", "firm"]

# （ORG_NAME）案件结案审批表：注册到结算模板管理的默认自定义模板名称
CLOSE_APPROVAL_TEMPLATE_NAME = "(ORG_NAME)案件结案审批表"

# 审批表第1部分（案件基本信息）可配置字段：key 对应 cases 表字段
APPROVAL_PART1_FIELDS = [
    {"key": "plaintiff", "label": "原告"},
    {"key": "defendant", "label": "被告"},
    {"key": "section", "label": "板块"},
    {"key": "case_cause", "label": "案由"},
    {"key": "subject_amount", "label": "标的"},
    {"key": "project", "label": "项目"},
    {"key": "fee_standard", "label": "收益标准"},
    {"key": "expected_profit", "label": "预期收益"},
]


def ensure_close_approval_template(db):
    """将（ORG_NAME）案件结案审批表整合进结算模板管理（自定义模板），幂等。
    - config：结算单部分配置（title/fields/show_*/footer，type=close_approval 标记）；
    - fields：第1部分案件基本信息字段配置（全部可见，按 APPROVAL_PART1_FIELDS 顺序）；
    - commission_config：第2部分提成比例（承办人 5% / 维护人 8%）。"""
    try:
        row = db.execute(
            "SELECT id,name,is_default,config,type,is_system,title,fields,commission_config "
            "FROM settle_templates WHERE name=?", (CLOSE_APPROVAL_TEMPLATE_NAME,)).fetchone()
        # 系统内置模板保证库中始终有默认模板（若已无任何默认，则由系统模板担任）
        has_default = db.execute(
            "SELECT 1 FROM settle_templates WHERE is_default=1 LIMIT 1").fetchone()
        config = {
            "title": CLOSE_APPROVAL_TEMPLATE_NAME,
            "type": "close_approval",
            "show_batch": False,
            "show_date": False,
            "show_user": False,
            "fields": [{"key": f["key"], "label": f["label"]} for f in SETTLE_TMPL_FIELDS],
            "footer": "财务经理签字：______________    领款人签字：______________",
        }
        part1 = [{"key": f["key"], "label": f["label"], "visible": True,
                  "sort_order": i} for i, f in enumerate(APPROVAL_PART1_FIELDS)]
        cc = json.dumps({"handler_rate": 5, "maintainer_rate": 8}, ensure_ascii=False)
        if row:
            old = {}
            try:
                old = json.loads(row[3] or "{}")
            except Exception:
                pass
            upd = []
            if old.get("type") != "close_approval":
                old.update(config)
                upd.append(("config", json.dumps(old, ensure_ascii=False)))
            if row[4] != "case_approval":
                upd.append(("type", "case_approval"))
            if not row[5]:
                upd.append(("is_system", 1))
            if not (row[6] or "").strip():
                upd.append(("title", "（ORG_NAME）案件结案审批表"))
            if not row[2] and not has_default:
                upd.append(("is_default", 1))
            try:
                if not json.loads(row[7] or "[]"):
                    upd.append(("fields", json.dumps(part1, ensure_ascii=False)))
            except Exception:
                upd.append(("fields", json.dumps(part1, ensure_ascii=False)))
            try:
                if not json.loads(row[8] or "{}"):
                    upd.append(("commission_config", cc))
            except Exception:
                upd.append(("commission_config", cc))
            if upd:
                upd_sql = ",".join("%s=?" % k for k, _ in upd)
                vals = [v for _, v in upd] + ["system", now(), row[0]]
                db.execute("UPDATE settle_templates SET %s,updated_by=?,updated_at=? WHERE id=?"
                           % upd_sql, vals)
        else:
            db.execute(
                "INSERT INTO settle_templates(name,is_default,type,is_system,title,fields,"
                "commission_config,config,created_by,created_at,updated_by,updated_at)"
                " VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                (CLOSE_APPROVAL_TEMPLATE_NAME, 1 if not has_default else 0, "case_approval", 1, "（ORG_NAME）案件结案审批表",
                 json.dumps(part1, ensure_ascii=False), cc,
                 json.dumps(config, ensure_ascii=False),
                 "system", now(), "system", now()))
        print("[迁移] 结算模板管理已注册自定义模板：%s" % CLOSE_APPROVAL_TEMPLATE_NAME)
    except Exception as e:
        print("[迁移] 注册（ORG_NAME）案件结案审批表模板跳过:", e)

# 结算单可配置字段（模板设计器用）：key 为可显示字段，label 为默认中文名
SETTLE_TMPL_FIELDS = [
    {"key": "case_no", "label": "案号"},
    {"key": "client_name", "label": "客户名称"},
    {"key": "developer", "label": "案件开发人"},
    {"key": "owner", "label": "负责人"},
    {"key": "members", "label": "小组成员"},
    {"key": "agent_fee", "label": "代理费金额"},
    {"key": "court_fee_advance", "label": "代垫诉讼费"},
    {"key": "invoice_amount", "label": "开票金额"},
    {"key": "user_name", "label": "结算人员"},
    {"key": "amount", "label": "结算金额"},
]

# 结算模板默认 JSON 配置
DEFAULT_SETTLE_CONFIG = {
    "title": "薪酬结算单",
    "show_batch": True,
    "show_date": True,
    "fields": [
        {"key": "case_no", "label": "案号"},
        {"key": "client_name", "label": "客户名称"},
        {"key": "developer", "label": "案件开发人"},
        {"key": "owner", "label": "负责人"},
        {"key": "members", "label": "小组成员"},
        {"key": "agent_fee", "label": "代理费金额"},
        {"key": "court_fee_advance", "label": "代垫诉讼费"},
        {"key": "invoice_amount", "label": "开票金额"},
    ],
    "footer": "财务经理签字：______________    领款人签字：______________",
}


def render_settlement_from_config(config, out_path, ctx):
    """按 JSON 模板配置渲染结算单（无需外部 docx 文件）"""
    from docx import Document
    doc = Document()
    firm = ctx.get("firm", "")
    title = config.get("title") or "薪酬结算单"
    doc.add_heading("%s %s" % (firm, title), level=1)
    meta = []
    if config.get("show_batch"):
        meta.append("批次号：%s" % ctx.get("batch_no", ""))
    if config.get("show_date"):
        # 日期留白：ctx 未提供 date 时显示下划线占位，由用户自行填写
        meta.append("结算日期：%s" % (ctx.get("date") or "____年__月__日"))
    if config.get("show_user"):
        meta.append("结算人员：%s" % ctx.get("user_name", ""))
    if meta:
        doc.add_paragraph("        ".join(meta))
    fields = config.get("fields") or []
    # 结算金额独立高亮展示
    amount = float(ctx.get("amount") or 0)
    doc.add_paragraph("")
    doc.add_heading("结算金额：¥ %.2f" % amount, level=2)
    if fields:
        tb = doc.add_table(rows=len(fields), cols=2)
        tb.style = "Table Grid"
        for i, f in enumerate(fields):
            key = f.get("key")
            label = f.get("label") or key
            tb.rows[i].cells[0].text = label
            tb.rows[i].cells[1].text = str(ctx.get(key, ""))
    doc.add_paragraph("")
    doc.add_paragraph(config.get("footer") or "")
    doc.save(out_path)


def render_settlement_docx(tpl_path, out_path, ctx):
    """模板占位符 {{field}} 替换；无自定义模板时用内置格式生成"""
    from docx import Document
    if tpl_path and os.path.exists(tpl_path):
        doc = Document(tpl_path)
        for p in doc.paragraphs:
            for k in SETTLE_FIELDS:
                if "{{%s}}" % k in p.text:
                    for run in p.runs:
                        if "{{%s}}" % k in run.text:
                            run.text = run.text.replace("{{%s}}" % k, str(ctx.get(k, "")))
        for t in doc.tables:
            for rowx in t.rows:
                for cell in rowx.cells:
                    for p in cell.paragraphs:
                        for k in SETTLE_FIELDS:
                            if "{{%s}}" % k in p.text:
                                for run in p.runs:
                                    if "{{%s}}" % k in run.text:
                                        run.text = run.text.replace(
                                            "{{%s}}" % k, str(ctx.get(k, "")))
    else:
        doc = Document()
        doc.add_heading("%s 薪酬结算单" % ctx["firm"], level=1)
        doc.add_paragraph("批次号：%s        结算日期：%s" % (ctx["batch_no"], ctx.get("date") or "____年__月__日"))
        doc.add_paragraph("结算人员：%s" % ctx["user_name"])
        tb = doc.add_table(rows=8, cols=2)
        tb.style = "Table Grid"
        pairs = [("案号", ctx["case_no"]), ("客户名称", ctx["client_name"]),
                 ("案件开发人", ctx["developer"]), ("负责人", ctx["owner"]),
                 ("小组成员", ctx["members"]),
                 ("代理费金额（¥）", ctx["agent_fee"]),
                 ("代垫诉讼费（¥）", ctx["court_fee_advance"]),
                 ("开票金额（¥）", ctx["invoice_amount"])]
        for i, (k, v) in enumerate(pairs):
            tb.rows[i].cells[0].text = k
            tb.rows[i].cells[1].text = str(v)
        doc.add_paragraph("")
        doc.add_heading("结算金额：¥ %.2f" % float(ctx["amount"]), level=2)
        doc.add_paragraph("财务经理签字：______________    领款人签字：______________")
    doc.save(out_path)


def render_close_approval_docx(case, recoveries, commissions, out_path, fields=None):
    """按纸质《（ORG_NAME）案件结案审批表》版式生成 Word（四段式）：
    第一行：维护人 / 承办人；
    第一部分 案件基本信息（原告/被告/板块/案由（案件类型）/标的/项目/收益标准/预期收益）；
    第二部分 案件办结情况（1 费用收回+出纳审核意见 2 收益 3 备注 4 提成计算+维护人签字）；
    第三部分 审批意见（财务主管 / 执行主任）；第四部分 说明文字。
    数据由案件管理 / 财务信息 / 提成明细索引。fields 参数保留兼容（不再分段控制）。"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    # 全局字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    def cv(k, default=''):
        try:
            return case[k] if k in case.keys() else default
        except Exception:
            return default

    def add_table(rows, cols, headers=None):
        t = doc.add_table(rows=(len(rows) + (1 if headers else 0)), cols=cols)
        t.style = 'Table Grid'
        if headers:
            for j, h in enumerate(headers):
                t.rows[0].cells[j].text = h
        for i, row in enumerate(rows):
            for j, v in enumerate(row):
                t.rows[i + (1 if headers else 0)].cells[j].text = ('' if v is None else str(v))
        return t

    # 标题
    title = doc.add_heading('（ORG_NAME）案件结案审批表', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第一行：维护人 / 承办人
    add_table([['维护人：%s' % (cv('maintain_user') or ''),
                '承办人：%s' % (cv('owner') or '')]], 2)
    doc.add_paragraph('')

    # 第一部分 案件基本信息
    doc.add_heading('一、案件基本信息', level=2)
    cause = cv('case_cause') or cv('case_type') or ''
    exp_text = (cv('expected_profit_text') or '').strip()
    exp_val = exp_text if exp_text else ('%.2f' % float(cv('expected_profit') or 0))
    add_table([
        ['原告', cv('plaintiff') or '', '被告', cv('defendant') or ''],
        ['板块', cv('section') or '', '案由（案件类型）', cause],
        ['标的', '%.2f' % float(cv('subject_amount') or 0), '项目', cv('project') or ''],
        ['收益标准', cv('fee_standard') or '', '预期收益', exp_val],
    ], 4)

    # 第二部分 案件办结情况
    doc.add_heading('二、案件办结情况', level=2)
    income, expense = 0.0, 0.0
    for r in recoveries:
        amt = float(r['amount'] or 0)
        if amt > 0:
            income += amt
        else:
            expense += -amt
    # 1、费用收回 + 出纳审核意见
    doc.add_paragraph('1、费用收回：')
    if recoveries:
        add_table([[r['recovery_date'] or '', r['item'] or '',
                    '%.2f' % float(r['amount'] or 0)] for r in recoveries], 3,
                  headers=['日期', '款项说明', '金额（¥）'])
    else:
        doc.add_paragraph('（暂无费用收回明细）')
    doc.add_paragraph('出纳审核意见：________________________')
    # 2、收益（索引案件信息中预期收益文本框内容）
    doc.add_paragraph('2、收益（预期收益）：%s'
                      % (exp_text if exp_text else ('%.2f 元' % float(cv('expected_profit') or 0))))
    # 3、备注行
    doc.add_paragraph('3、备注：%s' % ((cv('finance_remark') or '').strip() or '________________________'))
    # 4、提成计算（按现有薪酬结算模板填入）+ 维护人签字行及日期
    doc.add_paragraph('4、提成计算：')
    if commissions:
        add_table([[c['role'] or '', c['person'] or '',
                    ('%.0f%%' % (float(c['rate'] or 0) * 100)) if c['rate'] else '-',
                    c['calc_process'] or '', '%.2f' % float(c['amount'] or 0),
                    (c['remark'] or '')] for c in commissions], 6,
                  headers=['角色', '姓名', '比例', '计算过程', '金额（¥）', '备注'])
    else:
        doc.add_paragraph('（暂无提成明细）')
    add_table([['维护人签字：________________', '签字日期：______年____月____日']], 2)

    # 第三部分 审批意见
    doc.add_heading('三、审批意见', level=2)
    doc.add_paragraph('财务主管意见：________________________')
    add_table([['财务主管签字：________________', '签字日期：______年____月____日']], 2)
    doc.add_paragraph('执行主任意见：________________________')
    add_table([['执行主任签字：________________', '签字日期：______年____月____日']], 2)

    # 第四部分 说明
    doc.add_heading('四、说明', level=2)
    for line in [
        '1、本表是案件办结归档的凭证，不入卷；',
        '2、本表由财务人员负责填写报批，交财务部备存；',
        '3、各项费用必须写明回款时间和金额；',
        '4、必须附有收益计算表、代理协议、诉讼费用票据、文书（调解书、和解协议、撤诉裁定）、结算说明、收据发票；',
        '5、此表是诉讼部和市场部计算提成的唯一依据。']:
        p = doc.add_paragraph(line)
        try:
            p.runs[0].font.size = Pt(9)
        except Exception:
            pass
    doc.save(out_path)


# ---------------------------------------------------------------- 案件结案审批表生成
@app.route("/api/cases/<int:cid>/close-approval")
@perm_required("settle.view")
def close_approval(cid):
    """生成当前案件的《案件结案审批表》Word 文档。
    支持 query string ?fields=case_no,client_name,... 控制输出哪些段落（与前端字段自定义一致）。"""
    db = get_db()
    case = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not case:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    recoveries = db.execute("SELECT * FROM fee_recoveries WHERE case_id=? ORDER BY id",
                            (cid,)).fetchall()
    commissions = db.execute("SELECT * FROM case_commissions WHERE case_id=? ORDER BY id",
                             (cid,)).fetchall()
    # 解析 fields（白名单校验：只接受已知字段，非法值忽略）
    raw = (request.args.get("fields") or "").strip()
    fields = None
    if raw:
        allow = {'case_no','client_name','case_type','status','developer','owner',
                 'recoveries','profit_calc','total_profit','commissions',
                 'commission_total','remark','sign_column'}
        fields = [f.strip() for f in raw.split(',') if f.strip() in allow]
        if not fields:
            fields = None
    fname = "%s_%s_结案审批表.docx" % (date.today().isoformat(), case["case_no"])
    fpath = os.path.join(SETTLE_DIR, fname)
    render_close_approval_docx(case, recoveries, commissions, fpath, fields)
    audit("export", "close_approval", case["case_no"],
          "生成案件结案审批表 fields=%s" % (",".join(fields) if fields else "default"))
    return send_file(fpath, as_attachment=True,
                     download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ---------------------------------------------------------------- （ORG_NAME）案件结案审批表（合并：审批表 + 按成员薪酬结算单）
@app.route("/api/cases/<int:cid>/close-approval-bundle")
@perm_required("settle.view")
def close_approval_bundle(cid):
    """合并「批量生成薪酬结算单」与「案件结案审批表」：
    一键生成《（ORG_NAME）案件结案审批表》（含提成计算）+ 按成员《薪酬结算单》，打包 ZIP 返回。
    数据更新后直接调用即可，样式参考纸质单据。"""
    db = get_db()
    case = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not case:
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    recoveries = db.execute("SELECT * FROM fee_recoveries WHERE case_id=? ORDER BY id",
                            (cid,)).fetchall()
    commissions = db.execute("SELECT * FROM case_commissions WHERE case_id=? ORDER BY id",
                             (cid,)).fetchall()
    batch_no = "CGSP" + datetime.now().strftime("%Y%m%d%H%M%S")
    tpl = get_setting("settle_template")
    tpl_path = os.path.join(TEMPLATE_DIR, tpl) if tpl else None
    tpl_config = None
    tr = db.execute("SELECT * FROM settle_templates WHERE is_default=1 LIMIT 1").fetchone()
    if tr:
        try:
            tpl_config = json.loads(tr["config"] or "{}")
        except Exception:
            tpl_config = None

    files = []
    # 1) （ORG_NAME）案件结案审批表（四段式，含提成计算）
    fname = "%s_%s_（ORG_NAME）案件结案审批表.docx" % (date.today().isoformat(), case["case_no"])
    fpath = os.path.join(SETTLE_DIR, fname)
    render_close_approval_docx(case, recoveries, commissions, fpath)
    files.append(fpath)

    # 2) 按成员薪酬结算单（复用结算单生成逻辑）
    recovered = float(case["confirmed_amount"] or 0)
    if recovered <= 0:
        rec = db.execute(
            "SELECT COALESCE(SUM(amount),0) s FROM fee_recoveries WHERE case_id=? AND confirm=1",
            (cid,)).fetchone()["s"]
        recovered = float(rec or 0)
    members = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()
    for m in members:
        rate = float(m["rate"] or 0)
        gua = m["guaranteed_amount"]
        gua = default_guaranteed(m["role"]) if gua is None else float(gua or 0)
        fixed = float(m["fixed_amount"] or 0)
        # 统一薪酬公式：max(到账金额 × 提成率, 保底金额) + 固定金额
        base = max(recovered * rate, gua)
        amount = round_half_up(base + fixed)
        if amount <= 0:
            continue
        sfname = "%s_%s_%s_薪酬结算单.docx" % (batch_no, case["case_no"], m["person"])
        sfpath = os.path.join(SETTLE_DIR, sfname)
        ctx = {k: case[k] if k in case.keys() else "" for k in SETTLE_FIELDS}
        ctx.update(user_name=m["person"], amount=amount, batch_no=batch_no,
                   date="", firm=get_setting("firm_name"))  # date 留白，由用户自行填写
        if tpl_config:
            render_settlement_from_config(tpl_config, sfpath, ctx)
        else:
            render_settlement_docx(tpl_path, sfpath, ctx)
        files.append(sfpath)

    if not files:
        return jsonify({"ok": False, "msg": "该案件暂无可生成内容（无成员或回款金额）"}), 400
    import zipfile
    zip_path = os.path.join(EXPORT_DIR, batch_no + ".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, os.path.basename(f))
    audit("export", "close_approval_bundle", case["case_no"],
          "合并生成（ORG_NAME）案件结案审批表 + 结算单 %d 份" % len(files))
    return send_file(zip_path, as_attachment=True, download_name=batch_no + ".zip",
                     mimetype="application/zip")


# ---------------------------------------------------------------- （ORG_NAME）案件结案审批表合并渲染（新规格版式 + 自定义结算模板）
def build_commission_calc_detail(profit, rate_pct, guaranteed=0.0, fixed=0.0):
    """构造提成计算过程展开文本（逐行），供审批表第2部分提成区块与前端预览复用。

    统一公式：数额 = max(回款金额 × 提成率, 保底金额) + 固定金额，四舍五入取整。
    - profit：回款金额（本案收益 = 收款合计 - 垫付合计）
    - rate_pct：提成率（%，如 5 表示 5%）
    - guaranteed：保底金额（默认 0；维护人 0、其他成员 600）
    - fixed：固定金额（默认 0）
    返回：计算过程多行文本（首行公式 + 逐项代入数值展开 + 最终结果）。
    """
    p = float(profit or 0)
    r = float(rate_pct or 0)
    g = float(guaranteed or 0)
    f = float(fixed or 0)
    base = round_half_up(p * r / 100.0)      # 回款金额 × 提成率
    maxed = max(base, g)                     # 与保底金额取大
    exact = maxed + f                        # 加固定金额
    amount = round_half_up(exact)            # 四舍五入取整
    rate_txt = ("%.0f" % r) if r == int(r) else ("%.2f" % r)
    return [
        "数额 = max(回款金额 × 提成率, 保底金额) + 固定金额",
        "    = max(%.2f × %s%%, %.2f) + %.2f" % (p, rate_txt, g, f),
        "    = max(%.2f, %.2f) + %.2f" % (base, g, f),
        "    = %.2f + %.2f" % (maxed, f),
        "    = %.2f 元（四舍五入取整）" % amount,
    ]


def render_close_approval_bundle_docx(case, recoveries, commissions, tpl_config, members, out_path,
                                      fee_details=None, tpl_row=None, blank=False):
    """合并渲染《（ORG_NAME）案件结案审批表》单文档：
    - 审批表部分按新规格版式（A4/宋体小四/黑体标题/0.5pt 黑色实线边框/灰底字段名）动态填充数据：
      · 第1部分 案件基本信息：按模板 fields 配置（visible + sort_order）从案件管理数据自动带入；
      · 第2部分 案件办结情况：费用收回文字 + 明细表（负数红色）+ 收益自动公式 + 提成自动计算
        （承办人=收益×handler_rate、维护人=收益×maintainer_rate，比例取模板 commission_config）；
      · 第3部分 审批意见：维护人/财务主管/执行主任签字位，姓名自动带入，日期留白「____年__月__日」待用户自行填写；
      · 第4部分 固定说明。
    - 结算单部分合并结算模板管理（自定义模板）的 config 内容（标题/字段表/落款），每成员一段。
    输出文件名统一为「xxx_（ORG_NAME）案件结案审批表.docx」。
    fee_details：费用明细 [{fee_date,description,amount}]，None 时由 recoveries（fee_recoveries）转换；
    tpl_row：settle_templates 整行（含 fields/commission_config/type/title），None 时用全部 8 字段 + commissions 明细；
    tpl_config：结算单部分配置 dict，None 时使用内置默认；
    blank：True 时输出空白模板（全部下划线占位、日期不填充、无结算单）。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FILL = "______"
    # ---- 数据源：费用明细（fee_details 优先，其次 fee_recoveries 转换） ----
    if fee_details is None:
        fee_details = [{"fee_date": r["recovery_date"] or "", "description": r["item"] or "",
                        "amount": float(r["amount"] or 0)} for r in (recoveries or [])]
    else:
        fee_details = [{"fee_date": str(d.get("fee_date") or ""),
                        "description": str(d.get("description") or ""),
                        "amount": float(d.get("amount") or 0)} for d in fee_details]
    # ---- 模板行解析：第1部分字段配置 + 提成比例 ----
    part1_fields = None
    commission_cfg = {}
    if tpl_row is not None:
        try:
            flds = tpl_row["fields"]
        except (KeyError, TypeError):
            flds = getattr(tpl_row, "fields", "[]")
        try:
            part1_fields = json.loads(flds or "[]") or None
        except Exception:
            part1_fields = None
        try:
            cc = tpl_row["commission_config"]
        except (KeyError, TypeError):
            cc = getattr(tpl_row, "commission_config", "{}")
        try:
            commission_cfg = json.loads(cc or "{}") or {}
        except Exception:
            commission_cfg = {}
    doc = Document()
    # ---- 页面与全局样式：A4 纵向 / 边距 / 宋体小四 1.5 倍行距 ----
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    st.paragraph_format.line_spacing = 1.5

    def cv(k, default=""):
        try:
            return case[k] if k in case.keys() else default
        except Exception:
            return default

    def _set_run(run, cn="宋体", en="Times New Roman", size=12, bold=False):
        run.font.name = en
        run.font.size = Pt(size)
        run.font.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), en)
        rFonts.set(qn("w:hAnsi"), en)
        rFonts.set(qn("w:eastAsia"), cn)

    def _para(text="", size=12, bold=False, align=None, before=0, after=0, cn="宋体"):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        if align is not None:
            p.alignment = align
        pf.line_spacing = 1.5
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if text:
            r = p.add_run(text)
            _set_run(r, cn=cn, size=size, bold=bold)
        return p

    def _borders(t, sz=4):
        tblPr = t._tbl.tblPr
        b = tblPr.find(qn("w:tblBorders"))
        if b is None:
            b = OxmlElement("w:tblBorders")
            tblPr.append(b)
        for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = b.find(qn("w:" + e))
            if el is None:
                el = OxmlElement("w:" + e)
                b.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")

    def _margins(t, top=57, bottom=57, left=85, right=85):
        tblPr = t._tbl.tblPr
        m = tblPr.find(qn("w:tblCellMar"))
        if m is None:
            m = OxmlElement("w:tblCellMar")
            tblPr.append(m)
        for tag, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            el = m.find(qn("w:" + tag))
            if el is None:
                el = OxmlElement("w:" + tag)
                m.append(el)
            el.set(qn("w:w"), str(v))
            el.set(qn("w:type"), "dxa")

    def _shade(cell, fill="F2F2F2"):
        tcPr = cell._tc.get_or_add_tcPr()
        s = tcPr.find(qn("w:shd"))
        if s is None:
            s = OxmlElement("w:shd")
            tcPr.append(s)
        s.set(qn("w:val"), "clear")
        s.set(qn("w:color"), "auto")
        s.set(qn("w:fill"), fill)

    def _cell(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(text)
        _set_run(r, size=size, bold=bold)
        if color:
            r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _new_table(rows, cols, widths=None):
        t = doc.add_table(rows=rows, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _borders(t)
        _margins(t)
        if widths:
            t.autofit = False
            t.allow_autofit = False
            for row in t.rows:
                for i, w in enumerate(widths):
                    if i < len(row.cells):
                        row.cells[i].width = w
        return t

    def _row_height(row, cm):
        trPr = row._tr.get_or_add_trPr()
        trH = trPr.find(qn("w:trHeight"))
        if trH is None:
            trH = OxmlElement("w:trHeight")
            trPr.append(trH)
        trH.set(qn("w:val"), str(int(cm * 567)))
        trH.set(qn("w:hRule"), "atLeast")

    def _merge_row(t, ri, text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                   fill=None, h_cm=None):
        row = t.rows[ri]
        m = row.cells[0]
        for c in row.cells[1:]:
            m = m.merge(c)
        _cell(m, text, size=size, bold=bold, align=align)
        if fill:
            _shade(m)
        if h_cm:
            _row_height(row, h_cm)
        return m

    # ================= 一、标题：黑体二号 22pt 加粗居中 =================
    _para("（ORG_NAME）案件结案审批表", size=22, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=12, cn="黑体")

    # ================= 二、表头信息行：维护人/承办人/卷号（无边框三列）=================
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.tblPr
    b = tblPr.find(qn("w:tblBorders"))
    if b is None:
        b = OxmlElement("w:tblBorders")
        tblPr.append(b)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn("w:" + e))
        if el is None:
            el = OxmlElement("w:" + e)
            b.append(el)
        el.set(qn("w:val"), "none")
    for j, txt in enumerate(["维护人：%s" % (cv("maintain_user") or FILL),
                             "承办人：%s" % (cv("owner") or FILL),
                             "卷号：%s" % (cv("case_no") or FILL)]):
        c = t.rows[0].cells[j]
        c.width = Cm(4.89)
        _cell(c, txt)
    _row_height(t.rows[0], 0.8)

    # ================= 三、案件基本信息（第1部分：按模板 fields 配置，数据自动带入）=================
    cause = cv("case_cause") or cv("case_type") or ""
    exp_text = (cv("expected_profit_text") or "").strip()
    exp_val = exp_text if exp_text else ("%.2f" % float(cv("expected_profit") or 0))
    _p1_fields = {
        "plaintiff": ("原告", cv("plaintiff")),
        "defendant": ("被告", cv("defendant")),
        "section": ("板块", cv("section")),
        "case_cause": ("案由", cause),
        "subject_amount": ("标的", "%.2f" % float(cv("subject_amount") or 0)),
        "project": ("项目", cv("project")),
        "fee_standard": ("收益标准", cv("fee_standard")),
        "expected_profit": ("预期收益", exp_val),
    }
    info_items = []
    if part1_fields:
        _shown = [f for f in part1_fields if f.get("visible", True)]
        _shown.sort(key=lambda f: int(f.get("sort_order") or 0))
        for f in _shown:
            key = f.get("key") or f.get("field_key")
            if key in _p1_fields:
                info_items.append((f.get("label") or _p1_fields[key][0], _p1_fields[key][1]))
    if not info_items:  # 无配置时回退全部 8 字段
        info_items = [(v[0], v[1]) for v in _p1_fields.values()]
    _rows4 = [info_items[i:i + 2] for i in range(0, len(info_items), 2)]
    t = _new_table(1 + len(_rows4), 4, widths=[Cm(2.2), Cm(5.13), Cm(2.2), Cm(5.13)])
    _merge_row(t, 0, "案件基本信息", h_cm=0.8)
    for i, pair in enumerate(_rows4, start=1):
        row = t.rows[i]
        _row_height(row, 0.8)
        f1, v1 = pair[0]
        _cell(row.cells[0], f1, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(row.cells[0])
        _cell(row.cells[1], (FILL if blank else v1) or FILL)
        if len(pair) > 1:
            f2, v2 = pair[1]
            _cell(row.cells[2], f2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade(row.cells[2])
            _cell(row.cells[3], (FILL if blank else v2) or FILL)

    # ================= 四、案件办结情况（第2部分：大表格，系统自动生成）=================
    t = _new_table(1, 1, widths=[Cm(14.66)])
    _merge_row(t, 0, "案件办结情况", h_cm=0.8)

    def _trow(text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, h=0.8, fill=None, color=None):
        row = t.add_row()
        _cell(row.cells[0], text, size=size, bold=bold, align=align, color=color)
        if fill:
            _shade(row.cells[0])
        _row_height(row, h)
        return row

    # 收益基础数据：正=收款，负=垫付
    income_amt, expense_amt = 0.0, 0.0
    pos, neg = [], []
    for d in fee_details:
        a = float(d["amount"] or 0)
        if a >= 0:
            pos.append(d)
            income_amt += a
        else:
            neg.append(d)
            expense_amt += -a

    # 子区块1 费用收回：自动文字描述 + 明细表格（负数红色）+ 出纳审核意见（右下角）
    _trow("1、费用收回：", bold=True, h=0.7)
    if fee_details:
        for d in fee_details:
            a = float(d["amount"] or 0)
            _trow("%s    %s    %.2f 元" % (d["fee_date"] or "", d["description"] or "", a),
                  h=0.6, color=("FF0000" if a < 0 else None))
    else:
        for _ in range(3):
            _trow(FILL, h=0.9)
    _narr = []
    for d in pos:
        _narr.append("于%s收款%.2f元%s" % (d["fee_date"] or "____年__月__日", d["amount"],
                                          ("（%s）" % d["description"]) if d["description"] else ""))
    for d in neg:
        _desc = d["description"]
        if _desc.startswith("垫付"):
            _desc = _desc[2:]
        _narr.append("即垫付的%s%.2f元" % (_desc or "费用", -d["amount"]))
    if _narr:
        _trow("；".join(_narr) + "。", h=0.8)
    _trow("出纳审核意见：%s" % FILL, align=WD_ALIGN_PARAGRAPH.RIGHT, h=0.7)

    # 子区块2 收益：自动计算公式（各项收款之和 - 各项垫付 = 最终收益，保留2位小数）
    if not pos and not neg:
        formula = "本案收益为 %s+%s-%s-%s-%s-%s-%s-%s=%s 元" % tuple([FILL] * 9)
    else:
        parts = []
        for i, d in enumerate(pos):
            parts.append(("+" if i else "") + "%.2f" % d["amount"])
        for d in neg:
            parts.append("-" + "%.2f" % (-d["amount"]))
        profit = round_half_up(income_amt - expense_amt)
        formula = "本案收益为 %s=%s 元" % (" ".join(parts) if parts else "0", "%.2f" % profit)
    _trow("2、收益：", bold=True, h=0.7)
    _trow(formula, h=0.8)

    # 子区块3 备注
    _trow("3、备注（如有）：", bold=True, h=0.7)
    _trow((cv("finance_remark") or "").strip() or FILL, h=0.8)

    # 子区块4 提成（优先按模板 commission_config 自动计算并展开完整计算过程，其次用提成明细表）
    _trow("4、提成：", bold=True, h=0.7)
    h_rate = commission_cfg.get("handler_rate")
    m_rate = commission_cfg.get("maintainer_rate")

    def _fmt_rate(v):
        v = float(v)
        return ("%.0f" % v) if v == int(v) else ("%.2f" % v)

    # 成员配置查找：按角色取该成员的保底/固定金额（缺失用角色默认）
    def _member_cfg(role_key):
        for m in members or []:
            try:
                if norm_role(m["role"]) == role_key:
                    gua = m["guaranteed_amount"]
                    gua = default_guaranteed(role_key) if gua is None else float(gua or 0)
                    return gua, float(m["fixed_amount"] or 0)
            except (KeyError, TypeError, IndexError):
                continue
        return default_guaranteed(role_key), 0.0

    if blank or (h_rate is None and m_rate is None and not commissions):
        _trow("〈1〉承办人：%s   标准：__%%    数额 %s 元" % (FILL, FILL), h=0.8)
        _trow("〈2〉维护人：%s   标准：__%%    数额 %s 元" % (FILL, FILL), h=0.8)
    elif h_rate is not None or m_rate is not None:
        profit = round_half_up(income_amt - expense_amt)
        if h_rate is not None:
            h_gua, h_fixed = _member_cfg("owner")
            h_amt = round_half_up(max(profit * float(h_rate) / 100.0, h_gua) + h_fixed)
            _trow("〈1〉承办人：%s   标准：%s%%   数额 %.2f 元" % (
                (cv("owner") or FILL), _fmt_rate(h_rate), h_amt), h=0.8)
            for _ln in build_commission_calc_detail(profit, h_rate, h_gua, h_fixed):
                _trow(_ln, size=10.5, h=0.55)
        else:
            _trow("〈1〉承办人：%s   标准：__%%    数额 %s 元" % (FILL, FILL), h=0.8)
        if m_rate is not None:
            m_gua, m_fixed = _member_cfg("maintainer")
            m_amt = round_half_up(max(profit * float(m_rate) / 100.0, m_gua) + m_fixed)
            _trow("〈2〉维护人：%s   标准：%s%%   数额 %.2f 元" % (
                (cv("maintain_user") or FILL), _fmt_rate(m_rate), m_amt), h=0.8)
            for _ln in build_commission_calc_detail(profit, m_rate, m_gua, m_fixed):
                _trow(_ln, size=10.5, h=0.55)
        else:
            _trow("〈2〉维护人：%s   标准：__%%    数额 %s 元" % (FILL, FILL), h=0.8)
    else:
        for c in commissions:
            rate_s = ("%.0f%%" % (float(c["rate"] or 0) * 100)) if c["rate"] else "-"
            no = "1" if (c["role"] or "").find("承办") >= 0 else "2"
            _trow("〈%s〉%s：%s   标准：%s   数额 %.2f 元   （%s）" % (
                no, c["role"] or "", c["person"] or "", rate_s,
                float(c["amount"] or 0), c["calc_process"] or ""), h=0.7)

    # ================= 五、签字与审批区（第3部分：3行×2列，姓名自动带入，日期留白待填）=================
    # 日期不自动填充：统一留白「____年__月__日」下划线占位，由审批人后续自行填写，保持排版完整。
    # 如需恢复自动填充今天日期，将下行改为：
    #   _d_str = "" if blank else date.today().strftime("%Y年%m月%d日")
    _d_str = ""
    _maint_n = (cv("maintain_user") or "") if not blank else ""
    _owner_n = (cv("owner") or "") if not blank else ""
    t = _new_table(3, 2, widths=[Cm(8.0), Cm(6.66)])
    sign_rows = [
        ("维护人：%s" % (_maint_n or FILL), "%s" % (_d_str or "____年__月__日")),
        ("审批意见：%s" % FILL, "财务主管：%s   %s" % (FILL, _d_str or "____年__月__日")),
        ("执行主任：%s" % FILL, "%s" % (_d_str or "____年__月__日")),
    ]
    for i, (left, right) in enumerate(sign_rows):
        _row_height(t.rows[i], 1.0)
        _cell(t.rows[i].cells[0], left)
        _cell(t.rows[i].cells[1], right, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ================= 六、说明（标题加粗 + 5 条宋体五号 10.5pt）=================
    _para("说明：", bold=True, before=6)
    for line in [
        "1、本表是案件办结归档的凭证，不入卷；",
        "2、本表由财务人员负责填写报批，交财务部备存；",
        "3、各项费用必须写明回款时间和金额；",
        "4、必须附有收益计算表、代理协议、诉讼费用票据、文书（调解书、和解协议、撤诉裁定）、结算说明、收据发票；",
        "5、此表是诉讼部和市场部计算提成的唯一依据。",
    ]:
        _para(line, size=10.5)

    # ================= 七、薪酬结算单（自定义模板内容合并，每成员一段）=================
    recovered = float(cv("confirmed_amount") or 0)
    if recovered <= 0:
        rec = sum(float(r["amount"] or 0) for r in recoveries if r.get("confirm"))
        recovered = float(rec or 0)
    batch_no = "CGJD" + datetime.now().strftime("%Y%m%d%H%M%S")
    cfg = tpl_config or {}
    for m in members or []:
        rate = float(m["rate"] or 0)
        gua = m["guaranteed_amount"]
        gua = default_guaranteed(m["role"]) if gua is None else float(gua or 0)
        fixed = float(m["fixed_amount"] or 0)
        base = max(recovered * rate, gua)
        amount = round_half_up(base + fixed)
        if amount <= 0:
            continue
        doc.add_page_break()
        # 结算单部分标题：审批表模板(config.type=close_approval)合并时用「薪酬结算单」，
        # 其他自定义模板用其自身 title（实现"新建模板内容与审批表内容合并"）
        title = ("薪酬结算单" if cfg.get("type") in ("close_approval", "case_approval")
                 else (cfg.get("title") or "薪酬结算单"))
        _para("%s %s" % (get_setting("firm_name"), title), size=16, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
        meta = []
        if cfg.get("show_batch"):
            meta.append("批次号：%s" % batch_no)
        if cfg.get("show_date"):
            # 结算日期留白（下划线占位），由用户自行填写；如需自动填充改为：
            #   meta.append("结算日期：%s" % date.today().isoformat())
            meta.append("结算日期：____年__月__日")
        if cfg.get("show_user"):
            meta.append("结算人员：%s" % (m["person"] or ""))
        if meta:
            _para("        ".join(meta))
        t = _new_table(0, 2, widths=[Cm(4.5), Cm(10.16)])
        for f in cfg.get("fields") or []:
            k = f.get("key")
            label = f.get("label") or k
            if k == "user_name":
                val = m["person"] or ""
            elif k == "amount":
                val = "%.2f" % amount
            elif k in ("batch_no", "date", "firm"):
                continue
            else:
                val = str(cv(k))
            row = t.add_row()
            _row_height(row, 0.7)
            _cell(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade(row.cells[0])
            _cell(row.cells[1], val if val else FILL)
        _para("")
        _para("结算金额：¥ %.2f" % amount, size=14, bold=True)
        if cfg.get("footer"):
            _para(cfg["footer"])

    doc.save(out_path)


@app.route("/api/settlements/generate-bundle", methods=["POST"])
@perm_required("settle.edit")
def settle_generate_bundle():
    """一键生成结算文档（合并）：批量勾选案件 → 每案输出《（ORG_NAME）案件结案审批表》docx
    （审批表内容 + 结算模板管理自定义模板的薪酬结算单内容合并），打包 ZIP 返回。
    文档名称统一为「xxx_（ORG_NAME）案件结案审批表.docx」。"""
    d = request.get_json(force=True)
    case_ids = d.get("case_ids") or []
    if not case_ids:
        return jsonify({"ok": False, "msg": "请勾选案件"}), 400
    db = get_db()
    # 结算模板：优先显式指定 template_id，其次默认模板
    tpl_config = None
    tid = d.get("template_id")
    if tid:
        tr = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
        if tr:
            try:
                tpl_config = json.loads(tr["config"] or "{}")
            except Exception:
                tpl_config = None
    if tpl_config is None:
        dr = db.execute("SELECT * FROM settle_templates WHERE is_default=1 LIMIT 1").fetchone()
        if dr:
            try:
                tpl_config = json.loads(dr["config"] or "{}")
            except Exception:
                tpl_config = None
    batch_no = "JS" + datetime.now().strftime("%Y%m%d%H%M%S")
    files = []
    for cid in case_ids:
        case = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
        if not case:
            continue
        recoveries = db.execute("SELECT * FROM fee_recoveries WHERE case_id=? ORDER BY id",
                                (cid,)).fetchall()
        commissions = db.execute("SELECT * FROM case_commissions WHERE case_id=? ORDER BY id",
                                 (cid,)).fetchall()
        members = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()
        fname = "%s_（ORG_NAME）案件结案审批表.docx" % case["case_no"]
        fpath = os.path.join(SETTLE_DIR, fname)
        render_close_approval_bundle_docx(case, recoveries, commissions, tpl_config,
                                          members, fpath)
        files.append(fpath)
    if not files:
        return jsonify({"ok": False, "msg": "所选案件未生成任何文档"}), 400
    import zipfile
    zip_path = os.path.join(EXPORT_DIR, batch_no + ".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, os.path.basename(f))
    audit("export", "settlement_bundle", batch_no,
          "一键生成结算文档（合并（ORG_NAME）案件结案审批表） %d 份" % len(files))
    return jsonify({"ok": True, "batch_no": batch_no, "count": len(files),
                    "download": "/api/settlements/download/" + batch_no})


@app.route("/api/settle/templates/fields")
@perm_required("settle.view")
def settle_template_fields():
    """返回结算模板设计器可配置字段与默认配置"""
    return jsonify({"ok": True, "fields": SETTLE_TMPL_FIELDS,
                    "default": DEFAULT_SETTLE_CONFIG})


@app.route("/api/settle/templates", methods=["GET", "POST"])
@perm_required("settle.view")
def settle_templates():
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM settle_templates ORDER BY is_default DESC, id").fetchall()
        data = []
        for r in rows:
            d = dict(r)
            try:
                d["config"] = json.loads(d["config"] or "{}")
            except Exception:
                d["config"] = {}
            data.append(d)
        return jsonify({"ok": True, "data": data})
    if not has_perm("settle.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    name = (d.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "msg": "请填写模板名称"}), 400
    config = d.get("config") or {}
    cur = db.execute(
        "INSERT INTO settle_templates(name,is_default,config,created_by,created_at,updated_by,updated_at)"
        " VALUES(?,?,?,?,?,?,?)",
        (name, 1 if d.get("is_default") else 0, json.dumps(config, ensure_ascii=False),
         session["username"], now(), session["username"], now()))
    db.commit()
    audit("create", "settle_template", cur.lastrowid, "新建结算模板 %s" % name)
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/settle/templates/<int:tid>", methods=["PUT", "DELETE"])
@perm_required("settle.edit")
def settle_template_op(tid):
    db = get_db()
    r = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "模板不存在"}), 404
    if request.method == "DELETE":
        db.execute("DELETE FROM settle_templates WHERE id=?", (tid,))
        db.commit()
        audit("delete", "settle_template", tid, "删除结算模板 %s" % r["name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    sets, vals = [], []
    if "name" in d:
        sets.append("name=?"); vals.append((d.get("name") or "").strip())
    if "config" in d:
        sets.append("config=?"); vals.append(json.dumps(d.get("config") or {}, ensure_ascii=False))
    if "is_default" in d:
        sets.append("is_default=?"); vals.append(1 if d.get("is_default") else 0)
    if sets:
        vals += [session["username"], now(), tid]
        db.execute("UPDATE settle_templates SET %s,updated_by=?,updated_at=? WHERE id=?"
                   % ",".join(sets), vals)
        db.commit()
        audit("update", "settle_template", tid, "修改结算模板 %s" % r["name"])
    return jsonify({"ok": True})


# ================================================================ 结算模板管理（Word 四部分结构）
def _tpl_to_dict(r):
    """模板行 → dict，并解析 JSON 列（config/fields/commission_config）"""
    d = dict(r)
    for k, empty in (("config", {}), ("fields", []), ("commission_config", {})):
        try:
            v = json.loads(d.get(k) or "")
        except Exception:
            v = empty
        d[k] = v
    return d


def _stub_case():
    """空白模板渲染用的空案件数据"""
    return {"case_no": "", "maintain_user": "", "owner": "", "client_name": "",
            "plaintiff": "", "defendant": "", "section": "", "case_cause": "",
            "case_type": "", "subject_amount": 0, "project": "", "fee_standard": "",
            "expected_profit": 0, "expected_profit_text": "", "finance_remark": "",
            "confirmed_amount": 0}


def _load_fee_details(db, cid):
    """读取案件费用明细：优先 case_fee_detail，其次兼容旧 fee_recoveries 数据"""
    rows = db.execute("SELECT fee_date,description,amount FROM case_fee_detail "
                      "WHERE case_id=? ORDER BY id", (cid,)).fetchall()
    if rows:
        return [{"fee_date": r["fee_date"], "description": r["description"],
                 "amount": r["amount"]} for r in rows]
    recs = db.execute("SELECT recovery_date,item,amount FROM fee_recoveries "
                      "WHERE case_id=? ORDER BY id", (cid,)).fetchall()
    return [{"fee_date": r["recovery_date"], "description": r["item"],
             "amount": r["amount"]} for r in recs]


@app.route("/api/templates")
@perm_required("settle.view")
def api_templates_list():
    """获取模板列表（含 system 内置 (ORG_NAME)案件结案审批表）"""
    db = get_db()
    rows = db.execute("SELECT * FROM settle_templates "
                      "ORDER BY is_default DESC, is_system DESC, id").fetchall()
    return jsonify({"ok": True, "data": [_tpl_to_dict(r) for r in rows]})


@app.route("/api/templates", methods=["POST"])
@perm_required("settle.edit")
def api_templates_create():
    """新建自定义模板（按 Word 模板四部分结构：fields 第1部分、commission_config 第2部分比例）"""
    d = request.get_json(force=True)
    name = (d.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "msg": "请填写模板名称"}), 400
    ttype = d.get("type") or "salary_settlement"
    if ttype not in ("salary_settlement", "case_approval"):
        ttype = "salary_settlement"
    db = get_db()
    title = (d.get("title") or "").strip() or (
        "（ORG_NAME）案件结案审批表" if ttype == "case_approval" else name)
    fields = d.get("fields") or []
    commission_config = d.get("commission_config") or {}
    config = d.get("config")
    if config is None:
        if ttype == "case_approval":
            config = {"title": title, "type": "close_approval", "show_batch": False,
                      "show_date": False, "show_user": False,
                      "fields": [{"key": f["key"], "label": f["label"]}
                                 for f in SETTLE_TMPL_FIELDS],
                      "footer": "财务经理签字：______________    领款人签字：______________"}
        else:
            config = dict(DEFAULT_SETTLE_CONFIG)
            config["title"] = title
    cur = db.execute(
        "INSERT INTO settle_templates(name,is_default,type,is_system,title,fields,"
        "commission_config,config,created_by,created_at,updated_by,updated_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
        (name, 1 if d.get("is_default") else 0, ttype, 0, title,
         json.dumps(fields, ensure_ascii=False),
         json.dumps(commission_config, ensure_ascii=False),
         json.dumps(config, ensure_ascii=False),
         session["username"], now(), session["username"], now()))
    db.commit()
    if d.get("is_default"):
        db.execute("UPDATE settle_templates SET is_default=0 WHERE id<>?", (cur.lastrowid,))
        db.execute("UPDATE settle_templates SET is_default=1 WHERE id=?", (cur.lastrowid,))
        db.commit()
    audit("create", "settle_template", cur.lastrowid, "新建结算模板 %s（%s）" % (name, ttype))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/templates/<int:tid>", methods=["PUT", "DELETE"])
@perm_required("settle.edit")
def api_templates_op(tid):
    db = get_db()
    r = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "模板不存在"}), 404
    if request.method == "DELETE":
        if r["is_system"]:
            return jsonify({"ok": False, "msg": "系统内置模板禁止删除"}), 403
        db.execute("DELETE FROM settle_templates WHERE id=?", (tid,))
        db.commit()
        audit("delete", "settle_template", tid, "删除结算模板 %s" % r["name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    sets, vals = [], []
    if "name" in d and d.get("name") is not None:
        sets.append("name=?"); vals.append((d.get("name") or "").strip() or r["name"])
    if d.get("type") in ("salary_settlement", "case_approval"):
        sets.append("type=?"); vals.append(d["type"])
    if "title" in d:
        sets.append("title=?"); vals.append((d.get("title") or "").strip())
    if "fields" in d:
        sets.append("fields=?"); vals.append(json.dumps(d["fields"] or [], ensure_ascii=False))
    if "commission_config" in d:
        sets.append("commission_config=?"); vals.append(
            json.dumps(d["commission_config"] or {}, ensure_ascii=False))
    if "config" in d:
        sets.append("config=?"); vals.append(json.dumps(d.get("config") or {}, ensure_ascii=False))
    if "is_default" in d:
        sets.append("is_default=?"); vals.append(1 if d.get("is_default") else 0)
    if sets:
        vals += [session["username"], now(), tid]
        db.execute("UPDATE settle_templates SET %s,updated_by=?,updated_at=? WHERE id=?"
                   % ",".join(sets), vals)
        if d.get("is_default"):
            db.execute("UPDATE settle_templates SET is_default=0 WHERE id<>?", (tid,))
            db.execute("UPDATE settle_templates SET is_default=1 WHERE id=?", (tid,))
        db.commit()
        audit("update", "settle_template", tid, "修改结算模板 %s" % r["name"])
    return jsonify({"ok": True})


@app.route("/api/templates/<int:tid>/set-default", methods=["PUT"])
@perm_required("settle.edit")
def api_templates_set_default(tid):
    db = get_db()
    r = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "模板不存在"}), 404
    db.execute("UPDATE settle_templates SET is_default=0")
    db.execute("UPDATE settle_templates SET is_default=1,updated_by=?,updated_at=? WHERE id=?",
               (session["username"], now(), tid))
    db.commit()
    audit("update", "settle_template", tid, "设为默认模板 %s" % r["name"])
    return jsonify({"ok": True})


@app.route("/api/templates/<int:tid>/download")
@perm_required("settle.view")
def api_templates_download(tid):
    """下载空白 Word 模板，文件名「案件结案审批表.docx」（版式与原件一致）"""
    from io import BytesIO
    db = get_db()
    tpl_row = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    buf = BytesIO()
    render_close_approval_bundle_docx(_stub_case(), [], [], None, [], buf,
                                      fee_details=[], tpl_row=tpl_row, blank=True)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="案件结案审批表.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/api/templates/preview", methods=["POST"])
@perm_required("settle.view")
def api_templates_preview():
    """根据模板配置 + 已选案件费用明细生成临时预览 Word（不落库，返回下载流）"""
    d = request.get_json(force=True)
    db = get_db()
    tpl_row = None
    tid = d.get("template_id")
    if tid:
        tpl_row = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    config = d.get("config")
    cid = d.get("case_id")
    case = _stub_case()
    if cid:
        r = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
        if r:
            case = dict(r)
    fee_details = d.get("fee_details")
    if fee_details is None and cid:
        fee_details = _load_fee_details(db, cid)
    members = []
    if cid:
        members = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()
    if tpl_row is None and config is not None:
        from types import SimpleNamespace
        tpl_row = SimpleNamespace(
            type=(d.get("type") or "salary_settlement"),
            title=(d.get("title") or ""),
            fields=json.dumps(d.get("fields") or [], ensure_ascii=False),
            commission_config=json.dumps(d.get("commission_config") or {}, ensure_ascii=False),
            config=json.dumps(config, ensure_ascii=False))
    from io import BytesIO
    buf = BytesIO()
    render_close_approval_bundle_docx(case, [], [], config, members, buf,
                                      fee_details=fee_details or [], tpl_row=tpl_row)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="（ORG_NAME）案件结案审批表_预览.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/api/generate-settlement", methods=["POST"])
@perm_required("settle.edit")
def api_generate_settlement():
    """一键生成结算文档：每案输出一份合并 Word（审批表 + 按成员薪酬结算单），
    文档名称统一「（ORG_NAME）案件结案审批表.docx」，打包 ZIP 返回文件流
    （文件名「结算文档_YYYYMMDDHHmmss.zip」）。"""
    d = request.get_json(force=True)
    case_ids = d.get("case_ids") or []
    if not case_ids:
        return jsonify({"ok": False, "msg": "请勾选案件"}), 400
    db = get_db()
    tpl_row = None
    tid = d.get("template_id")
    if tid:
        tpl_row = db.execute("SELECT * FROM settle_templates WHERE id=?", (tid,)).fetchone()
    if tpl_row is None:
        tpl_row = db.execute("SELECT * FROM settle_templates WHERE is_default=1 LIMIT 1").fetchone()
    tpl_config = None
    if tpl_row is not None:
        try:
            tpl_config = json.loads(tpl_row["config"] or "{}")
        except Exception:
            tpl_config = None
    from io import BytesIO
    import zipfile
    zip_buf = BytesIO()
    batch_no = datetime.now().strftime("%Y%m%d%H%M%S")
    n = 0
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
        for cid in case_ids:
            case = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
            if not case:
                continue
            fee_details = _load_fee_details(db, cid)
            members = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()
            doc_buf = BytesIO()
            render_close_approval_bundle_docx(case, [], [], tpl_config, members, doc_buf,
                                              fee_details=fee_details, tpl_row=tpl_row)
            doc_buf.seek(0)
            fname = ("（ORG_NAME）案件结案审批表.docx" if len(case_ids) == 1
                     else "%s_（ORG_NAME）案件结案审批表.docx" % case["case_no"])
            z.writestr(fname, doc_buf.read())
            n += 1
    if not n:
        return jsonify({"ok": False, "msg": "所选案件未生成任何文档"}), 400
    zip_buf.seek(0)
    audit("export", "settlement_bundle", batch_no,
          "一键生成结算文档（（ORG_NAME）案件结案审批表） %d 份" % n)
    return send_file(zip_buf, as_attachment=True, download_name="结算文档_%s.zip" % batch_no,
                     mimetype="application/zip")


@app.route("/api/cases/<int:cid>/fee-details", methods=["GET", "POST"])
@perm_required("settle.view")
def case_fee_details(cid):
    """案件费用明细（第2部分数据源）：GET 查询；POST 支持 {rows:[...]} 整表替换或单行新增"""
    db = get_db()
    if not db.execute("SELECT 1 FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone():
        return jsonify({"ok": False, "msg": "案件不存在"}), 404
    if request.method == "GET":
        rows = db.execute("SELECT * FROM case_fee_detail WHERE case_id=? ORDER BY id",
                          (cid,)).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows]})
    if not has_perm("settle.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    ts = now()
    rows = d.get("rows")
    if isinstance(rows, list):
        db.execute("DELETE FROM case_fee_detail WHERE case_id=?", (cid,))
        for it in rows:
            db.execute(
                "INSERT INTO case_fee_detail(case_id,fee_date,description,amount,created_at,updated_at)"
                " VALUES(?,?,?,?,?,?)",
                (cid, (it.get("fee_date") or "").strip(), (it.get("description") or "").strip(),
                 float(it.get("amount") or 0), ts, ts))
    else:
        db.execute(
            "INSERT INTO case_fee_detail(case_id,fee_date,description,amount,created_at,updated_at)"
            " VALUES(?,?,?,?,?,?)",
            (cid, (d.get("fee_date") or "").strip(), (d.get("description") or "").strip(),
             float(d.get("amount") or 0), ts, ts))
    db.commit()
    return jsonify({"ok": True})


@app.route("/api/cases/<int:cid>/fee-details/<int:fid>", methods=["DELETE"])
@perm_required("settle.edit")
def case_fee_detail_del(cid, fid):
    db = get_db()
    db.execute("DELETE FROM case_fee_detail WHERE id=? AND case_id=?", (fid, cid))
    db.commit()
    return jsonify({"ok": True})


@app.route("/api/settlements/generate", methods=["POST"])
@perm_required("settle.edit")
def settle_generate():
    """依据案件台账 + 薪酬计算结果，批量生成 Word 结算单，打包 ZIP 返回"""
    d = request.get_json(force=True)
    case_ids = d.get("case_ids") or []
    if not case_ids:
        return jsonify({"ok": False, "msg": "请选择案件"}), 400
    db = get_db()
    batch_no = "JS" + datetime.now().strftime("%Y%m%d%H%M%S")
    tpl = get_setting("settle_template")
    tpl_path = os.path.join(TEMPLATE_DIR, tpl) if tpl else None
    # 结算模板：优先用显式指定 template_id，其次默认模板，最后 Word 文件/内置
    tpl_config = None
    template_id = d.get("template_id")
    if template_id:
        tr = db.execute("SELECT * FROM settle_templates WHERE id=?", (template_id,)).fetchone()
        if tr:
            try:
                tpl_config = json.loads(tr["config"] or "{}")
            except Exception:
                tpl_config = None
    if tpl_config is None:
        dr = db.execute("SELECT * FROM settle_templates WHERE is_default=1 LIMIT 1").fetchone()
        if dr:
            try:
                tpl_config = json.loads(dr["config"] or "{}")
            except Exception:
                tpl_config = None

    # 固定金额仅在案件成员(case_members)中维护

    files = []
    for cid in case_ids:
        c = db.execute("SELECT * FROM cases WHERE id=? AND deleted=0", (cid,)).fetchone()
        if not c:
            continue
        recovered = float(c["confirmed_amount"] or 0)
        if recovered <= 0:
            rec = db.execute(
                "SELECT COALESCE(SUM(amount),0) s FROM fee_recoveries WHERE case_id=? AND confirm=1",
                (cid,)).fetchone()["s"]
            recovered = float(rec or 0)
        members = db.execute("SELECT * FROM case_members WHERE case_id=?", (cid,)).fetchall()
        for m in members:
            rate = float(m["rate"] or 0)
            # 保底金额：成员自定义优先，否则按角色默认（维护人 0，其他 600）
            gua = m["guaranteed_amount"]
            gua = default_guaranteed(m["role"]) if gua is None else float(gua or 0)
            # 固定金额：仅取自案件成员配置
            fixed = float(m["fixed_amount"] or 0)
            # 统一薪酬公式：max(到账金额 × 提成率, 保底金额) + 固定金额
            base = max(recovered * rate, gua)
            amount = round_half_up(base + fixed)   # 四舍五入取整
            if amount <= 0:
                continue
            fname = "%s_%s_%s.docx" % (batch_no, c["case_no"], m["person"])
            fpath = os.path.join(SETTLE_DIR, fname)
            ctx = {k: c[k] if k in c.keys() else "" for k in SETTLE_FIELDS}
            ctx.update(user_name=m["person"], amount=amount, batch_no=batch_no,
                       date="", firm=get_setting("firm_name"))  # date 留白，由用户自行填写
            if tpl_config:
                render_settlement_from_config(tpl_config, fpath, ctx)
            else:
                render_settlement_docx(tpl_path, fpath, ctx)
            db.execute(
                "INSERT INTO settlements(batch_no,case_id,case_no,user_name,amount,doc_path,created_by,created_at)"
                " VALUES(?,?,?,?,?,?,?,?)",
                (batch_no, cid, c["case_no"], m["person"], amount, fpath,
                 session["username"], now()))
            files.append(fpath)
    db.commit()
    audit("create", "settlement", batch_no,
          "批量生成结算单 %d 份，案件数 %d" % (len(files), len(case_ids)))
    if not files:
        return jsonify({"ok": False, "msg": "所选案件未匹配到任何成员或回款金额，未生成结算单"}), 400
    # 打包 zip
    import zipfile
    zip_path = os.path.join(EXPORT_DIR, batch_no + ".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            z.write(f, os.path.basename(f))
    return jsonify({"ok": True, "batch_no": batch_no, "count": len(files),
                    "download": "/api/settlements/download/" + batch_no})


@app.route("/api/settlements/download/<batch_no>")
@perm_required("settle.view")
def settle_download(batch_no):
    if not re.fullmatch(r"JS\d+", batch_no):
        return jsonify({"ok": False, "msg": "非法批次号"}), 400
    zip_path = os.path.join(EXPORT_DIR, batch_no + ".zip")
    if not os.path.exists(zip_path):
        return jsonify({"ok": False, "msg": "文件不存在"}), 404
    audit("export", "settlement", batch_no, "下载结算单批次")
    return send_file(zip_path, as_attachment=True)


@app.route("/api/settlements")
@perm_required("settle.view")
def settle_list():
    """历史结算记录查询。支持按结算月（created_at YYYY-MM）、案号、人员名称模糊过滤。
    返回结果按 id DESC 顺序，便于最新批次优先展示。"""
    db = get_db()
    month = (request.args.get("month") or "").strip()
    kw = (request.args.get("kw") or "").strip()
    sql = "SELECT * FROM settlements WHERE 1=1"
    args = []
    if month:
        sql += " AND substr(created_at,1,7)=?"
        args.append(month)
    if kw:
        sql += " AND (case_no LIKE ? OR user_name LIKE ?)"
        args.extend(["%" + kw + "%"] * 2)
    sql += " ORDER BY id DESC LIMIT 1000"
    rows = db.execute(sql, args).fetchall()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


@app.route("/api/settlements/export")
@perm_required("settle.view")
def settle_export():
    """历史结算记录导出 CSV（与列表查询共用 month/kw 过滤条件）"""
    import csv as _csv
    db = get_db()
    month = (request.args.get("month") or "").strip()
    kw = (request.args.get("kw") or "").strip()
    sql = "SELECT * FROM settlements WHERE 1=1"
    args = []
    if month:
        sql += " AND substr(created_at,1,7)=?"
        args.append(month)
    if kw:
        sql += " AND (case_no LIKE ? OR user_name LIKE ?)"
        args.extend(["%" + kw + "%"] * 2)
    sql += " ORDER BY id DESC LIMIT 5000"
    rows = db.execute(sql, args).fetchall()
    fn = "settlements_{}_{}.csv".format(month or "all", now()[:19].replace(":", "").replace(" ", "_"))
    p = os.path.join(EXPORT_DIR, fn)
    with open(p, "w", encoding="utf-8-sig", newline="") as f:
        w = _csv.writer(f)
        w.writerow(["批次号", "案号", "人员", "金额(¥)", "文档路径", "生成人", "生成时间"])
        for r in rows:
            w.writerow([r["batch_no"], r["case_no"], r["user_name"],
                        r["amount"], r["doc_path"], r["created_by"], r["created_at"]])
    return jsonify({"ok": True, "download": "/api/settlements/csv-download/" + fn,
                    "count": len(rows)})


@app.route("/api/settlements/csv-download/<path:filename>")
def settle_csv_download(filename):
    """CSV 导出下载（与 export 接口配对）。使用 path 转换器允许文件名含下划线/连字符。"""
    from flask import send_file
    return send_file(os.path.join(EXPORT_DIR, filename), as_attachment=True)


# ================================================================ 导入导出
IMPORT_MAP = {"案号": "case_no", "客户名称": "client_name", "案件开发人": "developer",
              "负责人": "owner", "小组成员": "members", "案件类型": "case_type",
              "状态": "status", "代理费金额": "agent_fee", "代垫诉讼费": "court_fee_advance",
              "开票金额": "invoice_amount", "备注": "remark"}


# ---------------------------------------------------------------- 通用表格导入工具
def parse_tabular(file_bytes, filename):
    """解析 Excel/CSV，返回 (header, rows)。CSV 自动识别 UTF-8/GBK 编码"""
    name = (filename or "").lower()
    if name.endswith(".csv"):
        text = None
        for enc in ("utf-8-sig", "utf-8", "gbk"):
            try:
                text = file_bytes.decode(enc)
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if text is None:
            text = file_bytes.decode("gbk", errors="ignore")
        rows = [r for r in csv.reader(io.StringIO(text))]
        if not rows:
            return [], []
        return [str(h).strip() for h in rows[0]], rows[1:]
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    # 处理合并单元格：openpyxl 仅左上角有值，此处将左上值铺满整个合并区域，
    # 使「合并单元格内容」能随映射字段正确匹配入库
    merged_fill = {}
    for mr in ws.merged_cells.ranges:
        tl = ws.cell(row=mr.min_row, column=mr.min_col).value
        for r in range(mr.min_row, mr.max_row + 1):
            for c in range(mr.min_col, mr.max_col + 1):
                merged_fill[(r, c)] = tl
    rows = []
    for row in ws.iter_rows(values_only=False):
        row_vals = []
        for cell in row:
            v = cell.value
            if v is None and (cell.row, cell.column) in merged_fill:
                v = merged_fill[(cell.row, cell.column)]
            row_vals.append(v)
        rows.append(row_vals)
    if not rows:
        return [], []
    header = [str(h).strip() if h is not None else "" for h in rows[0]]
    return header, rows[1:]


def clean_value(v, ftype):
    """按类型清洗单元格：数字去逗号/¥，日期统一为 YYYY-MM-DD"""
    if v is None:
        return 0.0 if ftype == "num" else ""
    if isinstance(v, (datetime, date)):
        return v.strftime("%Y-%m-%d") if ftype == "date" else str(v)
    if ftype == "num":
        s = str(v).replace(",", "").replace("¥", "").replace("￥", "").replace("元", "").strip()
        try:
            return float(s)
        except ValueError:
            return v
    return str(v).strip()


def normalize_date(v):
    """日期标准化，失败返回 None"""
    if isinstance(v, (datetime, date)):
        return v.strftime("%Y-%m-%d")
    s = str(v).strip()
    m = re.match(r"^(\d{4})[-/年](\d{1,2})[-/月](\d{1,2})日?$", s)
    if m:
        try:
            return "%s-%02d-%02d" % (m.group(1), int(m.group(2)), int(m.group(3)))
        except ValueError:
            return None
    return None


def build_col_map(header, entity, mapping=None):
    """构建 {目标字段: 列索引}。mapping={源列名:目标字段} 优先，否则按字段中文名/英文键自动匹配"""
    fields = entity_fields(entity)
    label2key = {f["label"]: f["key"] for f in fields}
    key_set = {f["key"] for f in fields}
    col_map = {}
    for i, h in enumerate(header):
        h = (h or "").strip()
        if not h:
            continue
        if mapping and h in mapping and mapping[h] in key_set:
            col_map[mapping[h]] = i
        elif h in label2key:
            col_map[label2key[h]] = i
        elif h in key_set:
            col_map[h] = i
    return col_map


def validate_entity_row(entity, raw_row, col_map):
    """清洗 + 校验一行数据，返回 (data, errors)。校验必填项、数字格式、日期格式"""
    fields = entity_fields(entity)
    data, errors = {}, []
    for f in fields:
        key = f["key"]
        if key not in col_map:
            continue
        idx = col_map[key]
        v = raw_row[idx] if idx < len(raw_row) else ""
        v = clean_value(v, f["type"])
        if f["type"] == "date" and v not in ("", None):
            d = normalize_date(v)
            if d is None:
                errors.append("「%s」日期格式无效：%s" % (f["label"], v))
            else:
                v = d
        if f["type"] == "num" and not isinstance(v, (int, float)):
            errors.append("「%s」必须为数字：%s" % (f["label"], v))
            v = 0.0
        if f["required"] and (v in ("", None)):
            errors.append("「%s」为必填项，不能为空" % f["label"])
        data[key] = v
    return data, errors


# ---------------------------------------------------------------- 导入模板管理
@app.route("/api/import/fields/<entity>")
@perm_required("file.import")
def import_fields(entity):
    """返回实体字段定义（供导入模板创建界面展示可选字段）"""
    if entity not in ENTITY_DEFS:
        return jsonify({"ok": False, "msg": "未知实体类型"}), 400
    return jsonify({"ok": True, "name": ENTITY_DEFS[entity]["name"],
                    "fields": ENTITY_DEFS[entity]["fields"]})


@app.route("/api/import/templates", methods=["GET", "POST"])
@perm_required("file.import")
def import_templates():
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT * FROM import_templates ORDER BY id DESC").fetchall()
        data = []
        for r in rows:
            d = dict(r)
            try:
                d["mapping"] = json.loads(d["mapping"] or "{}")
            except Exception:
                d["mapping"] = {}
            d["entity_name"] = ENTITY_DEFS.get(d["entity"], {}).get("name", d["entity"])
            data.append(d)
        return jsonify({"ok": True, "data": data, "entities": {k: v["name"] for k, v in ENTITY_DEFS.items()}})
    d = request.get_json(force=True)
    name = (d.get("name") or "").strip()
    entity = d.get("entity") or "case"
    if not name:
        return jsonify({"ok": False, "msg": "请填写模板名称"}), 400
    if entity not in ENTITY_DEFS:
        return jsonify({"ok": False, "msg": "未知实体类型"}), 400
    mapping = d.get("mapping") or {}
    if not mapping:
        return jsonify({"ok": False, "msg": "请至少配置一个字段映射"}), 400
    cur = db.execute(
        "INSERT INTO import_templates(name,entity,mapping,created_by,created_at)"
        " VALUES(?,?,?,?,?)",
        (name, entity, json.dumps(mapping, ensure_ascii=False), session["username"], now()))
    db.commit()
    audit("create", "import_template", cur.lastrowid, "新建导入模板 %s（实体=%s）" % (name, entity))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/import/templates/<int:tid>", methods=["PUT", "DELETE"])
@perm_required("file.import")
def import_template_op(tid):
    db = get_db()
    r = db.execute("SELECT * FROM import_templates WHERE id=?", (tid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "模板不存在"}), 404
    if request.method == "DELETE":
        db.execute("DELETE FROM import_templates WHERE id=?", (tid,))
        db.commit()
        audit("delete", "import_template", tid, "删除导入模板 %s" % r["name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    sets, vals = [], []
    if "name" in d:
        sets.append("name=?"); vals.append((d.get("name") or "").strip())
    if "mapping" in d:
        sets.append("mapping=?"); vals.append(json.dumps(d.get("mapping") or {}, ensure_ascii=False))
    if sets:
        vals.append(tid)
        db.execute("UPDATE import_templates SET %s WHERE id=?" % ",".join(sets), vals)
        db.commit()
        audit("update", "import_template", tid, "修改导入模板 %s" % r["name"])
    return jsonify({"ok": True})


# ---------------------------------------------------------------- 人员花名册
@app.route("/api/staff", methods=["GET", "POST"])
@perm_required("case.view")
def staff_list():
    db = get_db()
    if request.method == "GET":
        kw = request.args.get("kw", "").strip()
        sql = "SELECT * FROM staff WHERE deleted=0"
        args = []
        if kw:
            sql += " AND (name LIKE ? OR staff_no LIKE ? OR department LIKE ? OR position LIKE ?)"
            args = ["%" + kw + "%"] * 4
        sql += " ORDER BY id DESC"
        rows = db.execute(sql, args).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows],
                        "fields": entity_fields("staff")})
    if not has_perm("case.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    if not (d.get("name") or "").strip():
        return jsonify({"ok": False, "msg": "姓名必填"}), 400
    cur = db.execute(
        "INSERT INTO staff(staff_no,name,department,position,hire_date,phone,id_card,status,commission_rate,remark,created_by,created_at,updated_by,updated_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (d.get("staff_no", "").strip(), d["name"].strip(), d.get("department", ""),
         d.get("position", ""), normalize_date(d.get("hire_date")) or "",
         d.get("phone", ""), d.get("id_card", ""), d.get("status") or "在职",
         float(d.get("commission_rate") or 0),
         d.get("remark", ""), session["username"], now(), session["username"], now()))
    db.commit()
    audit("create", "staff", cur.lastrowid, "新增人员 %s 提成率 %.2f%%"
          % (d["name"], float(d.get("commission_rate") or 0) * 100))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/staff/<int:sid>", methods=["PUT", "DELETE"])
@perm_required("case.edit")
def staff_op(sid):
    db = get_db()
    r = db.execute("SELECT * FROM staff WHERE id=? AND deleted=0", (sid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "人员不存在"}), 404
    if request.method == "DELETE":
        db.execute("UPDATE staff SET deleted=1, updated_at=? WHERE id=?", (now(), sid))
        db.commit()
        audit("delete", "staff", sid, "删除人员 %s" % r["name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    fields = {k: d[k] for k in ("staff_no", "name", "department", "position", "phone",
                                "id_card", "status", "remark") if k in d}
    if "commission_rate" in d:
        fields["commission_rate"] = float(d.get("commission_rate") or 0)
    if "hire_date" in d:
        fields["hire_date"] = normalize_date(d.get("hire_date")) or ""
    if not (fields.get("name") or r["name"]):
        return jsonify({"ok": False, "msg": "姓名必填"}), 400
    if fields:
        sets = ",".join(k + "=?" for k in fields)
        db.execute("UPDATE staff SET %s,updated_by=?,updated_at=? WHERE id=?"
                   % sets, list(fields.values()) + [session["username"], now(), sid])
        db.commit()
        audit("update", "staff", sid, "修改人员 %s: %s" % (r["name"], ",".join(fields.keys())))
    return jsonify({"ok": True})


@app.route("/api/staff/import", methods=["POST"])
@perm_required("file.import")
def staff_import():
    return _generic_import("staff", "花名册")


@app.route("/api/staff/batch-rate", methods=["POST"])
@perm_required("case.edit")
def staff_batch_rate():
    """按部门批量设置案件提成率"""
    d = request.get_json(force=True)
    department = (d.get("department") or "").strip()
    if not department:
        return jsonify({"ok": False, "msg": "请选择部门"}), 400
    rate = float(d.get("rate") or 0)
    db = get_db()
    cur = db.execute("UPDATE staff SET commission_rate=? WHERE department=? AND deleted=0",
                     (rate, department))
    db.commit()
    audit("update", "staff", department, "按部门批量设置提成率 %.2f%%，影响 %d 人"
          % (rate * 100, cur.rowcount))
    return jsonify({"ok": True, "count": cur.rowcount})


# ---------------------------------------------------------------- 客户信息
@app.route("/api/customers", methods=["GET", "POST"])
@perm_required("case.view")
def customers_list():
    db = get_db()
    if request.method == "GET":
        kw = request.args.get("kw", "").strip()
        sql = "SELECT * FROM customers WHERE deleted=0"
        args = []
        if kw:
            sql += " AND (name LIKE ? OR contact LIKE ? OR phone LIKE ? OR credit_code LIKE ?)"
            args = ["%" + kw + "%"] * 4
        sql += " ORDER BY id DESC"
        rows = db.execute(sql, args).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows],
                        "fields": entity_fields("customer")})
    if not has_perm("case.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    if not (d.get("name") or "").strip():
        return jsonify({"ok": False, "msg": "客户名称必填"}), 400
    cur = db.execute(
        "INSERT INTO customers(name,credit_code,contact,phone,cust_type,industry,address,remark,created_by,created_at,updated_by,updated_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
        (d["name"].strip(), d.get("credit_code", ""), d.get("contact", ""),
         d.get("phone", ""), d.get("cust_type", ""), d.get("industry", ""),
         d.get("address", ""), d.get("remark", ""), session["username"], now(),
         session["username"], now()))
    db.commit()
    audit("create", "customer", cur.lastrowid, "新增客户 %s" % d["name"])
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/customers/<int:cid>", methods=["PUT", "DELETE"])
@perm_required("case.edit")
def customers_op(cid):
    db = get_db()
    r = db.execute("SELECT * FROM customers WHERE id=? AND deleted=0", (cid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "客户不存在"}), 404
    if request.method == "DELETE":
        db.execute("UPDATE customers SET deleted=1, updated_at=? WHERE id=?", (now(), cid))
        db.commit()
        audit("delete", "customer", cid, "删除客户 %s" % r["name"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    fields = {k: d[k] for k in ("name", "credit_code", "contact", "phone", "cust_type",
                                "industry", "address", "remark") if k in d}
    if not (fields.get("name") or r["name"]):
        return jsonify({"ok": False, "msg": "客户名称必填"}), 400
    if fields:
        sets = ",".join(k + "=?" for k in fields)
        db.execute("UPDATE customers SET %s,updated_by=?,updated_at=? WHERE id=?"
                   % sets, list(fields.values()) + [session["username"], now(), cid])
        db.commit()
        audit("update", "customer", cid, "修改客户 %s: %s" % (r["name"], ",".join(fields.keys())))
    return jsonify({"ok": True})


@app.route("/api/customers/import", methods=["POST"])
@perm_required("file.import")
def customers_import():
    return _generic_import("customer", "客户信息")


def _generic_import(entity, entity_label):
    """通用批量导入：Excel/CSV → 字段映射（可指定导入模板）→ 逐行校验 → 去重入库 → 错误详情"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "msg": "请选择文件"}), 400
    file_bytes = f.read()
    # 指定导入模板时读取其字段映射
    mapping = None
    tid = request.form.get("template_id")
    if tid:
        tr = get_db().execute("SELECT * FROM import_templates WHERE id=?", (tid,)).fetchone()
        if tr and tr["entity"] == entity:
            try:
                mapping = json.loads(tr["mapping"] or "{}")
            except Exception:
                mapping = None
    header, rows = parse_tabular(file_bytes, f.filename)
    if not header:
        return jsonify({"ok": False, "msg": "文件无表头或无数据"}), 400
    fields = entity_fields(entity)
    required = [x["key"] for x in fields if x["required"]]
    col_map = build_col_map(header, entity, mapping)
    # 必填列检查
    missing = [next(x["label"] for x in fields if x["key"] == k) for k in required
               if k not in col_map]
    if missing:
        return jsonify({"ok": False, "msg": "缺少必需列：%s（请检查表头或字段映射）" % "、".join(missing)}), 400

    db = get_db()
    inserted, skipped, errors = 0, 0, []
    req_key = required[0] if required else None
    for idx, raw in enumerate(rows, start=2):
        data, row_errs = validate_entity_row(entity, raw, col_map)
        if row_errs:
            errors.append("第%d行：%s" % (idx, "；".join(row_errs)))
            skipped += 1
            continue
        if not data.get(req_key):
            skipped += 1
            continue
        # 去重
        if entity == "staff":
            dup = db.execute("SELECT 1 FROM staff WHERE deleted=0 AND name=? AND staff_no=?",
                             (data["name"], data.get("staff_no", ""))).fetchone()
        else:  # customer
            dup = db.execute("SELECT 1 FROM customers WHERE deleted=0 AND name=?",
                             (data["name"],)).fetchone()
        if dup:
            skipped += 1
            continue
        try:
            if entity == "staff":
                db.execute(
                    "INSERT INTO staff(staff_no,name,department,position,hire_date,phone,id_card,status,commission_rate,remark,created_by,created_at,updated_by,updated_at)"
                    " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (data.get("staff_no", ""), data["name"], data.get("department", ""),
                     data.get("position", ""), data.get("hire_date", ""), data.get("phone", ""),
                     data.get("id_card", ""), data.get("status") or "在职",
                     float(data.get("commission_rate") or 0), data.get("remark", ""),
                     session["username"], now(), session["username"], now()))
            else:
                db.execute(
                    "INSERT INTO customers(name,credit_code,contact,phone,cust_type,industry,address,remark,created_by,created_at,updated_by,updated_at)"
                    " VALUES(?,?,?,?,?,?,?,?,?,?,?,?)",
                    (data["name"], data.get("credit_code", ""), data.get("contact", ""),
                     data.get("phone", ""), data.get("cust_type", ""), data.get("industry", ""),
                     data.get("address", ""), data.get("remark", ""), session["username"], now(),
                     session["username"], now()))
            inserted += 1
        except Exception as e:
            errors.append("第%d行：%s" % (idx, e))
            skipped += 1
    db.commit()
    audit("import", entity, "", "导入%s：新增 %d 条，跳过 %d 条，错误 %d 条"
          % (entity_label, inserted, skipped, len(errors)))
    return jsonify({"ok": True, "inserted": inserted, "skipped": skipped,
                    "errors": errors[:50], "total_errors": len(errors)})


@app.route("/api/cases/import", methods=["POST"])
@perm_required("file.import")
def cases_import():
    """导入历史数据：Excel/CSV → 数据清洗 → 去重 → 入库。
    支持指定导入模板（template_id）使用已保存的字段映射。"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "msg": "请选择文件"}), 400
    file_bytes = f.read()
    header, body_rows = parse_tabular(file_bytes, f.filename)
    if len(body_rows) < 1:
        return jsonify({"ok": False, "msg": "文件无数据行"}), 400
    # 导入模板字段映射（可选）
    mapping = None
    tid = request.form.get("template_id")
    if tid:
        tr = get_db().execute("SELECT * FROM import_templates WHERE id=?", (tid,)).fetchone()
        if tr and tr["entity"] == "case":
            try:
                mapping = json.loads(tr["mapping"] or "{}")
            except Exception:
                mapping = None
    if mapping:
        key_set = {x["key"] for x in entity_fields("case")}
        col = {}
        for i, h in enumerate(header):
            if h in mapping and mapping[h] in key_set:
                col[mapping[h]] = i
    else:
        col = {IMPORT_MAP[h]: i for i, h in enumerate(header) if h in IMPORT_MAP}
    if "client_name" not in col:
        return jsonify({"ok": False, "msg": "缺少必需列：客户名称"}), 400

    db = get_db()
    inserted, skipped, errors = 0, 0, []

    def clean_num(v):
        if v is None or v == "":
            return 0.0
        try:
            return float(str(v).replace(",", "").replace("¥", "").strip())
        except ValueError:
            return 0.0

    def clean_str(v):
        return str(v).strip() if v is not None else ""

    for idx, r in enumerate(body_rows, start=2):
        client = clean_str(r[col["client_name"]]) if col["client_name"] < len(r) else ""
        if not client:
            skipped += 1
            continue
        data = {k: (clean_num(r[i]) if k in ("agent_fee", "court_fee_advance",
                                             "invoice_amount") and i < len(r)
                    else clean_str(r[i]) if i < len(r) else "")
                for k, i in col.items()}
        # 去重：已有案号跳过；无案号按 客户+开发人+负责人 判重
        if data.get("case_no") and db.execute(
                "SELECT 1 FROM cases WHERE case_no=?", (data["case_no"],)).fetchone():
            skipped += 1
            continue
        if not data.get("case_no"):
            data["case_no"] = next_case_no(db)
            if db.execute("SELECT 1 FROM cases WHERE client_name=? AND developer=? AND owner=?",
                          (client, data.get("developer", ""),
                           data.get("owner", ""))).fetchone():
                skipped += 1
                continue
        try:
            db.execute(
                """INSERT INTO cases(case_no,developer,owner,members,client_name,case_type,status,
                   agent_fee,court_fee_advance,invoice_amount,remark,created_by,created_at,updated_by,updated_at)
                   VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (data["case_no"], data.get("developer", ""), data.get("owner", ""),
                 data.get("members", ""), client, data.get("case_type", ""),
                 data.get("status") or "进行中", data.get("agent_fee", 0),
                 data.get("court_fee_advance", 0), data.get("invoice_amount", 0),
                 data.get("remark", ""), session["username"], now(),
                 session["username"], now()))
            inserted += 1
        except Exception as e:
            errors.append("第%d行: %s" % (idx, e))
    db.commit()
    audit("import", "case", "", "导入 Excel：新增 %d 条，跳过 %d 条，错误 %d 条"
          % (inserted, skipped, len(errors)))
    return jsonify({"ok": True, "inserted": inserted, "skipped": skipped,
                    "errors": errors[:20]})


@app.route("/api/cases/export")
@perm_required("file.export")
def cases_export():
    from openpyxl import Workbook
    db = get_db()
    rows = db.execute("SELECT * FROM cases WHERE deleted=0 ORDER BY id").fetchall()
    wb = Workbook()
    ws = wb.active
    ws.title = "案件台账"
    heads = list(IMPORT_MAP.keys())
    ws.append(heads + ["录入人", "更新时间"])
    for r in rows:
        ws.append([r[v] for v in IMPORT_MAP.values()] + [r["updated_by"], r["updated_at"]])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    audit("export", "case", "", "导出案件台账 Excel 共 %d 条" % len(rows))
    return send_file(buf, as_attachment=True,
                     download_name="案件台账_%s.xlsx" % date.today().isoformat(),
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ================================================================ 本地离线 OCR
# 基于 RapidOCR（ONNX Runtime + PP-OCRv3 模型），模型文件随包内置，全程离线，
# 发票图片数据不出本机，保障隐私安全。首次使用若未安装依赖可调用 /api/ocr/setup 自动安装。
_OCR_ENGINE = None
_OCR_LOCK = threading.Lock()
_OCR_ERR = None


def ocr_available():
    """检测本地 OCR 依赖是否可导入"""
    try:
        import rapidocr_onnxruntime  # noqa: F401
        return True
    except Exception:
        return False


def get_ocr_engine():
    """懒加载 OCR 引擎（线程安全），返回 (engine, err)"""
    global _OCR_ENGINE, _OCR_ERR
    if _OCR_ENGINE is None and _OCR_ERR is None:
        with _OCR_LOCK:
            if _OCR_ENGINE is None and _OCR_ERR is None:
                try:
                    from rapidocr_onnxruntime import RapidOCR
                    _OCR_ENGINE = RapidOCR()
                except Exception as e:
                    _OCR_ERR = str(e)
    return _OCR_ENGINE, _OCR_ERR


def ocr_read_text(img_bytes):
    """OCR 识别图片，返回识别文本列表（按位置排序）"""
    import numpy as np
    from PIL import Image
    engine, err = get_ocr_engine()
    if err:
        raise RuntimeError("OCR 引擎初始化失败: %s" % err)
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    arr = np.array(img)
    result, _ = engine(arr)
    # RapidOCR 返回 [[box, text, score], ...]
    texts = []
    if isinstance(result, list):
        for item in result:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                texts.append(str(item[1]))
            elif hasattr(item, "txts"):
                texts.append(str(item.txts))
    elif hasattr(result, "txts"):
        texts = [str(t) for t in result.txts]
    return texts


def pdfplumber_available():
    """检测 pdfplumber 是否可导入（文本型电子发票精确解析引擎）"""
    try:
        import pdfplumber  # noqa: F401
        return True
    except Exception:
        return False


def pdf_extract_text_pp(pdf_bytes):
    """用 pdfplumber 提取 PDF 文本层（逐页按行序，电子发票标签与值分离版式解析更稳）。
    返回非空行列表；无文本层（扫描件）返回 []。"""
    import pdfplumber
    lines = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            for line in txt.splitlines():
                line = line.strip()
                if line:
                    lines.append(line)
    return lines


def pdf_to_images(pdf_bytes):
    """PDF 转图片列表（每页一张 PNG），2x 放大提升 OCR 清晰度"""
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    imgs = []
    for page in doc:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        imgs.append(pix.tobytes("png"))
    doc.close()
    return imgs


def pdf_extract_text(pdf_bytes):
    """提取 PDF 文本层（文本型电子发票可直接提取，准确率远高于 OCR）。
    解析引擎优先级：pdfplumber（精确）→ PyMuPDF（兜底）；
    均以「有效字符数（中文+字母数字）≥10」为提取成功判据，防字体乱码导致错误文本；
    均无文本层返回 [] 由调用方转图 OCR。"""
    def _valid(lines):
        return len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", "".join(lines or [])))
    if pdfplumber_available():
        try:
            lines = pdf_extract_text_pp(pdf_bytes)
            if lines and _valid(lines) >= 10:
                return lines
        except Exception:
            pass  # pdfplumber 解析异常 → 回退 PyMuPDF
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    lines = []
    for page in doc:
        txt = page.get_text()
        for line in txt.splitlines():
            line = line.strip()
            if line:
                lines.append(line)
    doc.close()
    if _valid(lines) >= 10:
        return lines
    return []


def ocr_read_file(file_bytes, filename):
    """识别文件（图片或 PDF）。PDF 优先提取文本层（电子发票），扫描件才转图 OCR"""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        # 文本型 PDF：直接提取文本层，金额/号码零误差
        lines = pdf_extract_text(file_bytes)
        if lines and len("".join(lines)) >= 20:
            return lines
        # 扫描件 PDF：转图 OCR
        texts = []
        for img in pdf_to_images(file_bytes):
            texts.extend(ocr_read_text(img))
        return texts
    return ocr_read_text(file_bytes)


def fix_ocr_number(s):
    """修正 OCR 金额数字的常见误识别：O→0、l/I→1、全角符号转半角"""
    if not s:
        return s
    s = str(s).replace("O", "0").replace("o", "0") \
        .replace("l", "1").replace("I", "1") \
        .replace("，", ",").replace("．", ".").replace("。", ".")
    return s


# 发票费用科目关键词归集（默认科目 + 可扩展）
CATEGORY_KEYWORDS = {
    "诉讼费": ["诉讼", "案件受理", "法院", "立案"],
    "律师代理费": ["律师", "代理费", "法律服务", "咨询", "顾问", "诉讼代理"],
    "差旅费": ["差旅", "交通", "住宿", "机票", "火车", "动车", "打车", "出租车", "过路"],
    "公证费": ["公证"],
    "鉴定费": ["鉴定", "检测", "检验", "评估"],
    "办公费": ["办公", "文具", "打印", "复印", "耗材", "纸张", "电脑", "墨盒"],
}


def categorize_invoice(text, cats):
    """按发票品名/全文关键词归集费用科目，返回最匹配科目，无匹配归「其他」"""
    m = re.search(r"(?:货物或应税劳务、服务名称|项目名称|货物名称)[：:\s]*([^\n]+)", text)
    target = (m.group(1) if m else "") + "  " + text
    for cat in cats:
        if not cat or cat == "其他":
            continue
        kws = CATEGORY_KEYWORDS.get(cat, [cat])  # 自定义科目以科目名本身为关键词
        if any(k in target for k in kws):
            return cat
    return "其他" if "其他" in cats else (cats[0] if cats else "")


def detect_invoice_type(full):
    """从发票全文识别发票类型，返回票面标题原文（如「电子发票（增值税专用发票）」），
    与《发票AI识别提示词（整合版）》invoice_type 字段规范一致。
    兼容 OCR 丢字/换行：去空白后先匹配完整关键词，再匹配宽松关键词。返回 '' 表示未能识别。"""
    c = re.sub(r"\s+", "", full or "")
    # 完整票面标题优先（电子发票专用/普通）
    m = re.search(r"电子发票[（(]增值税[专普]用发票[)）]", c)
    if m:
        return m.group(0)
    if "增值税专用发票" in c or "专用发票" in c:
        return "增值税专用发票"
    if "增值税普通发票" in c or "普通发票" in c:
        return "增值税普通发票"
    return ""


def extract_item_name(full):
    """提取发票「项目名称」（含货物或应税劳务、服务名称/项目名称/货物名称标签）。
    返回去除标签后的项目名称原文（含税收分类编码前缀，如 *生产生活服务*法律咨询）。
    无标签时回退匹配行首为「*分类编码*」的货物明细行（表头标签在上、明细行直接以星号开头）。"""
    if not full:
        return ""
    # 优先：明细行以 *分类编码* 开头（规范要求项目名称必须含税收分类编码前缀），
    # 名称主体到第一个空白即止（金额/税率/税额等列与名称同行）
    m = re.search(r"^\*[^*\n]{1,40}\*[^\s\n]{0,60}", full, re.M)
    if m:
        return m.group(0).strip()
    # 兜底：带冒号的标签式（项目名称：xxx）
    m = re.search(r"(?:货物或应税劳务、服务名称|项目名称|货物名称)[：:]\s*([^\n]{1,80})", full)
    if m:
        return m.group(1).strip("：: ")
    return ""


def item_name_to_category(item_name):
    """项目名称 → 费用科目：去掉税收分类编码前缀（*xxx*）取主体；
    无前缀则直接用项目名称本身。空返回 ''。"""
    if not item_name:
        return ""
    cat = re.sub(r"^\*[^*]+\*", "", item_name).strip()
    return cat if cat else item_name.strip()


def parse_invoice_fields(texts, cats):
    """从发票文本（OCR 或 PDF 文本层）提取结构化字段。
    兼容文本型电子发票（标签与值分离）与 OCR 图片发票。"""
    full = "\n".join(texts)
    compact = re.sub(r"\s+", "", full)

    def _money(s):
        try:
            return float(fix_ocr_number(s).replace(",", ""))
        except (ValueError, TypeError):
            return None

    data = {"invoice_no": "", "buyer": "", "seller": "", "amount": 0.0, "tax_amount": 0.0,
            "invoice_date": "", "invoice_type": "", "item_name": "", "category": "",
            "total_amount": 0.0, "total_amount_cn": "", "tax_rate": "",
            "buyer_tax_id": "", "seller_tax_id": "", "drawer": "", "remark": "",
            "raw": full[:2000]}

    # 发票号码：20 位数字，标签与值可能分离（电子发票号码紧跟「开票人」）
    m = re.search(r"发票号码[：:\s]*([0-9]{8,20})", compact)
    if not m:
        m = re.search(r"开票人[：:]?[^0-9]{0,3}([0-9]{20})", compact)
    if not m:
        m = re.search(r"(?<![0-9])([0-9]{20})", compact)  # 仅要求前缀非数字，避免误匹配银行账号
    if m:
        data["invoice_no"] = m.group(1)

    # 购买方名称（发票「购 买 方 / 买方」区「名 称：」后完整内容）→ 校正为客户名称（client_name）
    # 注：发票排版中「名 称：」常为字间空格写法，用 名\s*称 容错匹配；
    # 竖排拆行（购/买/方 分列各行）或两列并排（购 名称：X 销 名称：Y 同行）时退化按「购…名称：」匹配
    m = re.search(r"购\s*买\s*方[^名\n]{0,20}名\s*称[：:\s]*([^\n]{2,40})", full)
    if not m:
        # 两列并排：购 名称：X 销 名称：Y 同行 → 非贪婪截到「销 名」前（销 前须有空白，防名称内含「销」字截断）
        m = re.search(r"购[^名\n]{0,6}名\s*称[：:]\s*([^\n]{2,40}?)(?=\s*销\s*名|$)", full)
    if not m:
        m = re.search(r"买方[信息]?[^名]{0,20}名\s*称[：:\s]*([^\n]{2,40})", full)
    if m:
        data["buyer"] = re.sub(r"[（(].*?[）)]", "", m.group(1)).strip("：: ")

    # 销售方名称（发票「销 售 方 / 卖方」区「名 称：」后完整内容）→ 销售方台账（默认取「ORG_NAME」）
    m = re.search(r"销\s*售\s*方[^名\n]{0,20}名\s*称[：:\s]*([^\n]{2,40})", full)
    if not m:
        m = re.search(r"销[^名\n]{0,6}名\s*称[：:]\s*([^\n]{2,40})", full)
    if not m:
        m = re.search(r"卖方[信息]?[^名]{0,20}名\s*称[：:\s]*([^\n]{2,40})", full)
    if m:
        data["seller"] = re.sub(r"[（(].*?[）)]", "", m.group(1)).strip("：: ")
    if not data["seller"]:
        for line in texts:
            line = line.strip()
            # 排除标签行（含「名称」），避免把「购 名称：X 销 名称：Y」整行当销售方
            if ("律师事务所" in line or "事务所" in line) and len(line) <= 40 and "名称" not in line:
                data["seller"] = line
                break

    # 所有 ¥ 金额（用于智能兜底）
    all_money = []
    for a in re.findall(r"[¥￥]\s*([0-9OolI,，.．]+)", compact):
        v = _money(a)
        if v is not None:
            all_money.append(v)

    # 价税合计（小写）：优先「（小写）」标签；兜底最大 ¥ 金额。
    # 金额口径与《发票AI识别提示词（整合版）》对齐：total_amount=价税合计、amount=不含税金额。
    m = re.search(r"[（(]小写[)）][：:\s]*[¥￥]?\s*([0-9OolI,，.．]+)", compact) or \
        re.search(r"价\s*税\s*合\s*计[^¥￥]{0,40}[¥￥]\s*([0-9OolI,，.．]+)", compact)
    if m:
        data["total_amount"] = _money(m.group(1)) or 0.0
    elif all_money:
        data["total_amount"] = max(all_money)  # 价税合计为最大金额

    # 不含税金额 amount：优先明细行（*编码*名称 金额 税率 税额）与合计行（合 计 ¥金额 ¥税额）的「金额」；
    # 取不到才兜底=价税合计（兼容无明细金额的旧版式）。
    m = re.search(r"^\*[^*\n]{1,40}\*[^\s\n]{0,60}\s+([0-9OolI,，.．]+)\s+[0-9.]+%\s+([0-9OolI,，.．]+)",
                  full, re.M)
    if not m:
        m = re.search(r"合\s*计[¥￥]?\s*([0-9OolI,，.．]+)\s*[¥￥]\s*([0-9OolI,，.．]+)", compact)
    if m:
        data["amount"] = _money(m.group(1)) or 0.0
        # 明细/合计行的税额优先于兜底（min 金额）
        if _money(m.group(2)) is not None:
            data["tax_amount"] = _money(m.group(2))
    if not data["amount"]:
        data["amount"] = data["total_amount"]

    # 价税合计大写：「（大写）」后内容（同行使截到「（小写」前）
    m = re.search(r"[（(]大写[)）][：:\s]*([^\n]{2,30})", full)
    if m:
        cn = m.group(1).strip()
        cn = re.split(r"[（(]小写[)）]", cn)[0].strip()
        data["total_amount_cn"] = cn

    # 税率/征收率：如 6%、13%、1%、免税。
    # 兼容两种写法：①「税率/征收率：6%」标签式；②明细行裸值（*编码*名称 金额 6% 税额，
    # 税率作为独立 token，前后为空白/行界，避免与金额粘连误抓）
    m = re.search(r"(?:税率|征收率)[：:\s]*([0-9.]+%|免税|不征税|零税率)", compact)
    if not m:
        m = re.search(r"(?:^|[ \t])([0-9]{1,3}(?:\.[0-9]{1,2})?%|免税|不征税|零税率)(?=[ \t]|$)",
                      full, re.M)
    if m:
        data["tax_rate"] = m.group(1)

    # 双方税号：统一社会信用代码 / 纳税人识别号（购买方区与销售方区各自取值；支持字间空格写法）
    # 兼容两列并排布局（购方/销方税号同行，如「信 统一社会信用代码/纳税人识别号：A 信 统一社会信用代码/纳税人识别号：B」）
    tax_id_label = r"统\s*一\s*社\s*会\s*信\s*用\s*代\s*码"
    # 标签后允许「/纳税人识别号：」等非数字字符（[^0-9]{0,20}），再取 15-20 位代码
    m = re.search(r"购\s*买\s*方[^0-9]{0,80}%s[^0-9]{0,20}([0-9A-Z]{15,20})" % tax_id_label, compact)
    if not m:
        m = re.search(r"(?<=购买方)[^0-9]{0,80}([0-9A-Z]{18})", compact)
    if not m:
        m = re.search(r"%s[^0-9]{0,20}([0-9A-Z]{15,20})" % tax_id_label, compact)
    if m:
        data["buyer_tax_id"] = m.group(1)
    m = re.search(r"销\s*售\s*方[^0-9]{0,80}%s[^0-9]{0,20}([0-9A-Z]{15,20})" % tax_id_label, compact)
    if not m:
        m = re.search(r"(?<=销售方)[^0-9]{0,80}([0-9A-Z]{18})", compact)
    if not m:
        # 第二个税号标签（两列并排时购方在前、销方在后；用 findall 取第 2 个）
        ids = re.findall(r"%s[^0-9]{0,20}([0-9A-Z]{15,20})" % tax_id_label, compact)
        if len(ids) >= 2:
            m = (ids[1],)
    if m:
        data["seller_tax_id"] = m[0] if isinstance(m, tuple) else m.group(1)
    # 兜底：未区分区域时，把第一个 18 位代码记给购买方（必须含字母，避免把纯数字发票号码前 18 位误当税号）
    if not data["buyer_tax_id"]:
        m = re.search(r"(?<![0-9A-Z])([0-9A-Z]*[A-Z][0-9A-Z]*)(?![0-9A-Z])", compact)
        if m and len(m.group(1)) == 18:
            data["buyer_tax_id"] = m.group(1)

    # 开票人 / 备注
    m = re.search(r"开票人[：:]*([^\n]{1,10})", full)
    if m:
        data["drawer"] = m.group(1).strip()
    m = re.search(r"备注[：:]*([^\n]{1,100})", full)
    if m:
        data["remark"] = m.group(1).strip()

    # 税额：明细行/合计行已提取则保留（见上方 amount 提取）；否则优先「税额」标签；兜底最小 ¥ 金额（专票税额通常最小）
    if not data["tax_amount"]:
        m = re.search(r"税额[：:\s]*[¥￥]?\s*([0-9OolI,，.．]+)", compact)
        if m:
            data["tax_amount"] = _money(m.group(1)) or 0.0
        elif len(all_money) >= 2:
            data["tax_amount"] = min(all_money)

    # 开票日期
    m = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", full)
    if m:
        data["invoice_date"] = "%s-%02d-%02d" % (m.group(1), int(m.group(2)), int(m.group(3)))
    else:
        m = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", full)
        if m:
            data["invoice_date"] = "%s-%02d-%02d" % (m.group(1), int(m.group(2)), int(m.group(3)))

    # 发票类型（专票/普票）：增值税专用发票 / 增值税普通发票
    data["invoice_type"] = detect_invoice_type(full)

    # 项目名称（发票「项目名称」/「货物或应税劳务、服务名称」）→ 原始值 + 费用科目
    item_name = extract_item_name(full)
    data["item_name"] = item_name
    cat_name = item_name_to_category(item_name)
    if cat_name:
        # 优先用项目名称主体作为费用科目（如 *生产生活服务*法律咨询 → 法律咨询）
        data["category"] = cat_name
    else:
        # 无项目名称时回退关键词归集（诉讼费/律师代理费/差旅费等）
        data["category"] = categorize_invoice(full, cats)
    return data


@app.route("/api/ocr/status")
@perm_required("invoice.edit")
def ocr_status():
    """返回本地 OCR 依赖与模型状态（无需联网）"""
    pp = pdfplumber_available()
    return jsonify({
        "ok": True,
        "available": ocr_available(),
        "pdfplumber": pp,
        "detail": ("本地离线 OCR（RapidOCR + PP-OCRv3）+ PDF 文本层引擎（pdfplumber）已就绪，"
                   "识别过程数据不出本机" if (ocr_available() and pp) else
                   "本地 OCR 依赖未安装，可点击「自动配置依赖」完成安装（需服务器联网一次）"),
    })


@app.route("/api/ocr/setup", methods=["POST"])
@perm_required("invoice.edit")
def ocr_setup():
    """自动安装本地 OCR 依赖：rapidocr_onnxruntime（模型随包内置，安装后离线可用）"""
    import subprocess
    d = request.get_json(silent=True) or {}
    mirror = d.get("mirror") or "https://pypi.tuna.tsinghua.edu.cn/simple"
    if ocr_available():
        return jsonify({"ok": True, "msg": "依赖已就绪，无需重复安装", "available": True})
    try:
        proc = subprocess.run(
            [sys.executable, "-m", "pip", "install", "-i", mirror,
             "rapidocr_onnxruntime", "pdfplumber"],
            capture_output=True, text=True, timeout=600)
    except Exception as e:
        audit("create", "system", "ocr", "OCR 依赖安装异常: %s" % e)
        return jsonify({"ok": False, "msg": "安装进程异常: %s" % e}), 500
    ok = ocr_available()
    audit("create", "system", "ocr", "OCR 依赖安装返回码=%d 就绪=%s"
          % (proc.returncode, ok))
    if not ok:
        tail = (proc.stderr or proc.stdout or "")[-500:]
        return jsonify({"ok": False, "msg": "依赖安装失败，请检查网络后重试", "detail": tail}), 500
    return jsonify({"ok": True, "msg": "本地 OCR 依赖已安装完成，可离线使用", "available": True})


@app.route("/api/invoices/ocr-recognize", methods=["POST"])
@perm_required("invoice.edit")
def ocr_recognize():
    """上传发票图片或 PDF → 本地离线 OCR 识别 → 提取字段并归集费用科目。
    PDF 自动转图逐页识别；金额经 OCR 误识别修正。全程本机处理，数据不出本地。"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "msg": "请上传发票图片或 PDF"}), 400
    if not ocr_available():
        return jsonify({"ok": False, "need_setup": True,
                        "msg": "本地 OCR 依赖未安装，请先点击「自动配置依赖」"}), 400
    raw = f.read()
    try:
        texts = ocr_read_file(raw, f.filename)
    except Exception as e:
        return jsonify({"ok": False, "msg": "OCR 识别失败（PDF 需可解析）: %s" % e}), 500
    cats = json.loads(get_setting("fee_categories", "[]") or "[]")
    data = parse_invoice_fields(texts, cats)
    if (not data.get("invoice_no") and not data.get("amount")
            and not data.get("buyer") and not data.get("seller")):
        return jsonify({"ok": False, "msg": "未能从文件中识别出发票要素，请确认清晰且为完整发票",
                        "raw": data["raw"][:500]}), 400

    # ---- 字段校正（购买方→客户名称；销售方默认ORG_NAME，不一致则标记待核对）----
    raw_buyer = (data.get("buyer") or "").strip()
    raw_seller = (data.get("seller") or "").strip()
    DEFAULT_SELLER = ORG_NAME
    client_name = raw_buyer
    seller = DEFAULT_SELLER
    review_flag = 0
    review_note = ""
    if raw_seller and DEFAULT_SELLER not in raw_seller and raw_seller not in DEFAULT_SELLER:
        # 识别销售方与默认不一致：以「ORG_NAME」为准，保留原始值追溯并标记待人工核对
        review_flag = 1
        review_note = ("识别销售方为「%s」，与默认销售方「%s」不一致，"
                       "已按默认记为%s，请人工核对" % (raw_seller, DEFAULT_SELLER, DEFAULT_SELLER))
    warnings = [review_note] if review_flag else []

    # 保存上传文件 + 发票记录（支持登记案件号）
    fname = "%d_%s" % (int(time.time()), re.sub(r"[^\w.\-]", "_", f.filename or "invoice"))
    fpath = os.path.join(UPLOAD_DIR, fname)
    with open(fpath, "wb") as fp:
        fp.write(raw)
    case_no = (request.form.get("case_no") or "").strip()
    db = get_db()
    total_amount = float(data.get("total_amount") or data.get("amount") or 0)
    cur = db.execute(
        "INSERT INTO invoices(case_id,case_no,invoice_no,client_name,seller,seller_raw,buyer_raw,review_flag,review_note,"
        "amount,tax_amount,total_amount,total_amount_cn,tax_rate,buyer_tax_id,seller_tax_id,drawer,remark,"
        "category,invoice_date,invoice_type,item_name,file_path,ai_raw,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (None, case_no, data.get("invoice_no", ""), client_name, seller, raw_seller, raw_buyer,
         review_flag, review_note,
         float(data.get("amount") or 0), float(data.get("tax_amount") or 0),
         total_amount, data.get("total_amount_cn", ""), data.get("tax_rate", ""),
         data.get("buyer_tax_id", ""), data.get("seller_tax_id", ""),
         data.get("drawer", ""), data.get("remark", ""),
         data.get("category", ""), data.get("invoice_date", ""),
         data.get("invoice_type", ""), data.get("item_name", ""),
         fpath, json.dumps(data, ensure_ascii=False)[:2000],
         session["username"], now()))
    db.commit()
    audit("create", "invoice", cur.lastrowid,
          "本地OCR识别发票 %s，类型=%s，科目=%s%s"
          % (data.get("invoice_no", ""), data.get("invoice_type", "") or "未识别",
             data.get("category", ""),
             ("；待核对：" + review_note) if review_flag else ""))
    result = {
        "invoice_no": data.get("invoice_no", ""),
        "client_name": client_name,
        "seller": seller,
        "seller_raw": raw_seller,
        "buyer_raw": raw_buyer,
        "amount": data.get("amount", 0.0),
        "tax_amount": data.get("tax_amount", 0.0),
        "total_amount": total_amount,
        "total_amount_cn": data.get("total_amount_cn", ""),
        "tax_rate": data.get("tax_rate", ""),
        "buyer_tax_id": data.get("buyer_tax_id", ""),
        "seller_tax_id": data.get("seller_tax_id", ""),
        "drawer": data.get("drawer", ""),
        "remark": data.get("remark", ""),
        "invoice_date": data.get("invoice_date", ""),
        "invoice_type": data.get("invoice_type", ""),
        "item_name": data.get("item_name", ""),
        "category": data.get("category", ""),
        "review_flag": review_flag,
        "review_note": review_note,
    }
    return jsonify({"ok": True, "id": cur.lastrowid, "data": result,
                    "warnings": warnings, "engine": "local-ocr"})


# ================================================================ 发票批量导入
_REQUIRED_INV_FIELDS = [
    ("invoice_no", "发票号码"),
    ("buyer", "购买方"),
    ("amount", "价税合计"),
    ("invoice_date", "开票日期"),
]


def _save_invoice_record(db, data, raw, fpath, case_no, username, extra_note=""):
    """发票记录入库（批量/单张共用）：购买方→客户名称；销售方默认ORG_NAME，不一致标记待核对；
    缺关键字段标记待补。返回 (id, review_flag, review_note, missing)。"""
    raw_buyer = (data.get("buyer") or "").strip()
    raw_seller = (data.get("seller") or "").strip()
    DEFAULT_SELLER = ORG_NAME
    client_name = raw_buyer
    seller = DEFAULT_SELLER
    review_flag = 0
    review_note = extra_note or ""
    if raw_seller and DEFAULT_SELLER not in raw_seller and raw_seller not in DEFAULT_SELLER:
        review_flag = 1
        note = ("识别销售方为「%s」，与默认销售方「%s」不一致，已按默认记为%s，请人工核对"
                % (raw_seller, DEFAULT_SELLER, DEFAULT_SELLER))
        review_note = (review_note + "；" + note) if review_note else note
    # 缺字段检测：缺失关键字段仍入库，但标记待补并提示
    missing = []
    for key, label in _REQUIRED_INV_FIELDS:
        v = data.get(key)
        if key in ("amount", "total_amount"):
            if not v or float(v or 0) <= 0:
                missing.append(label)
        elif not str(v or "").strip():
            missing.append(label)
    if missing:
        review_flag = 1
        note = "缺字段：" + "、".join(missing) + "，已登记待人工补录"
        review_note = (review_note + "；" + note) if review_note else note
    total_amount = float(data.get("total_amount") or data.get("amount") or 0)
    cur = db.execute(
        "INSERT INTO invoices(case_no,invoice_no,client_name,seller,seller_raw,buyer_raw,"
        "review_flag,review_note,amount,tax_amount,total_amount,total_amount_cn,tax_rate,"
        "buyer_tax_id,seller_tax_id,drawer,remark,category,invoice_date,invoice_type,"
        "item_name,file_path,ai_raw,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (case_no, data.get("invoice_no", ""), client_name, seller, raw_seller, raw_buyer,
         review_flag, review_note,
         float(data.get("amount") or 0), float(data.get("tax_amount") or 0),
         total_amount, data.get("total_amount_cn", ""), data.get("tax_rate", ""),
         data.get("buyer_tax_id", ""), data.get("seller_tax_id", ""),
         data.get("drawer", ""), data.get("remark", ""),
         data.get("category", ""), data.get("invoice_date", ""),
         data.get("invoice_type", ""), data.get("item_name", ""),
         fpath, json.dumps(data, ensure_ascii=False)[:2000],
         username, now()))
    return cur.lastrowid, review_flag, review_note, missing


@app.route("/api/invoices/batch-import", methods=["POST"])
@perm_required("invoice.edit")
def invoices_batch_import():
    """发票批量导入：一次上传多张发票（图片/PDF 混传），逐张本地识别并登记台账。
    异常分级处理：
      - ok      ：识别完整，直接入库；
      - partial ：已入库但缺关键字段（发票号码/购买方/价税合计/开票日期），标记待补并提示；
      - error   ：识别失败（无任何发票要素 / 文件损坏），不入库，提示原因。
    全程本机处理（pdfplumber 文本层 + RapidOCR 图像），数据不出本地。"""
    files = request.files.getlist("files") or ([request.files.get("file")] if request.files.get("file") else [])
    if not files or all(f.filename in ("", None) for f in files):
        return jsonify({"ok": False, "msg": "请至少选择一张发票文件（可多选）"}), 400
    if not ocr_available():
        return jsonify({"ok": False, "need_setup": True,
                        "msg": "本地 OCR 依赖未安装，请先点击「自动配置依赖」"}), 400
    case_no = (request.form.get("case_no") or "").strip()
    cats = json.loads(get_setting("fee_categories", "[]") or "[]")
    db = get_db()
    results = []
    n_ok = n_partial = n_error = 0
    for f in files:
        fname = f.filename or "未命名文件"
        item = {"filename": fname}
        try:
            raw = f.read()
            texts = ocr_read_file(raw, fname)
            data = parse_invoice_fields(texts, cats)
            if (not data.get("invoice_no") and not data.get("amount")
                    and not data.get("buyer") and not data.get("seller")
                    and not data.get("invoice_type")):
                item.update({"status": "error",
                             "message": "未能识别出任何发票要素，请确认文件清晰且为完整发票"})
                n_error += 1
                results.append(item)
                continue
            # 保存上传文件副本（可追溯）
            spath = "%d_%s" % (int(time.time()), re.sub(r"[^\w.\-]", "_", fname))
            fpath = os.path.join(UPLOAD_DIR, spath)
            with open(fpath, "wb") as fp:
                fp.write(raw)
            iid, _flag, _note, missing = _save_invoice_record(
                db, data, raw, fpath, case_no, session["username"])
            db.commit()
            audit("create", "invoice", iid, "批量导入发票 %s 类型=%s 科目=%s%s"
                  % (data.get("invoice_no", ""), data.get("invoice_type", "") or "未识别",
                     data.get("category", ""),
                     ("；" + _note) if _note else ""))
            item.update({
                "status": "partial" if missing else "ok",
                "id": iid,
                "invoice_no": data.get("invoice_no", ""),
                "invoice_type": data.get("invoice_type", ""),
                "client_name": (data.get("buyer") or "").strip(),
                "category": data.get("category", ""),
                "amount": data.get("amount", 0.0),
                "invoice_date": data.get("invoice_date", ""),
                "missing": missing,
                "message": "缺字段：" + "、".join(missing) + "，已登记待人工补录" if missing else "识别完整",
            })
            if missing:
                n_partial += 1
            else:
                n_ok += 1
        except Exception as e:
            item.update({"status": "error", "message": "识别失败：%s" % e})
            n_error += 1
        results.append(item)
    return jsonify({
        "ok": True, "total": len(results), "success": n_ok,
        "partial": n_partial, "error": n_error, "results": results,
        "engine": "local-ocr+pdfplumber",
    })


# ================================================================ AI 发票识别
@app.route("/api/invoices", methods=["GET", "POST"])
@login_required
def invoices():
    db = get_db()
    if request.method == "GET":
        if not has_perm("case.view"):
            return jsonify({"ok": False, "msg": "无权限"}), 403
        kw = request.args.get("kw", "").strip()
        case_no = request.args.get("case_no", "").strip()
        buyer = request.args.get("buyer", "").strip()  # 按购买方（客户名称）筛选
        sql = "SELECT * FROM invoices"
        clauses = []
        args = []
        if kw:
            clauses.append("(invoice_no LIKE ? OR client_name LIKE ? OR case_no LIKE ? OR category LIKE ?)")
            args += ["%" + kw + "%"] * 4
        if case_no:
            clauses.append("case_no LIKE ?")
            args.append("%" + case_no + "%")
        if buyer:
            clauses.append("client_name LIKE ?")
            args.append("%" + buyer + "%")
        if clauses:
            sql += " WHERE " + " AND ".join(clauses)
        sql += " ORDER BY id DESC LIMIT 500"
        rows = db.execute(sql, args).fetchall()
        return jsonify({"ok": True, "data": [dict(r) for r in rows]})
    if not has_perm("invoice.edit"):
        return jsonify({"ok": False, "msg": "无权限"}), 403
    d = request.get_json(force=True)
    # 手动登记：界面「购买方」→ 客户名称；销售方默认记「ORG_NAME」（无识别值，无需待核对）
    client_name = (d.get("client_name") or d.get("seller") or "").strip()
    DEFAULT_SELLER = ORG_NAME
    cur = db.execute(
        "INSERT INTO invoices(case_id,case_no,invoice_no,client_name,seller,seller_raw,buyer_raw,review_flag,review_note,"
        "amount,tax_amount,total_amount,total_amount_cn,tax_rate,buyer_tax_id,seller_tax_id,drawer,remark,"
        "category,invoice_date,invoice_type,item_name,confirm,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (d.get("case_id"), d.get("case_no", ""), d.get("invoice_no", ""), client_name,
         DEFAULT_SELLER, "", "",
         0, "",
         float(d.get("amount") or 0), float(d.get("tax_amount") or 0),
         float(d.get("total_amount") or 0), d.get("total_amount_cn", ""),
         d.get("tax_rate", ""), d.get("buyer_tax_id", ""), d.get("seller_tax_id", ""),
         d.get("drawer", ""), d.get("remark", ""),
         d.get("category", ""), d.get("invoice_date", ""),
         d.get("invoice_type", ""), d.get("item_name", ""),
         1 if d.get("confirm") else 0, session["username"], now()))
    db.commit()
    audit("create", "invoice", cur.lastrowid, "登记发票 %s 金额 %.2f"
          % (d.get("invoice_no", ""), float(d.get("amount") or 0)))
    return jsonify({"ok": True, "id": cur.lastrowid})


@app.route("/api/invoices/compare")
@perm_required("case.view")
def invoice_compare():
    """开票金额与案件开票金额比对：按案件号汇总发票「价税合计」金额 vs 案件 invoice_amount。
    金额口径与《发票AI识别提示词（整合版）》一致：价税合计 = total_amount（旧数据已迁移）。"""
    db = get_db()
    rows = db.execute(
        "SELECT i.case_no, COUNT(*) cnt, COALESCE(SUM(COALESCE(i.total_amount, i.amount)),0) inv_sum"
        " FROM invoices i WHERE i.case_no != '' GROUP BY i.case_no").fetchall()
    result = []
    for r in rows:
        c = db.execute("SELECT * FROM cases WHERE case_no=? AND deleted=0",
                       (r["case_no"],)).fetchone()
        case_inv = float(c["invoice_amount"] or 0) if c else 0.0
        inv_sum = float(r["inv_sum"] or 0)
        diff = round(case_inv - inv_sum, 2)
        result.append({
            "case_no": r["case_no"], "invoice_count": r["cnt"],
            "invoice_sum": round(inv_sum, 2),
            "case_invoice_amount": case_inv,
            "diff": diff,
            "matched": abs(diff) < 0.01,
        })
    return jsonify({"ok": True, "data": result})


@app.route("/api/invoices/<int:iid>", methods=["PUT", "DELETE"])
@perm_required("invoice.edit")
def invoice_op(iid):
    db = get_db()
    r = db.execute("SELECT * FROM invoices WHERE id=?", (iid,)).fetchone()
    if not r:
        return jsonify({"ok": False, "msg": "不存在"}), 404
    if request.method == "DELETE":
        db.execute("DELETE FROM invoices WHERE id=?", (iid,))
        db.commit()
        audit("delete", "invoice", iid, "删除发票 %s" % r["invoice_no"])
        return jsonify({"ok": True})
    d = request.get_json(force=True)
    # 全字段可编辑（含案件号 case_no、客户名称与销售方台账、价税合计/税率/税号/开票人/备注）
    fields = {k: d[k] for k in ("invoice_no", "client_name", "seller", "seller_raw",
                                "buyer_raw", "review_flag", "review_note", "category",
                                "invoice_date", "invoice_type", "item_name",
                                "case_id", "case_no", "total_amount", "total_amount_cn",
                                "tax_rate", "buyer_tax_id", "seller_tax_id",
                                "drawer", "remark") if k in d}
    for k in ("amount", "tax_amount"):
        if k in d:
            fields[k] = float(d[k] or 0)
    if "confirm" in d:
        fields["confirm"] = 1 if d["confirm"] else 0
    if fields:
        sets = ",".join(k + "=?" for k in fields)
        db.execute("UPDATE invoices SET %s WHERE id=?" % sets,
                   list(fields.values()) + [iid])
        db.commit()
        audit("update", "invoice", iid, "修改字段: %s" % ",".join(fields.keys()))
    return jsonify({"ok": True})


# ---------------------------------------------------------------- AI 识别后置校验（与《发票AI识别提示词（整合版）》失败兜底规则对应）
_ORG_MARKERS = ("公司", "事务所", "有限", "集团", "银行", "医院", "中心", "厂", "社",
                "院", "所", "部", "局", "学校", "协会", "合作社")


def _post_check_invoice(data):
    """系统端二次校验（模型端提示词已要求，此处双保险），返回 (review_flag, review_note)。

    校验项（对应提示词规范第四/六节）：
      1) 购买方/销售方名称全称校验：2~6 字且无行政区划/组织形式特征 → 疑似简称，提示人工核对；
      2) 项目名称税收分类编码前缀校验：缺失 *编码* 前缀 → 提示人工核对；
      3) 数据勾稽校验：amount + tax_amount ≈ total_amount；amount × 税率 ≈ tax_amount；
         不通过 → 追加「【勾稽异常：请人工复核】」。
    """
    notes = []
    # 1) 名称全称校验
    for label, key in (("购买方", "buyer_name"), ("销售方", "seller_name")):
        name = (data.get(key) or "").strip()
        if name and len(name) <= 6 and not any(m in name for m in _ORG_MARKERS):
            notes.append("%s名称疑似简称「%s」（应为企业全称），请人工核对" % (label, name))
    # 2) 项目名称编码前缀校验
    item = (data.get("item_name") or "").strip()
    if item and not re.match(r"^\*[^*]+\*", item):
        notes.append("项目名称「%s」缺少税收分类编码前缀（*编码*），请人工核对" % item)
    # 3) 勾稽校验
    try:
        amt = float(data.get("amount") or 0)
        tax = float(data.get("tax_amount") or 0)
        total = float(data.get("total_amount") or 0)
        if total and abs(amt + tax - total) > 0.01:
            notes.append("【勾稽异常：金额+税额≠价税合计，请人工复核】")
        rate_s = str(data.get("tax_rate") or "").strip()
        if rate_s.endswith("%") and amt and tax:
            r = float(rate_s[:-1]) / 100.0
            if round(amt * r, 2) != round(tax, 2) and abs(amt * r - tax) > 0.05:
                notes.append("【勾稽异常：金额×税率≠税额，请人工复核】")
    except Exception:
        pass
    return (1, "；".join(notes)) if notes else (0, "")


@app.route("/api/invoices/ai-recognize", methods=["POST"])
@perm_required("invoice.edit")
def ai_recognize():
    """上传电子发票图片 → 按所选云服务商（火山引擎/阿里百炼/腾讯云/兼容OpenAI）多模态识别 → 归集科目。
    识别提示词与字段映射严格遵循《发票AI识别提示词（整合版）》规范：
      购买方/销售方输出完整企业全称（buyer_name/seller_name）、项目名称保留税收分类编码前缀、
      金额/税额/价税合计/税率/税号/开票人/备注等 20 字段完整映射，并执行系统端后置校验。
    密钥仅来自环境变量或本地配置文件（ai_providers.json），绝不硬编码、不入库明文。"""
    f = request.files.get("file")
    if not f:
        return jsonify({"ok": False, "msg": "请上传发票文件"}), 400
    provider = (request.form.get("provider") or request.args.get("provider")
                or ai_providers.get_active_provider())
    cats = json.loads(get_setting("fee_categories", "[]"))
    try:
        import base64
        raw = f.read()
        fname = "%d_%s" % (int(time.time()), re.sub(r"[^\w.\-]", "_", f.filename))
        fpath = os.path.join(UPLOAD_DIR, fname)
        with open(fpath, "wb") as fp:
            fp.write(raw)
        mime = "image/png" if (f.filename or "").lower().endswith(".png") else "image/jpeg"
        data = ai_providers.recognize(provider_id=provider, image_bytes=raw,
                                       cats=cats, mime=mime)
    except RuntimeError as e:
        return jsonify({"ok": False, "need_config": True, "msg": str(e)}), 400
    except Exception as e:
        return jsonify({"ok": False, "msg": "AI 接口调用失败: %s" % e}), 502

    db = get_db()
    # 规范字段名映射：buyer_name→客户名称（购买方）、seller_name→销售方原始值；兼容旧字段名 buyer/seller
    raw_buyer = (data.get("buyer_name") or data.get("buyer") or "").strip()
    raw_seller = (data.get("seller_name") or data.get("seller") or "").strip()
    DEFAULT_SELLER = ORG_NAME
    client_name = raw_buyer
    seller = DEFAULT_SELLER
    review_flag = 0
    review_note = ""
    # 系统端后置校验（全称/编码前缀/勾稽），与模型端提示词约束双保险
    flag2, note2 = _post_check_invoice(data)
    if flag2:
        review_flag = 1
        review_note = note2
    if raw_seller and DEFAULT_SELLER not in raw_seller and raw_seller not in DEFAULT_SELLER:
        review_flag = 1
        review_note = ("识别销售方为「%s」，与默认销售方「%s」不一致，"
                       "已按默认记为%s，请人工核对" % (raw_seller, DEFAULT_SELLER, DEFAULT_SELLER)
                       + ("；" + note2 if note2 else ""))
    # remark 追加勾稽标记（若模型端未加，系统端按后置校验结果补充）
    remark = (data.get("remark") or "").strip()
    if "勾稽异常" in note2 and "勾稽异常" not in remark:
        remark = (remark + "【勾稽异常：请人工复核】").strip()
    cur = db.execute(
        "INSERT INTO invoices(invoice_no,client_name,seller,seller_raw,buyer_raw,review_flag,review_note,"
        "amount,tax_amount,total_amount,total_amount_cn,tax_rate,buyer_tax_id,seller_tax_id,drawer,remark,"
        "category,invoice_date,invoice_type,item_name,file_path,ai_raw,created_by,created_at)"
        " VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
        (data.get("invoice_no", ""), client_name, seller, raw_seller, raw_buyer,
         review_flag, review_note,
         float(data.get("amount") or 0), float(data.get("tax_amount") or 0),
         float(data.get("total_amount") or 0), data.get("total_amount_cn", ""),
         data.get("tax_rate", ""), data.get("buyer_tax_id", ""), data.get("seller_tax_id", ""),
         data.get("drawer", ""), remark,
         data.get("category", ""), data.get("invoice_date", ""),
         data.get("invoice_type", ""), data.get("item_name", ""),
         fpath, json.dumps(data, ensure_ascii=False)[:2000],
         session["username"], now()))
    db.commit()
    audit("create", "invoice", cur.lastrowid,
          "AI 识别发票 %s（%s），类型=%s，科目=%s%s"
          % (data.get("invoice_no", ""), provider,
             data.get("invoice_type", "") or "未识别", data.get("category", ""),
             ("；待核对：" + review_note) if review_flag else ""))
    result = {
        "invoice_no": data.get("invoice_no", ""),
        "client_name": client_name,
        "seller": seller,
        "seller_raw": raw_seller,
        "buyer_raw": raw_buyer,
        "buyer_tax_id": data.get("buyer_tax_id", ""),
        "seller_tax_id": data.get("seller_tax_id", ""),
        "amount": data.get("amount", 0.0),
        "tax_amount": data.get("tax_amount", 0.0),
        "total_amount": data.get("total_amount", 0.0),
        "total_amount_cn": data.get("total_amount_cn", ""),
        "tax_rate": data.get("tax_rate", ""),
        "drawer": data.get("drawer", ""),
        "remark": remark,
        "invoice_date": data.get("invoice_date", ""),
        "invoice_type": data.get("invoice_type", ""),
        "item_name": data.get("item_name", ""),
        "category": data.get("category", ""),
        "review_flag": review_flag,
        "review_note": review_note,
    }
    return jsonify({"ok": True, "id": cur.lastrowid, "data": result,
                    "warnings": [review_note] if review_flag else []})


# ================================================================ 多云 AI 服务商配置
@app.route("/api/ai/providers", methods=["GET", "PUT"])
@perm_required("setting.manage")
def ai_providers_api():
    """GET：返回各服务商配置状态（脱敏，绝不暴露明文密钥）；
    PUT：将密钥等敏感配置写入本地配置文件（ai_providers.json），不入库明文。"""
    if request.method == "GET":
        return jsonify({"ok": True, "data": ai_providers.public_provider_state()})
    d = request.get_json(force=True)
    active = d.get("active") or ai_providers.get_active_provider()
    providers = d.get("providers", {})
    try:
        state = ai_providers.save_provider_settings(active, providers)
    except Exception as e:
        return jsonify({"ok": False, "msg": "保存失败: %s" % e}), 400
    audit("update", "system", "ai_providers",
          "配置 AI 服务商（激活=%s）" % active)
    return jsonify({"ok": True, "data": state})


@app.route("/api/ai/protocols", methods=["GET"])
@login_required
def ai_protocols_api():
    """返回各云服务商接口协议说明（请求参数 / 鉴权方式 / 返回结果处理），供前端展示。"""
    return jsonify({"ok": True, "data": ai_providers.list_providers()})


# ================================================================ 审计日志
@app.route("/api/logs")
@perm_required("log.view")
def logs():
    kw = request.args.get("kw", "").strip()
    sql = "SELECT * FROM audit_logs"
    args = []
    if kw:
        sql += " WHERE username LIKE ? OR action LIKE ? OR object_id LIKE ? OR detail LIKE ?"
        args = ["%" + kw + "%"] * 4
    sql += " ORDER BY id DESC LIMIT 500"
    rows = get_db().execute(sql, args).fetchall()
    return jsonify({"ok": True, "data": [dict(r) for r in rows]})


# ================================================================ 系统设置
@app.route("/api/case-types")
@login_required
def case_types_api():
    """案件类型下拉选项（民事案由），存储于 settings.case_types；
    所有登录用户可读（案件弹窗需要），仅系统设置管理员可改。"""
    raw = get_setting("case_types", "")
    try:
        arr = json.loads(raw) if raw else []
    except Exception:
        arr = []
    if not isinstance(arr, list):
        arr = []
    return jsonify({"ok": True, "data": arr})


@app.route("/api/case-type-sections", methods=["GET", "PUT"])
@login_required
def case_type_sections_api():
    """案件类型「板块」二级选项配置。
    GET：返回 {type: [板块,...]}（所有登录用户可读，用于案件弹窗下拉）。
    PUT：
      - 传 {type, section}：在指定类型下追加一个板块（支持用户自定义录入）。
      - 传 {sections: {type:[...]}}：整体覆盖（系统设置管理员维护固定选项）。"""
    db = get_db()
    if request.method == "GET":
        raw = get_setting("case_type_sections", "")
        try:
            data = json.loads(raw) if raw else {}
        except Exception:
            data = {}
        if not isinstance(data, dict):
            data = {}
        return jsonify({"ok": True, "data": data})
    # ---- PUT ----
    if not has_perm("setting.manage"):
        return jsonify({"ok": False, "msg": "无系统设置权限"}), 403
    d = request.get_json(force=True)
    if "type" in d and "section" in d:
        data = {}
        raw = get_setting("case_type_sections", "")
        try:
            data = json.loads(raw) if raw else {}
        except Exception:
            data = {}
        if not isinstance(data, dict):
            data = {}
        t = (d.get("type") or "").strip()
        s = (d.get("section") or "").strip()
        if t and s:
            lst = list(data.get(t, []))
            if s not in lst:
                lst.append(s)
            data[t] = lst
        db.execute("INSERT INTO settings(key,value) VALUES(?,?) "
                   "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                   ("case_type_sections", json.dumps(data, ensure_ascii=False)))
        db.commit()
        audit("update", "system", "case_type_sections", "追加板块 %s => %s" % (t, s))
        return jsonify({"ok": True, "data": data})
    if "sections" in d:
        src = d["sections"]
        if not isinstance(src, dict):
            return jsonify({"ok": False, "msg": "sections 必须为对象"}), 400
        norm = {}
        for t, lst in src.items():
            if not t:
                continue
            norm[t] = [str(x) for x in (lst or [])]
        db.execute("INSERT INTO settings(key,value) VALUES(?,?) "
                   "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                   ("case_type_sections", json.dumps(norm, ensure_ascii=False)))
        db.commit()
        audit("update", "system", "case_type_sections", "整体更新板块配置")
        return jsonify({"ok": True, "data": norm})
    return jsonify({"ok": False, "msg": "缺少参数"}), 400


# ================================================================ 数据自动备份
_BACKUP_LOCK = threading.Lock()
_BACKUP_DEFAULTS = {
    "backup_enabled": "0",            # 总开关 1/0
    "backup_path": "",                # 备份目录（空 = DATA_DIR/backups）
    "backup_schedule": "daily",       # daily / weekly / interval / manual
    "backup_time": "23:00",           # daily/weekly 触发时刻 HH:MM
    "backup_weekday": "5",            # weekly 星期几 0-6（5=周六）
    "backup_interval_hours": "24",    # interval 间隔小时数
    "backup_keep": "10",              # 保留份数（0=不清理）
    "backup_scope": "{\"db\":1,\"uploads\":1,\"templates\":1,\"config\":1}",  # 备份范围
}


def default_backup_dir():
    return os.path.join(DATA_DIR, "backups")


def _bk_conn():
    """备份模块专用数据库连接：独立于 Flask 请求上下文，供后台调度线程使用"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def get_backup_cfg():
    """读取备份配置（合并默认值）。返回 dict，含 _scope 解析后的范围。"""
    cfg = dict(_BACKUP_DEFAULTS)
    db = _bk_conn()
    try:
        rows = db.execute(
            "SELECT key,value FROM settings WHERE key LIKE 'backup_%'").fetchall()
    finally:
        db.close()
    for r in rows:
        cfg[r["key"]] = r["value"]
    if not cfg.get("backup_path"):
        cfg["backup_path"] = default_backup_dir()
    try:
        scope = json.loads(cfg.get("backup_scope") or "{}")
    except Exception:
        scope = {}
    if not isinstance(scope, dict):
        scope = {}
    cfg["_scope"] = scope
    return cfg


def save_backup_cfg(d):
    """保存备份配置（白名单 key）。自动尝试创建备份目录，失败不阻断保存。"""
    allow = set(_BACKUP_DEFAULTS)
    db = _bk_conn()
    try:
        for k in allow:
            if k in d:
                v = str(d[k]).strip()
                db.execute("INSERT INTO settings(key,value) VALUES(?,?) "
                           "ON CONFLICT(key) DO UPDATE SET value=excluded.value", (k, v))
        db.commit()
    finally:
        db.close()
    p = str(d.get("backup_path") or "").strip() or default_backup_dir()
    try:
        os.makedirs(p, exist_ok=True)
    except Exception:
        pass
    audit("update", "system", "backup",
          "修改备份配置: %s" % ",".join(str(k) for k in d if k in allow))
    # 保存后重置调度窗口，让新配置立即生效
    _set_backup_state(last_check=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))


def _set_backup_state(**kw):
    """写入备份运行状态（settings 表 backup_* 键，供前端轮询；线程安全独立连接）"""
    db = _bk_conn()
    try:
        for k, v in kw.items():
            db.execute("INSERT INTO settings(key,value) VALUES(?,?) "
                       "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                       ("backup_" + k, str(v)))
        db.commit()
    finally:
        db.close()


def append_backup_log(line):
    """追加一行备份日志到 备份目录/backup.log；目录不可用时回退默认目录，日志写失败不影响备份结果"""
    for p in (get_backup_cfg()["backup_path"], default_backup_dir()):
        try:
            os.makedirs(p, exist_ok=True)
            with open(os.path.join(p, "backup.log"), "a", encoding="utf-8") as f:
                f.write("[%s] %s\n" % (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), line))
            return
        except Exception:
            continue


def _backup_db_safe(target):
    """SQLite 在线备份 API 复制数据库，全程一致、无需停机"""
    src = sqlite3.connect(DB_PATH)
    dst = sqlite3.connect(target)
    try:
        src.backup(dst)
    finally:
        dst.close()
        src.close()


def _dir_size(p):
    total = 0
    for root, _, files in os.walk(p):
        for fn in files:
            try:
                total += os.path.getsize(os.path.join(root, fn))
            except Exception:
                pass
    return total


def _readable_error(e, stage):
    """把底层异常翻译成可读错误提示（供界面直接展示）"""
    s = str(e)
    low = s.lower()
    if any(k in low for k in ("permission", "denied", "errno 13", "winerror 5", "拒绝访问")):
        return "备份目录无写入权限，请检查目录权限设置后重试（阶段：%s）" % stage
    if any(k in low for k in ("no space", "errno 28", "disk", "winerror 112")):
        return "磁盘空间不足，请清理磁盘空间后重试（阶段：%s）" % stage
    if any(k in low for k in ("no such", "not a directory", "errno 2", "winerror 3", "系统找不到指定的路径", "找不到")):
        return "备份路径无效或不存在，请检查备份路径配置（阶段：%s）" % stage
    if any(k in low for k in ("locked", "busy", "winerror 32", "winerror 33")):
        return "文件被占用（可能正在写入），请稍后重试（阶段：%s）" % stage
    return "备份失败：%s（阶段：%s）" % (s[:300], stage)


def cleanup_old_backups(cfg=None, log=False):
    """按保留份数清理过期备份，返回删除数量；keep<=0 表示不清理。"""
    cfg = cfg or get_backup_cfg()
    try:
        keep = int(cfg.get("backup_keep") or 0)
    except (TypeError, ValueError):
        keep = 0
    if keep <= 0:
        return 0
    out_dir = cfg["backup_path"]
    if not os.path.isdir(out_dir):
        return 0
    files = sorted(f for f in os.listdir(out_dir)
                   if f.startswith("coop_backup_") and f.endswith(".zip"))
    removed = 0
    while len(files) > keep:
        f = files.pop(0)
        try:
            os.remove(os.path.join(out_dir, f))
            removed += 1
            if log:
                append_backup_log("清理过期备份 %s（保留 %d 份）" % (f, keep))
        except Exception as e:
            if log:
                append_backup_log("清理失败 %s：%s" % (f, e))
    return removed


def list_backups(cfg=None):
    """列出备份文件（名称/大小/修改时间，新→旧排序）"""
    cfg = cfg or get_backup_cfg()
    out_dir = cfg["backup_path"]
    if not os.path.isdir(out_dir):
        return []
    items = []
    for fn in os.listdir(out_dir):
        if fn.startswith("coop_backup_") and fn.endswith(".zip"):
            fp = os.path.join(out_dir, fn)
            try:
                st = os.stat(fp)
                items.append({"name": fn, "size": st.st_size,
                              "mtime": datetime.fromtimestamp(st.st_mtime)
                              .strftime("%Y-%m-%d %H:%M:%S")})
            except Exception:
                continue
    items.sort(key=lambda x: x["name"], reverse=True)
    return items


def run_backup_once(cfg=None, trigger="manual"):
    """执行一次完整备份（线程安全，并发时第二个调用立即返回 running）。
    返回 (ok, result_dict)；成功/失败均更新状态与日志，并自动清理过期备份。"""
    if not _BACKUP_LOCK.acquire(blocking=False):
        return False, {"running": True, "msg": "已有备份任务在运行，请稍候"}
    tmp_zip = None
    try:
        cfg = cfg or get_backup_cfg()
        if trigger == "scheduled" and str(cfg.get("backup_enabled")) != "1":
            return True, {"skipped": True, "msg": "自动备份已停用，跳过本次调度"}
        scope = cfg.get("_scope") or {}
        out_dir = cfg["backup_path"]
        os.makedirs(out_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_name = "coop_backup_%s.zip" % stamp
        zip_path = os.path.join(out_dir, zip_name)
        tmp_zip = zip_path + ".tmp"
        t0 = time.time()

        _set_backup_state(running=1, progress="正在准备备份内容清单...")
        manifest = {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "trigger": trigger,
            "db": os.path.basename(DB_PATH),
            "scope": {k: 1 if scope.get(k, 1) else 0 for k in ("db", "uploads", "templates", "config")},
            "items": [],
        }
        with zipfile.ZipFile(tmp_zip, "w", zipfile.ZIP_DEFLATED) as zf:
            # 1) 数据库：在线备份到临时文件再入包，保证一致性
            if scope.get("db", 1):
                _set_backup_state(progress="正在备份数据库（在线快照）...")
                db_tmp = os.path.join(out_dir, ".backup_tmp_%s.db" % stamp)
                try:
                    _backup_db_safe(db_tmp)
                    zf.write(db_tmp, os.path.basename(DB_PATH))
                    manifest["items"].append({"name": os.path.basename(DB_PATH),
                                              "size": os.path.getsize(db_tmp)})
                finally:
                    try:
                        if os.path.exists(db_tmp):
                            os.remove(db_tmp)
                    except Exception:
                        pass
            # 2) 上传文件（发票图片/PDF 等）
            if scope.get("uploads", 1) and os.path.isdir(UPLOAD_DIR):
                _set_backup_state(progress="正在收集上传文件...")
                manifest["items"].append({"name": "uploads", "size": _dir_size(UPLOAD_DIR)})
                for root, _, files in os.walk(UPLOAD_DIR):
                    for fn in files:
                        fp = os.path.join(root, fn)
                        rel = os.path.join("uploads", os.path.relpath(fp, UPLOAD_DIR))
                        zf.write(fp, rel)
            # 3) 结算单 Word 模板
            if scope.get("templates", 1) and os.path.isdir(TEMPLATE_DIR):
                _set_backup_state(progress="正在收集结算模板...")
                manifest["items"].append({"name": "templates", "size": _dir_size(TEMPLATE_DIR)})
                for fn in os.listdir(TEMPLATE_DIR):
                    fp = os.path.join(TEMPLATE_DIR, fn)
                    if os.path.isfile(fp):
                        zf.write(fp, os.path.join("templates", fn))
            # 4) 配置文件（AI 服务商配置等）
            if scope.get("config", 1):
                p = os.path.join(DATA_DIR, "ai_providers.json")
                if os.path.exists(p):
                    manifest["items"].append({"name": "config/ai_providers.json",
                                              "size": os.path.getsize(p)})
                    zf.write(p, os.path.join("config", "ai_providers.json"))
            _set_backup_state(progress="正在打包压缩...")
            zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=1))

        os.replace(tmp_zip, zip_path)
        tmp_zip = None
        size = os.path.getsize(zip_path)
        cost = round(time.time() - t0, 1)
        result = {"ok": True, "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                  "file": zip_name, "size": size, "cost": cost}
        _set_backup_state(running=0, progress="",
                          last_result=json.dumps(result, ensure_ascii=False),
                          last_run=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        append_backup_log("成功  %s  %d 字节  耗时 %.1fs  (触发:%s)" % (zip_name, size, cost, trigger))
        cleaned = cleanup_old_backups(cfg, log=True)
        result["cleaned"] = cleaned
        return True, result
    except Exception as e:
        err = _readable_error(e, "打包压缩")
        result = {"ok": False, "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                  "error": err}
        _set_backup_state(running=0, progress="",
                          last_result=json.dumps(result, ensure_ascii=False),
                          last_run=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        append_backup_log("失败  %s" % err)
        try:
            if tmp_zip and os.path.exists(tmp_zip):
                os.remove(tmp_zip)
        except Exception:
            pass
        return False, result
    finally:
        _BACKUP_LOCK.release()


def _next_run_seconds(cfg):
    """计算距下次自动运行还有多少秒（daily/weekly 按时刻，interval 按上次运行）"""
    sched = cfg.get("backup_schedule", "daily")
    now = datetime.now()
    if sched == "interval":
        try:
            hours = float(cfg.get("backup_interval_hours") or 24)
        except (TypeError, ValueError):
            hours = 24
        if hours <= 0:
            hours = 24
        last = cfg.get("backup_last_run") or ""
        if not last:
            return 0  # 尚未运行过 → 立即执行
        try:
            last_dt = datetime.strptime(last, "%Y-%m-%d %H:%M:%S")
            return max(0.0, hours * 3600 - (now - last_dt).total_seconds())
        except ValueError:
            return 0
    try:
        hh, mm = (int(x) for x in (cfg.get("backup_time") or "23:00").split(":"))
    except Exception:
        hh, mm = 23, 0
    target = now.replace(hour=hh, minute=mm, second=0, microsecond=0)
    if sched == "weekly":
        try:
            wd = int(cfg.get("backup_weekday") or 5) % 7
        except (TypeError, ValueError):
            wd = 5
        target += timedelta(days=(wd - now.weekday()) % 7)
    delta = (target - now).total_seconds()
    return delta if delta > 0 else delta + 86400


def backup_scheduler_loop():
    """后台调度线程：每 60s 检查一次，到达触发点则执行备份（守护线程，不阻塞主服务）"""
    while True:
        try:
            cfg = get_backup_cfg()
            if str(cfg.get("backup_enabled")) == "1" and cfg.get("backup_schedule") != "manual":
                wait = _next_run_seconds(cfg)
                if wait <= 0:
                    run_backup_once(cfg, trigger="scheduled")
                    _set_backup_state(last_run=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
                    wait = _next_run_seconds(get_backup_cfg())
                time.sleep(min(max(wait, 5), 3600))
            else:
                time.sleep(60)
        except Exception:
            time.sleep(60)


def _run_backup_async():
    """异步执行立即备份（接口用），完成后更新 last_run 供 interval 模式计时"""
    run_backup_once(trigger="manual")
    _set_backup_state(last_run=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))


@app.route("/api/backup/config", methods=["GET", "PUT"])
@perm_required("setting.manage")
def backup_config_api():
    """读取/保存备份配置（路径、周期、范围、保留份数等）"""
    if request.method == "GET":
        cfg = get_backup_cfg()
        cfg.pop("_scope", None)
        return jsonify({"ok": True, "data": cfg, "default_path": default_backup_dir()})
    d = request.get_json(force=True)
    save_backup_cfg(d)
    return jsonify({"ok": True, "msg": "备份配置已保存"})


@app.route("/api/backup/status")
@perm_required("setting.manage")
def backup_status_api():
    """当前运行状态 + 上次结果（前端轮询进度用）"""
    rows = get_db().execute(
        "SELECT key,value FROM settings WHERE key LIKE 'backup_%'").fetchall()
    st = {r["key"]: r["value"] for r in rows}
    try:
        st["_last_result"] = json.loads(st.get("backup_last_result") or "{}")
    except Exception:
        st["_last_result"] = {}
    st["backup_keep"] = st.get("backup_keep") or "0"
    st["_running"] = 1 if str(st.get("backup_running")) == "1" else 0
    return jsonify({"ok": True, "data": st})


@app.route("/api/backup/run", methods=["POST"])
@perm_required("setting.manage")
def backup_run_api():
    """立即备份（异步执行，通过 /api/backup/status 轮询进度）"""
    if not _BACKUP_LOCK.acquire(blocking=False):
        return jsonify({"ok": False, "running": True, "msg": "已有备份任务正在运行，请稍候"})
    _BACKUP_LOCK.release()
    threading.Thread(target=_run_backup_async, daemon=True).start()
    return jsonify({"ok": True, "running": True, "msg": "备份已开始，正在执行..."})


@app.route("/api/backup/list")
@perm_required("setting.manage")
def backup_list_api():
    return jsonify({"ok": True, "data": list_backups()})


@app.route("/api/backup/download/<name>")
@perm_required("setting.manage")
def backup_download_api(name):
    if not re.fullmatch(r"coop_backup_\d{8}_\d{6}\.zip", name):
        return jsonify({"ok": False, "msg": "非法的备份文件名"}), 400
    fp = os.path.join(get_backup_cfg()["backup_path"], name)
    if not os.path.isfile(fp):
        return jsonify({"ok": False, "msg": "备份文件不存在，可能已被清理"}), 404
    audit("download", "backup", name, "下载备份文件")
    return send_file(fp, as_attachment=True, download_name=name)


@app.route("/api/backup/delete/<name>", methods=["DELETE"])
@perm_required("setting.manage")
def backup_delete_api(name):
    if not re.fullmatch(r"coop_backup_\d{8}_\d{6}\.zip", name):
        return jsonify({"ok": False, "msg": "非法的备份文件名"}), 400
    fp = os.path.join(get_backup_cfg()["backup_path"], name)
    if not os.path.isfile(fp):
        return jsonify({"ok": False, "msg": "备份文件不存在"}), 404
    try:
        os.remove(fp)
        audit("delete", "backup", name, "删除备份文件")
        return jsonify({"ok": True, "msg": "备份已删除"})
    except Exception as e:
        return jsonify({"ok": False, "msg": "删除失败：%s" % e}), 500


@app.route("/api/backup/cleanup", methods=["POST"])
@perm_required("setting.manage")
def backup_cleanup_api():
    """手动触发过期备份清理（按保留份数）"""
    n = cleanup_old_backups(log=True)
    audit("update", "backup", "cleanup", "手动清理过期备份 %d 个" % n)
    return jsonify({"ok": True, "removed": n, "msg": "已清理 %d 个过期备份" % n})


@app.route("/api/backup/log")
@perm_required("setting.manage")
def backup_log_api():
    # 优先读配置目录日志；配置目录不可用时回退默认备份目录（日志始终可追溯）
    for p in (get_backup_cfg()["backup_path"], default_backup_dir()):
        lp = os.path.join(p, "backup.log")
        if os.path.isfile(lp):
            with open(lp, encoding="utf-8") as f:
                return jsonify({"ok": True, "data": f.read().splitlines()[-200:]})
    return jsonify({"ok": True, "data": []})


@app.route("/api/settings", methods=["GET", "PUT"])
@perm_required("setting.manage")
def settings_api():
    db = get_db()
    if request.method == "GET":
        rows = db.execute("SELECT key,value FROM settings").fetchall()
        data = {r["key"]: r["value"] for r in rows}
        # 注意：ai_api_key 等密钥不再入库，统一从环境变量/配置文件读取，此处不再出现
        return jsonify({"ok": True, "data": data})
    d = request.get_json(force=True)
    for k in ("pair_code", "firm_name", "ai_provider", "ai_base_url", "ai_model",
              "fee_categories", "settle_template", "case_no_seq_start", "case_types"):
        if k in d:
            db.execute("INSERT INTO settings(key,value) VALUES(?,?) "
                       "ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                       (k, str(d[k])))
    db.commit()
    audit("update", "system", "settings", "修改系统设置: %s" % ",".join(d.keys()))
    return jsonify({"ok": True})


@app.route("/api/data-library")
@perm_required("case.view")
def data_library():
    """数据信息库综合查阅：按类别（案件/花名册/客户/发票/财务结算）统一查询数据库"""
    db = get_db()
    dtype = request.args.get("type", "case")
    kw = request.args.get("kw", "").strip()
    if dtype == "case":
        sql = "SELECT case_no,client_name,owner,developer,status,agent_fee,invoice_amount,archive_status FROM cases WHERE deleted=0"
        if kw:
            sql += (" AND (client_name LIKE ? OR case_no LIKE ? OR owner LIKE ?"
                    " OR developer LIKE ? OR case_cause LIKE ? OR plaintiff LIKE ? OR defendant LIKE ?)")
            args = ["%" + kw + "%"] * 7
        else:
            args = []
        sql += " ORDER BY id DESC LIMIT 500"
    elif dtype == "staff":
        sql = "SELECT staff_no,name,department,position,hire_date,phone,status FROM staff WHERE deleted=0"
        args = ["%" + kw + "%"] * 3 if kw else []
        if kw:
            sql += " AND (name LIKE ? OR staff_no LIKE ? OR department LIKE ?)"
        sql += " ORDER BY id DESC LIMIT 500"
    elif dtype == "customer":
        sql = "SELECT name,credit_code,contact,phone,cust_type,industry,address FROM customers WHERE deleted=0"
        args = ["%" + kw + "%"] * 3 if kw else []
        if kw:
            sql += " AND (name LIKE ? OR contact LIKE ? OR phone LIKE ?)"
        sql += " ORDER BY id DESC LIMIT 500"
    elif dtype == "invoice":
        sql = "SELECT invoice_no,seller,amount,tax_amount,category,invoice_date,confirm FROM invoices"
        args = ["%" + kw + "%"] * 3 if kw else []
        if kw:
            sql += " WHERE invoice_no LIKE ? OR seller LIKE ? OR category LIKE ?"
        sql += " ORDER BY id DESC LIMIT 500"
    elif dtype == "finance":
        rows = db.execute(
            "SELECT f.case_id, c.case_no, c.client_name, f.recovery_date, f.item, f.amount"
            " FROM fee_recoveries f LEFT JOIN cases c ON c.id=f.case_id"
            " ORDER BY f.id DESC LIMIT 500").fetchall()
        return jsonify({"ok": True, "type": "finance", "data": [dict(r) for r in rows]})
    else:
        return jsonify({"ok": False, "msg": "未知数据类型"}), 400
    rows = db.execute(sql, args).fetchall()
    return jsonify({"ok": True, "type": dtype, "data": [dict(r) for r in rows]})


# ---------------------------------------------------------------- 运营看板（独立子页面）
def _day(v):
    """统一取 YYYY-MM-DD，兼容 'YYYY-MM-DD HH:MM:SS' 存储格式"""
    return (v or "")[:10]


def _rate(rec, fee):
    """回款率 = 回款金额 ÷ 代理费（合同额）× 100%；分母为 0 返回 None（前端显示 —）"""
    if not fee:
        return None
    return round(rec / fee * 100, 1)


def _yoy(cur, prev):
    """同比增减百分比；上年为 0 时返回 None"""
    if not prev:
        return None
    return round((cur - prev) / abs(prev) * 100, 1)


@app.route("/api/dashboard")
@login_required
def dashboard():
    """运营看板
    口径说明（区间 = [start, end]）：
      立项   created_at ∈ 区间
      结项   status 归一化后为 closed/archived/mediation_closed 且 updated_at ∈ 区间
      暂停   status 归一化后为 terminated 且 updated_at ∈ 区间
      进行中 status 归一化后为 flow 类别且 created_at ≤ end（期末在办）
      期间案件：区间内有立项、状态更新或回款记录的案件
      回款率 = 期间回款金额 ÷ 期间案件代理费合计 × 100%
    """
    db = get_db()
    today = date.today()
    start = (request.args.get("start") or "").strip() or "%04d-%02d-01" % (today.year, today.month)
    end = (request.args.get("end") or "").strip() or today.strftime("%Y-%m-%d")
    if end < start:
        start, end = end, start
    dim = (request.args.get("dim") or "owner").strip()
    if dim not in ("owner", "developer", "member"):
        dim = "owner"

    cases = [dict(r) for r in db.execute("SELECT * FROM cases WHERE deleted=0").fetchall()]

    # 回款：全量累计 / 区间 / 按年月归集
    rec_all, rec_range, rec_month, rec_year = {}, {}, {}, {}
    for r in db.execute("SELECT case_id,recovery_date,amount FROM fee_recoveries").fetchall():
        amt = float(r["amount"] or 0)
        rd = _day(r["recovery_date"])
        rec_all[r["case_id"]] = rec_all.get(r["case_id"], 0.0) + amt
        if rd:
            rec_year[rd[:4]] = rec_year.get(rd[:4], 0.0) + amt
            if start <= rd <= end:
                rec_range[r["case_id"]] = rec_range.get(r["case_id"], 0.0) + amt
                rec_month[rd[:7]] = rec_month.get(rd[:7], 0.0) + amt

    # 小组成员（dim=member 时作为人员来源）
    members_by_case = {}
    for r in db.execute("SELECT case_id,person FROM case_members").fetchall():
        lst = members_by_case.setdefault(r["case_id"], [])
        if r["person"] and r["person"] not in lst:
            lst.append(r["person"])

    def person_keys(c):
        if dim == "developer":
            return [c.get("developer") or "未分配"]
        if dim == "member":
            return members_by_case.get(c["id"]) or [c.get("owner") or "未分配"]
        return [c.get("owner") or "未分配"]

    def blank():
        return {"total": 0, "new": 0, "closed": 0, "paused": 0, "doing": 0,
                "closed_now": 0, "listed": 0, "fee": 0.0, "recovery": 0.0,
                "invoice": 0.0, "subject": 0.0}

    def bucket(store, key):
        if key not in store:
            store[key] = blank()
        return store[key]

    proj, person, summary = {}, {}, blank()
    year_new, year_closed, year_fee = {}, {}, {}
    month_new, month_closed, month_fee = {}, {}, {}
    cum_listed = 0          # 期末在册案件数（截至 end）
    cum_doing = 0           # 期末在办（进行中）
    cum_paused = 0          # 期末暂停
    cum_closed = 0          # 期末已结案
    cum_fee = 0.0           # 累计合同额（代理费，created_at ≤ end）
    cum_rec = 0.0           # 累计回款

    for c in cases:
        cid = c["id"]
        cdt = _day(c.get("created_at"))
        udt = _day(c.get("updated_at") or c.get("created_at"))
        st = c.get("status") or ""
        code = norm_case_status(st)          # 兼容旧中文状态与新英文 code
        fee = float(c.get("agent_fee") or 0)
        inv = float(c.get("invoice_amount") or 0)
        sub = float(c.get("subject_amount") or 0)
        rec = rec_range.get(cid, 0.0)
        is_closed = code in ("closed", "archived", "mediation_closed")

        in_new = bool(cdt) and start <= cdt <= end
        in_closed = is_closed and bool(udt) and start <= udt <= end
        in_paused = (code == "terminated") and bool(udt) and start <= udt <= end
        doing_now = is_flow_status(code) and bool(cdt) and cdt <= end
        listed = bool(cdt) and cdt <= end and (not is_closed or (udt and udt > end))
        closed_now = bool(cdt) and cdt <= end and is_closed and bool(udt) and udt <= end
        active = in_new or (bool(udt) and start <= udt <= end) or (cid in rec_range)

        # ---- 期末 / 累计口径（不受区间案件筛选影响）
        if bool(cdt) and cdt <= end:
            cum_fee += fee
        if listed:
            cum_listed += 1
        if doing_now:
            cum_doing += 1
        if closed_now:
            cum_closed += 1
        if code == "terminated" and bool(cdt) and cdt <= end:
            cum_paused += 1
        cum_rec += rec_all.get(cid, 0.0)

        # ---- 历史年度（不受区间影响）
        if cdt:
            y = cdt[:4]
            year_new[y] = year_new.get(y, 0) + 1
            year_fee[y] = year_fee.get(y, 0.0) + fee
            if is_closed and udt:
                year_closed[udt[:4]] = year_closed.get(udt[:4], 0) + 1
        # ---- 区间月度
        if in_new and cdt:
            m = cdt[:7]
            month_new[m] = month_new.get(m, 0) + 1
            month_fee[m] = month_fee.get(m, 0.0) + fee
        if in_closed and udt:
            month_closed[udt[:7]] = month_closed.get(udt[:7], 0) + 1

        if not active:
            continue

        # ---- 汇总
        s = summary
        s["total"] += 1
        s["fee"] += fee
        s["recovery"] += rec
        s["invoice"] += inv
        s["subject"] += sub
        s["new"] += 1 if in_new else 0
        s["closed"] += 1 if in_closed else 0
        s["paused"] += 1 if in_paused else 0
        s["doing"] += 1 if doing_now else 0
        s["closed_now"] += 1 if closed_now else 0
        s["listed"] += 1 if listed else 0

        # ---- 项目维度
        p = bucket(proj, (c.get("project") or "").strip() or "未分类项目")
        p["total"] += 1
        p["fee"] += fee
        p["recovery"] += rec
        p["invoice"] += inv
        p["subject"] += sub
        p["new"] += 1 if in_new else 0
        p["closed"] += 1 if in_closed else 0
        p["paused"] += 1 if in_paused else 0
        p["doing"] += 1 if doing_now else 0
        p["closed_now"] += 1 if closed_now else 0
        p["listed"] += 1 if listed else 0

        # ---- 人员维度
        for name in person_keys(c):
            u = bucket(person, name)
            u["total"] += 1
            u["fee"] += fee
            u["recovery"] += rec
            u["invoice"] += inv
            u["subject"] += sub
            u["new"] += 1 if in_new else 0
            u["closed"] += 1 if in_closed else 0
            u["paused"] += 1 if in_paused else 0
            u["doing"] += 1 if doing_now else 0
            u["closed_now"] += 1 if closed_now else 0
            u["listed"] += 1 if listed else 0

    def pack(store, order_key="fee"):
        out = []
        for k, v in store.items():
            d = dict(v)
            d["name"] = k
            d["rate"] = _rate(v["recovery"], v["fee"])
            d["close_rate"] = round(v["closed_now"] / v["total"] * 100, 1) if v["total"] else 0.0
            d["fee"] = round(v["fee"], 2)
            d["recovery"] = round(v["recovery"], 2)
            d["invoice"] = round(v["invoice"], 2)
            d["subject"] = round(v["subject"], 2)
            out.append(d)
        out.sort(key=lambda x: (-x[order_key], -x["total"], x["name"]))
        return out

    # ---- 历史年度对比（含同比）
    years = sorted(set(list(year_new) + list(year_closed) + list(year_fee) + list(rec_year)),
                   reverse=True)
    yearly = []
    for i, y in enumerate(years):
        prev = yearly[i - 1] if i else None
        fee_y = round(year_fee.get(y, 0.0), 2)
        rec_y = round(rec_year.get(y, 0.0), 2)
        row = {"year": y,
               "new": year_new.get(y, 0),
               "closed": year_closed.get(y, 0),
               "fee": fee_y,
               "recovery": rec_y,
               "rate": _rate(rec_y, fee_y)}
        row["yoy_fee"] = None if prev is None else _yoy(fee_y, prev["fee"])
        row["yoy_recovery"] = None if prev is None else _yoy(rec_y, prev["recovery"])
        row["yoy_new"] = None if prev is None else _yoy(row["new"], prev["new"])
        yearly.append(row)

    # ---- 区间月度对比
    months = sorted(set(list(month_new) + list(month_closed) + list(month_fee) + list(rec_month)))
    monthly = []
    for m in months:
        f = round(month_fee.get(m, 0.0), 2)
        r = round(rec_month.get(m, 0.0), 2)
        monthly.append({"month": m, "new": month_new.get(m, 0),
                        "closed": month_closed.get(m, 0),
                        "fee": f, "recovery": r, "rate": _rate(r, f)})

    s = summary
    return jsonify({"ok": True, "data": {
        "range": {"start": start, "end": end, "dim": dim,
                  "dim_label": {"owner": "主承办人", "developer": "开发人",
                                "member": "小组成员"}[dim]},
        "summary": {
            "total": s["total"], "new": s["new"], "closed": s["closed"],
            "paused": s["paused"], "doing": s["doing"],
            "listed": cum_listed, "closed_now": s["closed_now"],
            "doing_all": cum_doing, "paused_all": cum_paused, "closed_all": cum_closed,
            "fee": round(s["fee"], 2), "recovery": round(s["recovery"], 2),
            "invoice": round(s["invoice"], 2), "subject": round(s["subject"], 2),
            "rate": _rate(s["recovery"], s["fee"]),
            "new_fee": round(sum(month_fee.values()), 2),
            "cum_fee": round(cum_fee, 2),
            "cum_recovery": round(cum_rec, 2),
            "cum_rate": _rate(cum_rec, cum_fee),
        },
        "yearly": yearly,
        "monthly": monthly,
        "projects": pack(proj),
        "persons": pack(person),
    }})


@app.route("/api/meta/case-status")
@login_required
def meta_case_status():
    """案件主状态枚举元数据（含 label / color / category / 可流转 next）。"""
    return jsonify({
        "ok": True,
        "data": CASE_STATUS_META,
        "legacy_map": LEGACY_STATUS_MAP,
        "categories": {
            "flow": "正常流转",
            "result": "结果类",
            "terminated": "终止类",
            "abnormal": "异常类",
        },
    })


@app.route("/api/meta/execution-status")
@login_required
def meta_execution_status():
    """执行状态枚举元数据。"""
    return jsonify({"ok": True, "data": EXECUTION_STATUS_META})


@app.route("/api/panel")
@login_required
def panel():
    """工作台案件面板：负责人承办案件、客户集中、案件总进度、开发人客户、助理承接案件"""
    db = get_db()
    cases = db.execute("SELECT * FROM cases WHERE deleted=0").fetchall()
    total = len(cases)
    done = sum(1 for c in cases if norm_case_status(c["status"]) in ("closed", "archived", "mediation_closed"))
    progress = round(done / total * 100) if total else 0
    total_fee = sum(float(c["agent_fee"] or 0) for c in cases)
    # 负责人承办案件
    owners = {}
    for c in cases:
        o = c["owner"] or "未分配"
        owners.setdefault(o, {"count": 0, "amount": 0.0, "clients": set()})
        owners[o]["count"] += 1
        owners[o]["amount"] += float(c["agent_fee"] or 0)
        owners[o]["clients"].add(c["client_name"])
    owner_list = sorted(
        [{"name": k, "count": v["count"], "amount": round(v["amount"], 2),
          "client_count": len(v["clients"])} for k, v in owners.items()],
        key=lambda x: -x["count"])
    # 开发人客户统计
    devs = {}
    for c in cases:
        d = c["developer"] or "未分配"
        devs.setdefault(d, {"clients": set(), "amount": 0.0})
        devs[d]["clients"].add(c["client_name"])
        devs[d]["amount"] += float(c["agent_fee"] or 0)
    month = date.today().strftime("%Y-%m")
    month_new = db.execute("SELECT COUNT(*) n FROM customers WHERE created_at LIKE ?",
                           (month + "%",)).fetchone()["n"]
    # 助理人员承接案件数（members 中的人员）
    assistants = {}
    for c in cases:
        for m in (c["members"] or "").split(","):
            m = m.strip()
            if m:
                assistants[m] = assistants.get(m, 0) + 1
    # 甘特图：各案件进度（按费用收回 / 代理费 估算；无代理费按状态折算）
    gantt = []
    for c in cases:
        rec = db.execute("SELECT COALESCE(SUM(amount),0) s FROM fee_recoveries WHERE case_id=?",
                         (c["id"],)).fetchone()["s"]
        base = float(c["agent_fee"] or 0)
        if base > 0:
            p = min(100, max(0, round(rec / base * 100)))
        else:
            p = 100 if c["status"] == "已结案" else (50 if c["status"] == "进行中" else 30)
        gantt.append({"case_no": c["case_no"], "client_name": c["client_name"],
                      "owner": c["owner"], "progress": p, "status": c["status"]})
    return jsonify({"ok": True, "data": {
        "total": total, "done": done, "progress": progress,
        "total_fee": round(total_fee, 2),
        "customer_total": len({c["client_name"] for c in cases}),
        "owners": owner_list,
        "month_new_customers": month_new,
        "assistants": [{"name": k, "count": v} for k, v in
                       sorted(assistants.items(), key=lambda x: -x[1])],
        "dev_customers": [{"name": k, "client_count": len(v["clients"]),
                           "amount": round(v["amount"], 2)}
                          for k, v in devs.items()],
        "gantt": gantt,
    }})


@app.route("/api/overview")
@login_required
def overview():
    db = get_db()
    c = db.execute("SELECT COUNT(*) n, COALESCE(SUM(agent_fee),0) fee,"
                   " COALESCE(SUM(invoice_amount),0) inv FROM cases WHERE deleted=0").fetchone()
    s = db.execute("SELECT COUNT(*) n FROM settlements").fetchone()
    inv = db.execute("SELECT COUNT(*) n FROM invoices WHERE confirm=0").fetchone()
    recent = db.execute("SELECT case_no,client_name,status,updated_by,updated_at"
                        " FROM cases WHERE deleted=0 ORDER BY updated_at DESC LIMIT 5").fetchall()
    return jsonify({"ok": True, "data": {
        "case_count": c["n"], "agent_fee_total": round(c["fee"], 2),
        "invoice_total": round(c["inv"], 2), "settle_count": s["n"],
        "invoice_pending": inv["n"], "recent": [dict(r) for r in recent]}})


# ---------------------------------------------------------------- 启动
def _migrate_legacy_ai_key():
    """遗留明文 ai_api_key 不再入库：若有值则先迁移到本地配置文件，随后删除该数据库行。
    无论是否有值，都确保 settings 表中不再保留 ai_api_key 行（新架构密钥仅来自环境变量/配置文件）。"""
    try:
        legacy = get_setting("ai_api_key", "")
        if legacy:
            pid = get_setting("ai_provider", "bailian")
            if pid not in ai_providers.PROVIDER_PRESETS:
                pid = "bailian"
            kind = ai_providers.PROVIDER_PRESETS[pid]["auth_kind"]
            if kind != "tc3":
                cur = ai_providers.load_config_file()
                prov = cur.get("providers", {})
                prov[pid] = prov.get(pid, {})
                prov[pid]["api_key"] = legacy
                ai_providers.save_config_file({"active": pid, "providers": prov})
                print("[AI] 已将遗留明文 ai_api_key 迁移至本地配置文件")
        # 无论如何都删除数据库中的 ai_api_key 行（新架构密钥不入库）
        db = get_db()
        db.execute("DELETE FROM settings WHERE key='ai_api_key'")
        db.commit()
        print("[AI] 已确保 settings 表中不存在 ai_api_key 明文行")
    except Exception as e:
        print("[AI] 遗留密钥迁移跳过:", e)


if __name__ == "__main__":
    init_db()
    # 启动时加载角色-权限矩阵配置（data/role_config.json），覆盖代码内默认值
    load_role_config()
    # 启动数据自动备份调度线程（守护线程，按配置周期执行）
    try:
        threading.Thread(target=backup_scheduler_loop, daemon=True).start()
        print("[备份] 自动备份调度线程已启动")
    except Exception as e:
        print("[备份] 调度线程启动失败:", e)
    # 注册 AI 密钥读取回调（仅用于非密钥项兜底），并迁移遗留的数据库明文密钥
    ai_providers.set_db_getter(get_setting)
    with app.app_context():
        _migrate_legacy_ai_key()
    db = sqlite3.connect(DB_PATH)
    pair = db.execute("SELECT value FROM settings WHERE key='pair_code'").fetchone()[0]
    db.close()
    # 启动局域网发现服务（UDP 广播响应）
    try:
        import discovery
        discovery.start_responder(pair, SERVER_PORT)
    except Exception as e:
        print("[配网] 发现服务启动失败:", e)
    print("=" * 56)
    print(f"  律师工作台  ({ORG_FULL_NAME})")
    print("  本机访问:   http://127.0.0.1:%d" % SERVER_PORT)
    print("  配网数字:   %s  （其他电脑运行「连接服务器」并输入此数字即可连接）" % pair)
    print("  默认账号:   admin / 123456  （请登录后立即修改）")
    print("  数据目录:   %s" % DATA_DIR)
    print("=" * 56)
    try:
        webbrowser.open("http://127.0.0.1:%d" % SERVER_PORT)
    except Exception:
        pass
    app.run(host="0.0.0.0", port=SERVER_PORT, threaded=True)
